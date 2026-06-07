import copy
from pathlib import Path
from typing import Iterable, List

from docx import Document
from docx.enum.text import WD_BREAK

from docx_layout import strip_document_list_numbering
from docx_numbering import generation_order, section_docx_path


COMBINED_SECTION_FILENAME = "现状调查与评价.docx"
PROJECT_AREA_OVERVIEW_FILENAME = "project_area_overview.docx"
NOISE_SECTION_FILENAME = "noise_section.docx"
SURFACE_WATER_SECTION_FILENAME = "surface_water_section.docx"


def build_combined_section_docx(output_dir: Path) -> Path:
    output_dir = Path(output_dir)
    combined_path = output_dir / COMBINED_SECTION_FILENAME
    existing = list_existing_section_docx_paths(output_dir)

    if not existing:
        raise FileNotFoundError("未找到可合并的项目区域概况、声环境或地表水 Word 章节")

    if len(existing) == 1:
        doc = Document(str(existing[0]))
        strip_document_list_numbering(doc)
        doc.save(combined_path)
        return combined_path

    combine_docx_files(existing, combined_path)
    return combined_path


def list_existing_section_docx_paths(output_dir: Path) -> List[Path]:
    output_dir = Path(output_dir)
    return [
        section_docx_path(output_dir, section_key)
        for section_key in generation_order()
        if section_docx_path(output_dir, section_key).exists()
    ]


def combine_docx_files(paths: List[Path], output_path: Path) -> None:
    first_doc = Document(str(paths[0]))
    for path in paths[1:]:
        add_page_break(first_doc)
        append_body_elements(first_doc, Document(str(path)))
    strip_document_list_numbering(first_doc)
    first_doc.save(str(output_path))


def append_docx(first_path: Path, second_path: Path, output_path: Path) -> None:
    combine_docx_files([Path(first_path), Path(second_path)], Path(output_path))


def add_page_break(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    paragraph.add_run().add_break(WD_BREAK.PAGE)


def append_body_elements(target_doc: Document, source_doc: Document) -> None:
    target_body = target_doc.element.body
    target_section_properties = target_body.sectPr
    if target_section_properties is not None:
        target_body.remove(target_section_properties)

    for element in iter_body_content_without_section_properties(source_doc):
        target_body.append(copy.deepcopy(element))

    source_section_properties = source_doc.element.body.sectPr
    if source_section_properties is not None:
        target_body.append(copy.deepcopy(source_section_properties))
    elif target_section_properties is not None:
        target_body.append(target_section_properties)


def iter_body_content_without_section_properties(doc: Document) -> Iterable[object]:
    for element in doc.element.body.iterchildren():
        if element.tag.endswith("}sectPr"):
            continue
        yield element
