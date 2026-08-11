"""Shared Word layout helpers aligned with the reference status-survey report."""

import json
import os
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.text.paragraph import Paragraph

BASE_DIR = Path(__file__).resolve().parent
REFERENCE_LAYOUT_PATH = Path(
    __import__("os").getenv(
        "EIA_REFERENCE_LAYOUT_CONFIG",
        BASE_DIR / "config" / "reference_layout.json",
    )
)

BLACK = RGBColor(0, 0, 0)
MARGIN_CM = 2.5
BODY_FONT_ASCII = "Times New Roman"
BODY_FONT_EAST_ASIA = "宋体"
BODY_SIZE_PT = 12
TABLE_SIZE_PT = 10.5
FIRST_LINE_INDENT_PT = 24
BODY_LINE_SPACING_PT = 24
CHAPTER_TITLE_SIZE_PT = 14
CHAPTER_TITLE_LINE_SPACING_PT = 30
SECTION_HEADING_SIZE_PT = 12
TABLE_CAPTION_SIZE_PT = 12
DOCUMENT_TITLE_SIZE_PT = 18
TABLE_STYLE = "Normal Table"
HEADER_FILL: Optional[str] = None
CELL_MARGIN_DXA = 80
WIDE_TABLE_MIN_COLUMNS = int(os.getenv("EIA_WIDE_TABLE_MIN_COLUMNS", "9"))
EIA_HEADING_STYLE_NAMES = {
    1: "EIAHeading1",
    2: "EIAHeading2",
    3: "EIAHeading3",
}


def load_reference_layout() -> Dict[str, Any]:
    path = Path(REFERENCE_LAYOUT_PATH)
    if not path.exists():
        return {}
    return json.loads(path.read_text(encoding="utf-8"))


def reference_docx_path() -> Optional[Path]:
    layout = load_reference_layout()
    source = layout.get("source_docx")
    if not source:
        return None
    path = Path(str(source))
    return path if path.exists() else None


def create_section_document() -> Document:
    # Blank document only: reference templates carry multilevel list definitions that
    # break after section merge and show reversed chapter numbers (3/2/1) in WPS.
    return Document()


def heading_style_candidates(level: int) -> List[str]:
    name = EIA_HEADING_STYLE_NAMES.get(level)
    return [name] if name else []


def resolve_heading_style(doc: Document, level: int):
    for name in heading_style_candidates(level):
        try:
            return doc.styles[name]
        except KeyError:
            continue
    fallback_name = f"EIAHeading{level}"
    try:
        return doc.styles[fallback_name]
    except KeyError:
        style = doc.styles.add_style(fallback_name, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = doc.styles["Normal"]
        _set_style_outline_level(style, level - 1)
        return style


def setup_document(doc: Document) -> None:
    _apply_section_layout(doc.sections[0], portrait=True)
    _configure_base_styles(doc)
    _ensure_heading_styles(doc)


def add_landscape_section(doc: Document) -> None:
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    _apply_section_layout(section, portrait=False)


def add_portrait_section(doc: Document) -> None:
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    _apply_section_layout(section, portrait=True)


def table_needs_landscape(headers: List[str]) -> bool:
    return len(headers or []) >= WIDE_TABLE_MIN_COLUMNS


def add_document_title(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.line_spacing = Pt(BODY_LINE_SPACING_PT)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    run = paragraph.add_run(text)
    apply_run_style(run, BODY_FONT_ASCII, BODY_FONT_EAST_ASIA, DOCUMENT_TITLE_SIZE_PT)


def add_chapter_title(doc: Document, text: str) -> None:
    _add_heading_paragraph(
        doc,
        text,
        level=1,
        size_pt=CHAPTER_TITLE_SIZE_PT,
        line_spacing_pt=CHAPTER_TITLE_LINE_SPACING_PT,
    )


def add_section_heading(doc: Document, text: str) -> None:
    _add_heading_paragraph(
        doc,
        text,
        level=2,
        size_pt=SECTION_HEADING_SIZE_PT,
        line_spacing_pt=BODY_LINE_SPACING_PT,
    )


def add_level3_heading(doc: Document, text: str) -> None:
    _add_heading_paragraph(
        doc,
        text,
        level=3,
        size_pt=BODY_SIZE_PT,
        line_spacing_pt=BODY_LINE_SPACING_PT,
        first_line_indent_pt=FIRST_LINE_INDENT_PT,
    )


def add_body_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(str(text or "-"))
    paragraph.paragraph_format.first_line_indent = Pt(FIRST_LINE_INDENT_PT)
    paragraph.paragraph_format.line_spacing = Pt(BODY_LINE_SPACING_PT)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    for run in paragraph.runs:
        apply_run_style(run, BODY_FONT_ASCII, BODY_FONT_EAST_ASIA, BODY_SIZE_PT)


def add_numbered_title(doc: Document, text: str) -> None:
    add_level3_heading(doc, text)


def add_subtitle(doc: Document, text: str) -> None:
    add_body_paragraph(doc, text)


def add_caption(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.line_spacing = Pt(BODY_LINE_SPACING_PT)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    run = paragraph.add_run(text)
    apply_run_style(run, BODY_FONT_ASCII, BODY_FONT_EAST_ASIA, TABLE_CAPTION_SIZE_PT, bold=False)


def add_table(doc: Document, headers: List[str], rows: List[Dict[str, Any]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = TABLE_STYLE
    table.autofit = True
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        header_cells[index].text = str(header)
        if HEADER_FILL:
            shade_cell(header_cells[index], HEADER_FILL)
        set_cell_text_style(header_cells[index], bold=True)

    for row in rows:
        cells = table.add_row().cells
        for index, header in enumerate(headers):
            cells[index].text = str(row.get(header, "-"))
            set_cell_text_style(cells[index], bold=False)

    finalize_table(table, header_row_count=1)
    doc.add_paragraph()


def add_formula(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.line_spacing = Pt(BODY_LINE_SPACING_PT)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    run = paragraph.add_run(text)
    apply_run_style(run, "Times New Roman", BODY_FONT_EAST_ASIA, BODY_SIZE_PT)


def finalize_table(table: Any, header_row_count: int = 0) -> None:
    apply_reference_table_borders(table)
    style_table(table, header_row_count=header_row_count)


def set_repeat_table_header(row: Any) -> None:
    properties = row._tr.get_or_add_trPr()
    marker = properties.find(qn("w:tblHeader"))
    if marker is None:
        marker = OxmlElement("w:tblHeader")
        marker.set(qn("w:val"), "true")
        properties.append(marker)


def set_table_row_cant_split(row: Any) -> None:
    properties = row._tr.get_or_add_trPr()
    marker = properties.find(qn("w:cantSplit"))
    if marker is None:
        properties.append(OxmlElement("w:cantSplit"))


def style_table(table: Any, header_row_count: int = 0) -> None:
    for row_index, row in enumerate(table.rows):
        is_header = row_index < header_row_count
        set_table_row_cant_split(row)
        if is_header:
            set_repeat_table_header(row)
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margin(cell)
            set_cell_text_style(cell, bold=is_header)


def apply_reference_table_borders(table: Any) -> None:
    tbl_pr = table._tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        table._tbl.insert(0, tbl_pr)

    existing = tbl_pr.find(qn("w:tblBorders"))
    if existing is not None:
        tbl_pr.remove(existing)

    borders = OxmlElement("w:tblBorders")
    border_specs = {
        "top": ("thinThickSmallGap", "12"),
        "left": ("thinThickSmallGap", "12"),
        "bottom": ("thickThinSmallGap", "12"),
        "right": ("thickThinSmallGap", "12"),
        "insideH": ("single", "6"),
        "insideV": ("single", "6"),
    }
    for edge, (value, size) in border_specs.items():
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), value)
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), "auto")
        element.set(qn("w:space"), "0")
        borders.append(element)
    tbl_pr.append(borders)


def shade_cell(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_text_style(cell: Any, bold: bool = False) -> None:
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.first_line_indent = Pt(0)
        for run in paragraph.runs:
            apply_run_style(
                run,
                BODY_FONT_ASCII,
                BODY_FONT_EAST_ASIA,
                TABLE_SIZE_PT,
                bold=bold,
            )


def set_cell_margin(cell: Any, margin: int = CELL_MARGIN_DXA) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side in ("top", "left", "bottom", "right"):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(margin))
        node.set(qn("w:type"), "dxa")


def apply_run_style(
    run: Any,
    font_ascii: str,
    font_east_asia: str,
    size_pt: float,
    bold: bool = False,
) -> None:
    run.font.name = font_ascii
    run._element.rPr.rFonts.set(qn("w:ascii"), font_ascii)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font_ascii)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_east_asia)
    run.font.size = Pt(size_pt)
    run.font.color.rgb = BLACK
    run.bold = bold


def _add_heading_paragraph(
    doc: Document,
    text: str,
    *,
    level: int,
    size_pt: float,
    line_spacing_pt: float,
    first_line_indent_pt: float = 0,
) -> Paragraph:
    style = resolve_heading_style(doc, level)
    paragraph = doc.add_paragraph(str(text or "").strip(), style=style)
    paragraph.paragraph_format.first_line_indent = Pt(first_line_indent_pt)
    paragraph.paragraph_format.line_spacing = Pt(line_spacing_pt)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    _clear_paragraph_numbering(paragraph)
    _set_paragraph_outline_level(paragraph, level - 1)
    for run in paragraph.runs:
        apply_run_style(run, BODY_FONT_ASCII, BODY_FONT_EAST_ASIA, size_pt)
    return paragraph


def _ensure_heading_styles(doc: Document) -> None:
    for level, size_pt, line_spacing_pt in (
        (1, CHAPTER_TITLE_SIZE_PT, CHAPTER_TITLE_LINE_SPACING_PT),
        (2, SECTION_HEADING_SIZE_PT, BODY_LINE_SPACING_PT),
        (3, BODY_SIZE_PT, BODY_LINE_SPACING_PT),
    ):
        style = resolve_heading_style(doc, level)
        _configure_style_fonts(style, size_pt)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        style.paragraph_format.line_spacing = Pt(line_spacing_pt)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)
        if level == 1:
            style.paragraph_format.first_line_indent = Pt(0)
        elif level == 2:
            style.paragraph_format.first_line_indent = Pt(0)
        else:
            style.paragraph_format.first_line_indent = Pt(FIRST_LINE_INDENT_PT)
        _set_style_outline_level(style, level - 1)
        _clear_style_numbering(style)


def _set_style_outline_level(style: Any, outline_level: int) -> None:
    pPr = style.element.get_or_add_pPr()
    existing = pPr.find(qn("w:outlineLvl"))
    if existing is not None:
        pPr.remove(existing)
    outline = OxmlElement("w:outlineLvl")
    outline.set(qn("w:val"), str(outline_level))
    pPr.append(outline)


def _set_paragraph_outline_level(paragraph: Paragraph, outline_level: int) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    existing = pPr.find(qn("w:outlineLvl"))
    if existing is not None:
        pPr.remove(existing)
    outline = OxmlElement("w:outlineLvl")
    outline.set(qn("w:val"), str(outline_level))
    pPr.append(outline)


def _clear_style_numbering(style: Any) -> None:
    pPr = style.element.get_or_add_pPr()
    numPr = pPr.find(qn("w:numPr"))
    if numPr is not None:
        pPr.remove(numPr)


def strip_document_list_numbering(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        _clear_paragraph_numbering(paragraph)


def finalize_section_document(doc: Document) -> None:
    strip_document_list_numbering(doc)


def _clear_paragraph_numbering(paragraph: Paragraph) -> None:
    pPr = paragraph._p.pPr
    if pPr is None:
        return
    numPr = pPr.find(qn("w:numPr"))
    if numPr is not None:
        pPr.remove(numPr)


def _configure_style_fonts(style: Any, size_pt: float) -> None:
    style.font.name = BODY_FONT_ASCII
    style.font.size = Pt(size_pt)
    style.font.color.rgb = BLACK
    style.font.bold = False
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), BODY_FONT_ASCII)
    rFonts.set(qn("w:hAnsi"), BODY_FONT_ASCII)
    rFonts.set(qn("w:eastAsia"), BODY_FONT_EAST_ASIA)


def _configure_base_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT_ASCII
    normal._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT_ASCII)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT_ASCII)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT_EAST_ASIA)
    normal.font.size = Pt(BODY_SIZE_PT)
    normal.font.color.rgb = BLACK
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    normal.paragraph_format.line_spacing = Pt(BODY_LINE_SPACING_PT)
    normal.paragraph_format.first_line_indent = Pt(FIRST_LINE_INDENT_PT)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.space_before = Pt(0)


def _apply_section_layout(section: Any, portrait: bool) -> None:
    section.top_margin = Cm(MARGIN_CM)
    section.bottom_margin = Cm(MARGIN_CM)
    section.left_margin = Cm(MARGIN_CM)
    section.right_margin = Cm(MARGIN_CM)
    if portrait:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        return
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
