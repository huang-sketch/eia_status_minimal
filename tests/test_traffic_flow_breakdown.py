from docx import Document
from openpyxl import Workbook

from noise_section_generator import (
    add_vehicle_flow_fields,
    build_road_vehicle_result_table_header,
    canonical_road_flow_text,
    result_headers,
    road_flow_breakdown_from_row,
)
from xlsx_monitoring_parser import parse_noise_sheet


def test_xlsx_vehicle_columns_are_mapped_by_name_not_source_order():
    workbook = Workbook()
    sheet = workbook.active
    sheet.append([
        "\u6d4b\u70b9\u7f16\u53f7",
        "\u6d4b\u70b9\u540d\u79f0",
        "\u6d4b\u70b9\u6869\u53f7",
        "\u6d4b\u70b9\u4f4d\u7f6e",
        "\u6d4b\u70b9\u697c\u5c42",
        "\u76d1\u6d4b\u65f6\u957f",
        "\u566a\u58f0\u6267\u884c\u6807\u51c6",
        "\u76d1\u6d4b\u65f6\u95f4",
        "",
        "",
        "\u76d1\u6d4b\u6570\u636e",
        "\u566a\u58f0\u6267\u884c\u6807\u51c6",
        "\u8f66\u6d41\u91cf\u6570\u636e",
        "",
        "",
    ])
    sheet.append([
        "", "", "", "", "", "", "",
        "\u65e5\u671f", "\u65f6\u95f4", "\u65f6\u6bb5", "", "",
        "\u5c0f\u578b\u8f66", "\u4e2d\u578b\u8f66", "\u5927\u578b\u8f66",
    ])
    sheet.append([
        "NJ1", "\u6d4b\u8bd5\u70b9", "K1+000", "\u9996\u6392", "2\u5c42",
        "20min", "4a\u7c7b", "2026.6.1", "10:00~10:20", "\u663c\u95f4",
        58, 70, 100, 20, 10,
    ])
    sheet.append([
        "", "", "", "", "", "", "", "2026.6.1", "22:00~22:20",
        "\u591c\u95f4", 48, 55, 80, 15, 5,
    ])



    records, errors, _warnings = parse_noise_sheet(sheet)

    assert errors == []
    assert records[0]["traffic_flow"] == "10/20/100"
    assert records[0]["traffic_flow_large"] == "10"
    assert records[0]["traffic_flow_medium"] == "20"
    assert records[0]["traffic_flow_small"] == "100"


def test_docx_flattened_vehicle_subheaders_map_to_large_medium_small():
    row = {
        "\u8f66\u6d41\u91cf_\u5c0f\u578b\u8f66": "48",
        "\u8f66\u6d41\u91cf_\u5927\u578b\u8f66": "42",
        "\u8f66\u6d41\u91cf_\u4e2d\u578b\u8f66": "24",
    }

    breakdown = road_flow_breakdown_from_row(row, "/")

    assert breakdown == {"large": "42", "medium": "24", "small": "48"}
    assert canonical_road_flow_text(breakdown) == "42/24/48"


def test_vehicle_fields_keep_missing_type_as_slash():
    output = {}
    pairs = [{
        "day": {
            "traffic_flow": "42//48",
            "traffic_flow_large": "42",
            "traffic_flow_medium": "/",
            "traffic_flow_small": "48",
        }
    }]

    add_vehicle_flow_fields(output, pairs)

    assert output["traffic_flow_day1_day_large"] == "42"
    assert output["traffic_flow_day1_day_medium"] == "/"
    assert output["traffic_flow_day1_day_small"] == "48"
    assert output["traffic_flow_day2_night_large"] == "/"


def test_road_vehicle_word_header_has_25_columns_and_four_rows():
    document = Document()
    table = document.add_table(rows=4, cols=25)

    build_road_vehicle_result_table_header(
        table, "\u8f66\u6d41\u91cf\uff08\u8f86/20min\uff09"
    )

    assert len(table.columns) == 25
    assert len(table.rows) == 4
    assert table.cell(0, 13).text == "\u8f66\u6d41\u91cf\uff08\u8f86/20min\uff09"
    assert table.cell(1, 13).text == "\u7b2c\u4e00\u5929"
    assert table.cell(1, 19).text == "\u7b2c\u4e8c\u5929"
    assert table.cell(2, 13).text == "\u663c"
    assert table.cell(2, 16).text == "\u591c"
    assert [table.cell(3, column).text for column in (13, 14, 15)] == [
        "\u5927", "\u4e2d", "\u5c0f"
    ]
    assert len(result_headers(include_flow=True, vehicle_breakdown=True)) == 25

def test_finalize_table_marks_all_four_header_rows_for_page_repetition():
    from docx.oxml.ns import qn
    from docx_layout import finalize_table

    document = Document()
    table = document.add_table(rows=5, cols=25)
    finalize_table(table, header_row_count=4)

    flags = [
        row._tr.get_or_add_trPr().find(qn("w:tblHeader")) is not None
        for row in table.rows
    ]
    assert flags == [True, True, True, True, False]
    cant_split_flags = [
        row._tr.get_or_add_trPr().find(qn("w:cantSplit")) is not None
        for row in table.rows
    ]
    assert cant_split_flags == [True, True, True, True, True]

