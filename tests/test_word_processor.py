from docx import Document

from word_processor import load_docx_chunks


def test_docx_chunks_keep_title_context_and_merged_cells(tmp_path):
    path = tmp_path / "anonymized.docx"
    doc = Document()
    doc.add_paragraph("Example monitoring results")
    table = doc.add_table(rows=2, cols=3)
    table.cell(0, 0).merge(table.cell(0, 1))
    table.cell(0, 0).text = "Location"
    table.cell(0, 2).text = "Result"
    table.cell(1, 0).text = "WJ1"
    table.cell(1, 1).text = "pH"
    table.cell(1, 2).text = "7.5"
    doc.save(path)

    chunks = load_docx_chunks(path)
    table_chunk = next(item for item in chunks if item["kind"] == "table")

    assert table_chunk["metadata"]["table_title"] == "Example monitoring results"
    assert table_chunk["metadata"]["row_count"] == 2
    assert table_chunk["metadata"]["col_counts"] == [3, 3]
    assert "row 1:" in table_chunk["text"]
    assert "col 1: Location" in table_chunk["text"]
    assert "col 2:" in table_chunk["text"]
    assert "col 3: Result" in table_chunk["text"]
