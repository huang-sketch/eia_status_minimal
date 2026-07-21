import json
from pathlib import Path

from docx import Document

from noise_section_generator import parse_table_rows as parse_noise_rows
from section_docx_combiner import build_combined_section_docx
from docx_numbering import DocxNumbering
from surface_water_pipeline import (
    evaluate_compliance,
    parse_standard_config,
    parse_surface_water_records,
)
from surface_water_section_generator import (
    build_compliance_table,
    build_conclusion,
    build_docx,
    build_monitor_points_table,
    build_monitor_results_table,
    build_rule_texts,
    detect_evaluated_factors,
    detect_factors,
    register_surface_tables,
)
from table_schema_mapper import resolve_table_schema
from word_processor import load_docx_chunks


FIXTURE_PATH = Path(__file__).parent / "fixtures" / "anonymized_eia_case.json"


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    for index, header in enumerate(headers):
        table.cell(0, index).text = header
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = value


def build_input_docx(path, payload):
    doc = Document()
    doc.add_paragraph(payload["title"])
    add_table(doc, payload["headers"], payload["rows"])
    doc.save(path)


def stable_record(record):
    return {
        "point_code": record["point_code"],
        "sample_date": record["sample_date"],
        "factor": record["factor"],
        "value": record["value"],
        "unit": record["unit"],
    }


def test_anonymized_plan_report_to_json_and_docx(tmp_path, monkeypatch):
    fixture = json.loads(FIXTURE_PATH.read_text(encoding="utf-8"))
    expected = fixture["expected"]
    input_dir = tmp_path / "input"
    output_dir = tmp_path / "output"
    input_dir.mkdir()
    output_dir.mkdir()
    plan_path = input_dir / "plan.docx"
    report_path = input_dir / "report.docx"
    build_input_docx(plan_path, fixture["plan"])
    noise_path = input_dir / "noise_report.docx"
    build_input_docx(report_path, fixture["report"])
    monkeypatch.setenv("EIA_OUTPUT_DIR", str(output_dir))
    build_input_docx(noise_path, expected["noise_report"])

    standard_config = parse_standard_config(plan_path)
    noise_chunk = next(
        item for item in load_docx_chunks(noise_path) if item["kind"] == "table"
    )
    noise_rows = parse_noise_rows(noise_chunk["text"])
    noise_schema = resolve_table_schema(
        "noise_result",
        noise_rows[0],
        sample_rows=noise_rows[1:],
        output_dir=output_dir,
        source_table=noise_chunk["chunk_id"],
        enable_llm=False,
    )
    assert noise_schema["validation"]["valid"] is True

    records = parse_surface_water_records(report_path)
    compliance = evaluate_compliance(records, standard_config)

    assert standard_config["points"]["WJ1"]["river_name"] == expected["river_name"]
    assert standard_config["points"]["WJ1"]["standard_class"] == expected["standard_class"]
    assert len(records) == expected["record_count"]
    assert [record["factor"] for record in records] == expected["factors"]
    assert [stable_record(record) for record in records] == [
        {
            "point_code": "WJ1",
            "sample_date": "2026-01-01",
            "factor": "\u6c34\u6e29",
            "value": "20",
            "unit": "\u2103",
        },
        {
            "point_code": "WJ1",
            "sample_date": "2026-01-01",
            "factor": "pH\u503c",
            "value": "7.5",
            "unit": "\u65e0\u91cf\u7eb2",
        },
        {
            "point_code": "WJ1",
            "sample_date": "2026-01-01",
            "factor": "\u6eb6\u89e3\u6c27",
            "value": "6.0",
            "unit": "mg/L",
        },
    ]
    compliance_by_factor = {item["factor"]: item["is_compliant"] for item in compliance}
    assert compliance_by_factor["pH\u503c"] is expected["compliance"]["pH\u503c"]
    assert compliance_by_factor["\u6eb6\u89e3\u6c27"] is expected["compliance"]["\u6eb6\u89e3\u6c27"]

    factors = detect_factors(records)
    evaluated_factors = detect_evaluated_factors(compliance)
    table1 = build_monitor_points_table(records, standard_config, factors)
    table2 = build_monitor_results_table(records, standard_config, factors)
    table3 = build_compliance_table(compliance, standard_config, evaluated_factors)
    numbering = DocxNumbering(tmp_path / "numbering.json")
    numbering.reset()
    numbering.begin_section("surface_water", "\u5730\u8868\u6c34\u73af\u5883\u73b0\u72b6\u8c03\u67e5\u4e0e\u8bc4\u4ef7")
    register_surface_tables(numbering, table1, table2, table3)
    texts = build_rule_texts(
        table1,
        table2,
        factors,
        compliance,
        build_conclusion(compliance),
        numbering,
    )
    doc = build_docx(table1, table2, table3, factors, evaluated_factors, texts, numbering)
    output_path = output_dir / "surface_water_section.docx"
    doc.save(output_path)

    combined_path = build_combined_section_docx(output_dir)
    assert combined_path.exists()
    rendered = Document(combined_path)
    assert len(rendered.tables) == 3
    paragraph_text = "\n".join(paragraph.text for paragraph in rendered.paragraphs)
    assert "\u5730\u8868\u6c34\u73af\u5883\u73b0\u72b6\u8c03\u67e5\u4e0e\u8bc4\u4ef7" in paragraph_text
    assert any("WJ1" in cell.text for table in rendered.tables for row in table.rows for cell in row.cells)
