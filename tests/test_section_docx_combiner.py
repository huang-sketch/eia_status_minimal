from docx import Document
from docx.enum.section import WD_ORIENT

from section_docx_combiner import combine_docx_files


def write_section(path, text, orientation):
    document = Document()
    section = document.sections[0]
    if orientation == WD_ORIENT.LANDSCAPE:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
    document.add_paragraph(text)
    document.save(path)


def test_combiner_preserves_each_source_section_orientation(tmp_path):
    portrait_before = tmp_path / "before.docx"
    landscape = tmp_path / "wide.docx"
    portrait_after = tmp_path / "after.docx"
    combined = tmp_path / "combined.docx"
    write_section(portrait_before, "before", WD_ORIENT.PORTRAIT)
    write_section(landscape, "wide", WD_ORIENT.LANDSCAPE)
    write_section(portrait_after, "after", WD_ORIENT.PORTRAIT)

    combine_docx_files(
        [portrait_before, landscape, portrait_after],
        combined,
    )

    result = Document(combined)
    assert [section.orientation for section in result.sections] == [
        WD_ORIENT.PORTRAIT,
        WD_ORIENT.LANDSCAPE,
        WD_ORIENT.PORTRAIT,
    ]

