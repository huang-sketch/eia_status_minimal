import json
import os
import re
import time
from collections import OrderedDict
from datetime import date
from decimal import Decimal, ROUND_HALF_UP
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Tuple

from dotenv import load_dotenv
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from docx_layout import (
    HEADER_FILL,
    TABLE_STYLE,
    add_body_paragraph,
    add_caption,
    add_chapter_title,
    add_landscape_section,
    add_level3_heading,
    add_section_heading,
    add_table,
    create_section_document,
    finalize_section_document,
    finalize_table,
    setup_document,
    set_cell_text_style,
    shade_cell,
)
from docx_numbering import DocxNumbering, load_numbering
from formal_text_skill import build_noise_formal_text_validation, write_formal_text_validation
from llm_client import (
    LLM_TEXT_BATCH_PAUSE_SECONDS,
    LlmProfile,
    build_rule_text_fallback_validation,
    chat_completion_json_object,
    chat_completion_json_object_with_recovery,
    is_network_error,
)
from text_polish_utils import build_text_polish_prompt, load_text_polish_guidance
from table_schema_mapper import apply_header_mapping, record_data_validation, resolve_table_schema
from word_processor import load_docx_chunks

load_dotenv()

BASE_DIR = Path(__file__).resolve().parent
INPUT_DIR = Path(os.getenv("EIA_INPUT_DIR", "input"))
OUTPUT_DIR = Path(os.getenv("EIA_OUTPUT_DIR", "output"))
DEBUG_DIR = OUTPUT_DIR / "debug_tables"
NOISE_STANDARD_LIMITS_PATH = Path(
    os.getenv("EIA_NOISE_STANDARD_LIMITS", BASE_DIR / "config" / "noise_standard_limits.json")
)
OUTPUT_DOCX = OUTPUT_DIR / "noise_section.docx"
ENABLE_LLM_TEXT_POLISH = os.getenv("ENABLE_LLM_TEXT_POLISH", "false").lower() == "true"
TEXT_POLISH_GUIDANCE_PATH = Path(
    os.getenv("EIA_NOISE_TEXT_POLISH_GUIDANCE", BASE_DIR / "config" / "noise_text_polish_guidance.json")
)

POINT_KEY = "点位"
TIME_KEY = "监测时间"
FLOW_KEY = "车流量（辆/20min）"
ROW_POSITION_MARKER_PATTERN = r"(?:首排|第?[一二三四五六七八九十]+排)"
ROAD_VEHICLE_KEYS = ("large", "medium", "small")
ROAD_VEHICLE_LABELS = {"large": "\u5927", "medium": "\u4e2d", "small": "\u5c0f"}
LEGACY_TRAFFIC_FLOW_FIELDS = (
    "traffic_flow_day1_day", "traffic_flow_day1_night",
    "traffic_flow_day2_day", "traffic_flow_day2_night",
)

DATE_TUPLE = Tuple[int, int, int]
LAB_NAME_RE = re.compile(
    r"^[\u4e00-\u9fffA-Za-z0-9（）()·\-]{4,80}(?:检测中心|环境监测中心|环境检测中心)$"
)
DATE_TOKEN_RE = re.compile(r"(\d{4})[-/.年](\d{1,2})[-/.月](\d{1,2})")
DATE_RANGE_RE = re.compile(
    r"(\d{4})[-/.](\d{1,2})[-/.](\d{1,2})\s*[-~–—至]\s*(\d{4})[-/.](\d{1,2})[-/.](\d{1,2})"
)

_NOISE_LIMITS_CACHE: Optional[Dict[str, Any]] = None

NOISE_TABLE_MONITOR_POINTS = "noise_monitor_points"
NOISE_TABLE_SENSITIVE = "noise_sensitive_results"
NOISE_TABLE_INDOOR = "noise_indoor_results"
NOISE_TABLE_ATTENUATION = "noise_attenuation"


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    DEBUG_DIR.mkdir(parents=True, exist_ok=True)

    plan_rows = parse_noise_plan()
    area_records, traffic_records = load_available_flattened_noise()
    all_records = [*area_records, *traffic_records]
    if not all_records:
        raise FileNotFoundError(f"未找到可用的噪声监测结果扁平化表: {DEBUG_DIR}")
    apply_noise_result_meta_to_plan_rows(plan_rows, all_records)

    mismatch_warnings = validate_plan_report_match(all_records, plan_rows)
    standard_warnings = build_noise_standard_warnings(plan_rows)
    plan_meta_warnings = build_noise_plan_meta_warnings(plan_rows)
    flow_unit_warnings = build_flow_unit_warnings(all_records)
    position_warnings = build_plan_report_position_warnings(all_records, plan_rows)
    record_data_validation(
        OUTPUT_DIR,
        "noise_data",
        valid=not standard_warnings and not plan_meta_warnings and not flow_unit_warnings and not mismatch_warnings and not position_warnings,
        warnings=[*standard_warnings, *plan_meta_warnings, *flow_unit_warnings, *mismatch_warnings, *position_warnings],
    )
    table1 = build_monitor_points_table(plan_rows, all_records)
    table2 = build_sensitive_result_table(all_records, plan_rows)
    table_indoor = build_indoor_result_table(table2)
    table3 = build_attenuation_result_table(traffic_records, plan_rows)
    summary = build_compliance_summary(table2, table3)

    numbering = load_numbering(OUTPUT_DIR)
    numbering.begin_section("noise", "声环境现状调查与评价")
    register_noise_tables(numbering, table1, table2, table3, table_indoor)

    write_json(DEBUG_DIR / "noise_monitor_points_table.json", table1)
    write_json(DEBUG_DIR / "noise_sensitive_points_result_table.json", table2)
    write_json(DEBUG_DIR / "noise_indoor_result_table.json", table_indoor)
    write_json(DEBUG_DIR / "traffic_noise_attenuation_table.json", table3)
    write_json(DEBUG_DIR / "noise_result_name_position_debug.json", collect_result_name_position_debug(table2, table3))
    monitoring_meta = build_noise_monitoring_meta(all_records, plan_rows)
    summary["monitoring_meta"] = monitoring_meta
    write_json(DEBUG_DIR / "noise_compliance_summary.json", summary)
    write_json(DEBUG_DIR / "noise_monitoring_meta.json", monitoring_meta)

    texts = build_rule_texts(table1, table2, table3, summary, plan_rows, numbering)
    texts = polish_noise_text_with_llm(texts, table1, table2, table3, summary, numbering)
    formal_validation = build_noise_formal_text_validation(
        plan_rows,
        table1,
        table2,
        table3,
        texts,
        data_warnings=[*standard_warnings, *plan_meta_warnings, *flow_unit_warnings, *mismatch_warnings, *position_warnings],
    )
    write_formal_text_validation(DEBUG_DIR, formal_validation)

    doc = build_docx(table1, table2, table3, table_indoor, summary, texts, numbering)
    finalize_section_document(doc)
    doc.save(OUTPUT_DOCX)
    write_json(DEBUG_DIR / "noise_section_texts.json", texts)
    print(f"generated: {OUTPUT_DOCX}")
    print(f"table4.2-1 rows: {len(table1['rows'])}")
    print(f"table4.2-2 rows: {len(table2['rows'])}")
    print(f"table4.2-3 rows: {len(table3['rows'])}")
    print(f"sensitive exceed count: {summary['sensitive']['exceed_count']}")
    print(f"attenuation exceed count: {summary['attenuation']['exceed_count']}")


def parse_noise_plan() -> List[Dict[str, Any]]:
    plan_path = next(
        path for path in INPUT_DIR.glob("*.docx")
        if "方案" in path.name and not path.name.startswith("~$")
    )
    chunks = load_docx_chunks(plan_path)
    defaults = extract_noise_plan_defaults(chunks)
    detection: List[Dict[str, Any]] = []
    for chunk in chunks:
        if chunk.get("kind") != "table":
            continue
        title = (chunk.get("metadata") or {}).get("table_title") or ""
        title_score = noise_plan_title_score(title)
        if title_score < 0:
            detection.append(
                {
                    "source_table": chunk.get("chunk_id"),
                    "table_title": title,
                    "skipped_reason": "excluded_by_title",
                }
            )
            continue
        rows = parse_table_rows(chunk.get("text") or "")
        if len(rows) < 2:
            detection.append(
                {
                    "source_table": chunk.get("chunk_id"),
                    "table_title": title,
                    "skipped_reason": "too_few_rows",
                    "row_count": len(rows),
                }
            )
            continue
        headers = rows[0]
        schema_result = resolve_table_schema(
            "noise_plan",
            headers,
            sample_rows=rows[1:4],
            table_title=title,
            output_dir=OUTPUT_DIR,
            source_table=chunk["chunk_id"],
            job_id=OUTPUT_DIR.parent.name,
        )
        if not schema_result["validation"]["valid"]:
            detection.append(
                {
                    "source_table": chunk.get("chunk_id"),
                    "table_title": title,
                    "headers": headers,
                    "title_score": title_score,
                    "schema_validation": schema_result["validation"],
                    "skipped_reason": "schema_invalid",
                }
            )
            continue
        parsed = [
            normalize_plan_row(
                {
                    **row_to_dict(headers, row),
                    **apply_header_mapping(headers, row, schema_result["mapping"]),
                },
                source_file=str(plan_path),
                source_table=chunk["chunk_id"],
                defaults=defaults,
            )
            for row in rows[1:]
            if len(row) >= 5
        ]
        parsed = [row for row in parsed if is_valid_noise_point_code(row.get("point_code"))]
        detection.append(
            {
                "source_table": chunk.get("chunk_id"),
                "table_title": title,
                "headers": headers,
                "title_score": title_score,
                "schema_validation": schema_result["validation"],
                "parsed_rows": len(parsed),
                "accepted": bool(parsed),
            }
        )
        if parsed:
            write_json(DEBUG_DIR / "noise_plan_table_detection.json", detection)
            return parsed
    write_json(DEBUG_DIR / "noise_plan_table_detection.json", detection)
    raise RuntimeError("未找到声环境现状监测方案表")


def noise_plan_title_score(title: str) -> int:
    text = str(title or "")
    if any(word in text for word in ("振动", "底泥", "土壤", "地下水", "地表水", "水质", "环境空气", "大气")):
        return -1
    score = 0
    if any(word in text for word in ("声环境", "噪声", "声现状")):
        score += 3
    if "监测" in text:
        score += 1
    if "方案" in text or "布点" in text or "布设" in text:
        score += 1
    return score


def normalize_plan_row(
    row: Dict[str, Any],
    source_file: str,
    source_table: str,
    defaults: Optional[Dict[str, str]] = None,
) -> Dict[str, Any]:
    defaults = defaults or {}
    raw_code = cell_value(row, "point_code", "\u76d1\u6d4b\u70b9\u7f16\u53f7", "\u70b9\u4f4d\u7f16\u53f7", "\u7f16\u53f7")
    code = extract_point_code(raw_code) or raw_code
    raw_standard = cell_value(
        row,
        "standard_class",
        "\u73b0\u72b6\u6807\u51c6",
        "\u73b0\u72b6\u6267\u884c\u6807\u51c6",
        "\u58f0\u73b0\u72b6\u6267\u884c\u6807\u51c6",
        "\u58f0\u73af\u5883\u73b0\u72b6\u6267\u884c\u6807\u51c6",
        "\u58f0\u73af\u5883\u6807\u51c6",
        "\u566a\u58f0\u6267\u884c\u6807\u51c6",
        "\u6807\u51c6\u7c7b\u522b",
        "\u6267\u884c\u6807\u51c6",
    )
    standard_day, standard_night = parse_custom_noise_standard_limits(raw_standard)
    raw_frequency = cell_value(row, "frequency", "\u76d1\u6d4b\u9891\u6b21", "\u76d1\u6d4b\u65f6\u95f4", "\u76d1\u6d4b\u8981\u6c42")
    raw_duration = cell_value(row, "monitor_duration", "\u76d1\u6d4b\u65f6\u957f", "\u5355\u6b21\u76d1\u6d4b\u65f6\u95f4")
    monitor_duration = extract_noise_monitor_duration(raw_duration or raw_frequency)
    frequency = strip_noise_monitor_duration(raw_frequency) or defaults.get("frequency", "")
    return {
        "point_code": code,
        "code_detail": extract_code_detail(raw_code),
        "station": cell_value(row, "station", "\u6869\u53f7"),
        "point_name": cell_value(row, "point_name", "\u76d1\u6d4b\u70b9\u540d\u79f0", "\u654f\u611f\u76ee\u6807\u540d\u79f0", "\u540d\u79f0"),
        "district": cell_value(row, "district", "\u884c\u653f\u533a\u5212", "\u6240\u5728\u884c\u653f\u533a"),
        "position": cell_value(row, "position", "\u76d1\u6d4b\u70b9\u4f4d\u7f6e", "\u70b9\u4f4d", "\u76d1\u6d4b\u4f4d\u7f6e"),
        "standard_class": normalize_noise_standard_class(raw_standard),
        "standard_class_raw": raw_standard,
        "standard_day": standard_day,
        "standard_night": standard_night,
        "factor": cell_value(row, "factor", "\u76d1\u6d4b\u56e0\u5b50") or defaults.get("factor", ""),
        "frequency": frequency,
        "monitor_duration": monitor_duration,
        "is_attenuation": is_attenuation_text(" ".join(str(v) for v in row.values())),
        "source_file": source_file,
        "source_table": source_table,
    }


def cell_value(row: Dict[str, Any], *keys: str) -> str:
    for key in keys:
        value = str(row.get(key) or "").strip()
        if value:
            return value
    return ""


def normalize_noise_standard_class(value: Any) -> str:
    text = re.sub(r"\s+", "", str(value or "")).lower()
    if not text:
        return ""
    if parse_custom_noise_standard_limits(text) != (None, None):
        return ""
    text = text.replace("类区", "类")
    match = re.search(r"(0|1|2|3|4a|4b|4|ⅰ|ⅱ|ⅲ|ⅳ|ⅴ|i{1,3}|iv|v)\s*类?", text, flags=re.IGNORECASE)
    token = match.group(1) if match else text
    mapping = {
        "0": "0类",
        "1": "1类",
        "2": "2类",
        "3": "3类",
        "4": "4a类",
        "4a": "4a类",
        "4b": "4b类",
        "ⅰ": "1类",
        "ⅱ": "2类",
        "ⅲ": "3类",
        "ⅳ": "4a类",
        "ⅴ": "4b类",
        "i": "1类",
        "ii": "2类",
        "iii": "3类",
        "iv": "4a类",
        "v": "4b类",
    }
    if token in mapping:
        return mapping[token]
    return str(value or "").strip()


def parse_custom_noise_standard_limits(value: Any) -> Tuple[Optional[int], Optional[int]]:
    text = str(value or "")
    match = re.search(r"(?<!\d)(\d{2,3})(?:\.0)?\s*/\s*(\d{2,3})(?:\.0)?(?!\d)", text)
    if not match:
        return None, None
    return int(match.group(1)), int(match.group(2))


def load_noise_standard_limits() -> Dict[str, Any]:
    global _NOISE_LIMITS_CACHE
    if _NOISE_LIMITS_CACHE is not None:
        return _NOISE_LIMITS_CACHE
    try:
        payload = json.loads(NOISE_STANDARD_LIMITS_PATH.read_text(encoding="utf-8"))
        limits = payload.get("limits") if isinstance(payload, dict) else None
        if not isinstance(limits, dict):
            raise ValueError("limits must be an object")
        _NOISE_LIMITS_CACHE = payload
    except Exception as exc:
        _NOISE_LIMITS_CACHE = {"limits": {}, "error": str(exc)}
    return _NOISE_LIMITS_CACHE


def noise_standard_limits(plan: Dict[str, Any]) -> Tuple[Optional[Decimal], Optional[Decimal]]:
    custom_day = plan.get("standard_day")
    custom_night = plan.get("standard_night")
    if custom_day is not None and custom_night is not None:
        return parse_decimal(custom_day), parse_decimal(custom_night)
    item = (load_noise_standard_limits().get("limits") or {}).get(plan.get("standard_class")) or {}
    return parse_decimal(item.get("day")), parse_decimal(item.get("night"))


def noise_standard_display(plan: Dict[str, Any]) -> str:
    custom_day, custom_night = noise_standard_limits(plan)
    if plan.get("standard_day") is not None and plan.get("standard_night") is not None:
        return f"{custom_day}/{custom_night}"
    return str(plan.get("standard_class") or plan.get("standard_class_raw") or "").strip()


def build_noise_standard_warnings(plan_rows: List[Dict[str, Any]]) -> List[str]:
    warnings: List[str] = []
    for row in plan_rows:
        day_limit, night_limit = noise_standard_limits(row)
        raw_standard = str(row.get("standard_class_raw") or row.get("standard_class") or "空").strip()
        if day_limit is None or night_limit is None:
            warnings.append(f"{row.get('point_code')} 未识别现状标准: {raw_standard}")
        elif row.get("standard_day") is not None and row.get("standard_night") is not None:
            warnings.append(f"{row.get('point_code')} 使用自定义昼/夜限值{day_limit}/{night_limit}，需复核")
    return warnings


def extract_noise_plan_defaults(chunks: List[Dict[str, Any]]) -> Dict[str, str]:
    text = noise_plan_meta_source_text(chunks)
    defaults: Dict[str, str] = {}
    status: Dict[str, Any] = {
        "rule": {},
        "llm_status": "not_needed",
        "warnings": [],
        "final": {},
    }

    factor = extract_noise_factor_rule(text)
    if factor:
        defaults["factor"] = factor
        status["rule"]["factor"] = factor
    frequency = extract_noise_frequency_rule(text)
    if frequency:
        defaults["frequency"] = frequency
        status["rule"]["frequency"] = frequency

    missing = [
        field
        for field in ("factor", "frequency")
        if not is_valid_noise_plan_meta(field, defaults.get(field, ""))
    ]
    if missing:
        llm_defaults, llm_status = extract_noise_plan_defaults_with_llm(chunks, defaults, missing)
        status.update(llm_status)
        for field, value in llm_defaults.items():
            if field in missing and is_valid_noise_plan_meta(field, value):
                defaults[field] = value

    for field in ("factor", "frequency"):
        if not is_valid_noise_plan_meta(field, defaults.get(field, "")):
            status["warnings"].append(
                "监测因子未识别" if field == "factor" else "监测频次未识别"
            )
    status["final"] = {field: defaults.get(field, "") for field in ("factor", "frequency")}
    write_json(DEBUG_DIR / "noise_plan_meta_status.json", status)
    return defaults


def noise_plan_meta_source_text(chunks: List[Dict[str, Any]]) -> str:
    parts: List[str] = []
    for chunk in chunks:
        kind = chunk.get("kind")
        if kind == "paragraphs":
            parts.append(str(chunk.get("text") or ""))
        elif kind == "table":
            title = (chunk.get("metadata") or {}).get("table_title") or ""
            if noise_plan_title_score(title) >= 0:
                parts.append(str(title))
                parts.append(str(chunk.get("text") or "")[:2000])
    return "\n".join(part for part in parts if part).strip()


def extract_noise_factor_rule(text: str) -> str:
    patterns = [
        r"(?:监测因子|监测项目)\s*[:：]?\s*([^\n。；;]+)",
        r"(?:因子|项目)\s*[:：]\s*([^\n。；;]+)",
    ]
    for pattern in patterns:
        match = re.search(pattern, text or "")
        if match:
            value = clean_plan_meta_value(match.group(1))
            if is_valid_noise_plan_meta("factor", value):
                return value
    if re.search(r"等效连续\s*A\s*声级", text or "", flags=re.IGNORECASE):
        return "等效连续A声级（LAeq）"
    if re.search(r"\bLAeq\b", text or "", flags=re.IGNORECASE):
        return "LAeq"
    return ""


def extract_noise_frequency_rule(text: str) -> str:
    source = str(text or "")
    if re.search(r"连续监测\s*2\s*天", source) and re.search(r"昼间[、,，和及]?夜间.*各监测\s*1\s*次", source, flags=re.DOTALL):
        suffix = "，每次测量时间不低于20min" if re.search(r"20\s*min|20\s*分钟", source, flags=re.IGNORECASE) else ""
        return f"连续监测2天，每天昼间、夜间各监测1次{suffix}。"
    sentences = re.split(r"(?<=[。；;])|\n+", source)
    candidates = []
    for sentence in sentences:
        value = clean_plan_meta_value(sentence)
        if not value:
            continue
        if any(marker in value for marker in ("监测频次", "监测时间", "连续监测", "昼间", "夜间", "每次测量")):
            candidates.append(value)
    combined = "；".join(candidates[:3])
    return combined if is_valid_noise_plan_meta("frequency", combined) else ""


def clean_plan_meta_value(value: str) -> str:
    text = re.sub(r"\s+", " ", str(value or "")).strip(" ：:;；。")
    text = re.sub(r"^[:：]+", "", text).strip()
    return text[:240]


def is_valid_noise_plan_meta(field: str, value: str) -> bool:
    text = str(value or "").strip()
    if not text or text in {"-", "/", "无", "None", "none", "null"}:
        return False
    if re.search(r"\bN(?:J)?\d+(?:-\d+)?\b", text, flags=re.IGNORECASE):
        return False
    if field == "factor":
        return len(text) >= 2 and len(text) <= 120
    if field == "frequency":
        return len(text) >= 6 and any(marker in text for marker in ("监测", "昼", "夜", "次", "天", "min", "分钟"))
    return bool(text)


def extract_noise_plan_defaults_with_llm(
    chunks: List[Dict[str, Any]],
    defaults: Dict[str, str],
    missing: List[str],
) -> Tuple[Dict[str, str], Dict[str, Any]]:
    status: Dict[str, Any] = {
        "llm_status": "not_needed",
        "llm_missing_fields": missing,
        "llm_error": None,
        "warnings": [],
    }
    if not os.getenv("EIA_LLM_API_KEY"):
        status["llm_status"] = "skipped_no_api_key"
        status["warnings"].append("LLM方案元信息兜底跳过：未配置 API Key")
        return {}, status

    payload = {
        "task": "extract_noise_plan_metadata",
        "missing_fields": missing,
        "existing_rule_values": defaults,
        "constraints": [
            "只提取监测方案中的监测因子和监测频次/监测时间。",
            "不要提取或生成监测结果数值。",
            "不要返回监测点位编号、点位名称或点位列表。",
            "如果原文没有相关信息，对应字段返回空字符串。",
        ],
        "source_text": noise_plan_meta_source_text(chunks)[:5000],
    }
    write_json(DEBUG_DIR / "noise_plan_meta_llm_input.json", payload)
    try:
        result = chat_completion_json_object(
            [
                {"role": "system", "content": "Return strict JSON only."},
                {
                    "role": "user",
                    "content": (
                        "你是环评监测方案元信息抽取器。返回 JSON："
                        "{\"factor\":\"\",\"frequency\":\"\",\"confidence\":0,\"warnings\":[]}。\n"
                        + json.dumps(payload, ensure_ascii=False)
                    ),
                },
            ],
            profile=LlmProfile.fast,
            max_retries=1,
            max_tokens=700,
            label="noise_plan_meta",
        )
        write_json(DEBUG_DIR / "noise_plan_meta_llm_output.json", result)
        accepted: Dict[str, str] = {}
        for field in ("factor", "frequency"):
            value = clean_plan_meta_value(str(result.get(field) or ""))
            if field in missing and is_valid_noise_plan_meta(field, value):
                accepted[field] = value
            elif field in missing and value:
                status["warnings"].append(
                    f"LLM方案元信息兜底结果未通过校验: {field}={value}"
                )
        status["llm_status"] = "success" if accepted else "validation_failed"
        status["llm_confidence"] = result.get("confidence")
        if isinstance(result.get("warnings"), list):
            status["warnings"].extend(str(item) for item in result.get("warnings") if item)
        return accepted, status
    except Exception as exc:
        status["llm_status"] = "failed"
        status["llm_error"] = str(exc)
        status["warnings"].append(f"LLM方案元信息兜底失败: {exc}")
        return {}, status


def load_flattened_noise(filename: str, noise_type: str) -> List[Dict[str, Any]]:
    path = DEBUG_DIR / filename
    if not path.exists():
        raise FileNotFoundError(f"缺少噪声扁平化表: {path}")
    data = json.loads(path.read_text(encoding="utf-8"))
    return parse_flattened_noise_records(data, noise_type, table_order=flattened_table_sort_key(path)[0])


def load_available_flattened_noise() -> Tuple[List[Dict[str, Any]], List[Dict[str, Any]]]:
    area_records: List[Dict[str, Any]] = []
    traffic_records: List[Dict[str, Any]] = []
    debug_items: List[Dict[str, Any]] = []
    paths = sorted(DEBUG_DIR.glob("flattened_table_*.json"), key=flattened_table_sort_key)
    for path in paths:
        data = json.loads(path.read_text(encoding="utf-8"))
        noise_type = infer_flattened_noise_type(data)
        records = parse_flattened_noise_records(data, noise_type, table_order=flattened_table_sort_key(path)[0])
        debug_items.append(
            {
                "file": path.name,
                "noise_type": noise_type,
                "record_count": len(records),
                "table_title": data.get("table_title") or "",
                "flattened_headers": data.get("flattened_headers") or [],
            }
        )
        if noise_type == "traffic_noise":
            traffic_records.extend(records)
        else:
            area_records.extend(records)
    write_json(DEBUG_DIR / "noise_flattened_table_detection.json", debug_items)
    return area_records, traffic_records


def flattened_table_sort_key(path: Path) -> Tuple[int, str]:
    match = re.search(r"flattened_table_(\d+)\.json$", path.name)
    return (int(match.group(1)) if match else 9999, path.name)


def infer_flattened_noise_type(data: Dict[str, Any]) -> str:
    text_parts = [
        str(data.get("table_title") or ""),
        " ".join(str(item) for item in data.get("flattened_headers") or []),
        " ".join(
            str(key)
            for row in (data.get("records") or [])[:5]
            for key in row.keys()
        ),
    ]
    text = "\n".join(text_parts)
    if "铁路边界噪声" in text or "列车流量" in text:
        return "railway_boundary_noise"
    traffic_markers = ("交通噪声", "车流量", "衰减断面", "交通量", "距路")
    if any(marker in text for marker in traffic_markers):
        return "traffic_noise"
    return "area_environment_noise"


def parse_road_flow_breakdown(value: Any) -> Dict[str, str]:
    text = str(value or "").strip()
    if text in {"", "-", "/", "\u2014", "None", "none", "null"}:
        return {}
    parts = [part.strip() for part in re.split(r"[/\uff0f,\uff0c\s]+", text) if part.strip()]
    if len(parts) != 3:
        return {}
    return dict(zip(ROAD_VEHICLE_KEYS, parts))


def vehicle_flow_value_from_row(row: Dict[str, Any], vehicle: str) -> str:
    aliases = {
        "large": ("\u5927\u578b\u8f66", "\u5927\u8f66"),
        "medium": ("\u4e2d\u578b\u8f66", "\u4e2d\u8f66"),
        "small": ("\u5c0f\u578b\u8f66", "\u5c0f\u8f66"),
    }[vehicle]
    for key, value in row.items():
        header = re.sub(r"\s+", "", str(key or ""))
        if any(alias in header for alias in aliases):
            return str(value or "").strip()
    return ""


def road_flow_breakdown_from_row(row: Dict[str, Any], raw_flow: Any) -> Dict[str, str]:
    explicit = {
        vehicle: vehicle_flow_value_from_row(row, vehicle)
        for vehicle in ROAD_VEHICLE_KEYS
    }
    if any(value not in {"", "-", "/", "\u2014"} for value in explicit.values()):
        return {vehicle: explicit[vehicle] or "/" for vehicle in ROAD_VEHICLE_KEYS}
    return parse_road_flow_breakdown(raw_flow)


def canonical_road_flow_text(breakdown: Dict[str, str], fallback: Any = "/") -> str:
    if not breakdown:
        return str(fallback or "/")
    return "/".join(str(breakdown.get(vehicle) or "/") for vehicle in ROAD_VEHICLE_KEYS)

def parse_flattened_noise_records(data: Dict[str, Any], noise_type: str, table_order: int = 9999) -> List[Dict[str, Any]]:
    records: List[Dict[str, Any]] = []
    raw_records = data.get("records") or []
    headers = list(data.get("flattened_headers") or [])
    if not headers and raw_records:
        headers = [str(key) for key in raw_records[0].keys() if not str(key).startswith("_")]
    schema_result = resolve_table_schema(
        "noise_result",
        headers,
        sample_rows=[
            [str(row.get(header) or "") for header in headers]
            for row in raw_records[:3]
        ],
        table_title=str(data.get("table_title") or ""),
        output_dir=OUTPUT_DIR,
        source_table=str(data.get("chunk_id") or ""),
        job_id=OUTPUT_DIR.parent.name,
    )
    mapping = schema_result.get("mapping") or {}
    point_key = mapping.get("point") or POINT_KEY
    time_key = mapping.get("monitor_time") or TIME_KEY
    mapped_laeq_key = mapping.get("laeq")
    flow_key = mapping.get("traffic_flow") or FLOW_KEY
    flow_unit = flow_label_from_header(flow_key)
    factor_items = noise_factor_items_from_headers(headers)
    for index, row in enumerate(raw_records):
        point_text = str(row.get(point_key) or "").strip()
        laeq_key = mapped_laeq_key or next((key for key in row if key.endswith("_LAeq")), None)
        if not point_text or not laeq_key:
            continue
        raw_point_code = extract_point_code(point_text)
        if not raw_point_code and not has_noise_point_identity(point_text):
            continue
        raw_flow = row.get(FLOW_KEY) or row.get(flow_key) or "/"
        explicit_vehicle_flow = any(
            vehicle_flow_value_from_row(row, vehicle) not in {"", "-", "/", "\u2014"}
            for vehicle in ROAD_VEHICLE_KEYS
        )
        flow_breakdown = road_flow_breakdown_from_row(row, raw_flow) if "\u8f66" in flow_unit or explicit_vehicle_flow else {}
        traffic_flow = canonical_road_flow_text(flow_breakdown, raw_flow)
        records.append(
            {
                "noise_type": noise_type,
                "point_text": point_text,
                "raw_point_code": raw_point_code,
                "monitor_time": row.get(time_key) or "",
                "period": infer_period(row.get(time_key) or ""),
                "laeq": row.get(laeq_key),
                "laeq_key": laeq_key,
                "traffic_flow": traffic_flow,
                "traffic_flow_large": flow_breakdown.get("large", "/"),
                "traffic_flow_medium": flow_breakdown.get("medium", "/"),
                "traffic_flow_small": flow_breakdown.get("small", "/"),
                "traffic_flow_unit": flow_unit,
                "noise_factor_items": factor_items,
                "source_headers": headers,
                "source_table_order": table_order,
                "source_order": index,
                "raw": row,
            }
        )
    return records


def apply_noise_result_meta_to_plan_rows(plan_rows: List[Dict[str, Any]], records: List[Dict[str, Any]]) -> None:
    factor = infer_noise_factor_from_records(records)
    frequency = infer_noise_frequency_from_records(
        records,
        [str(row.get("frequency") or "") for row in plan_rows if row.get("frequency")],
    )
    status_path = DEBUG_DIR / "noise_plan_meta_status.json"
    status = read_json(status_path) if status_path.exists() else {}
    status.setdefault("result_header_rule", {})
    if factor:
        status["result_header_rule"]["factor"] = factor
    if frequency:
        status["result_header_rule"]["frequency"] = frequency
    for row in plan_rows:
        if factor:
            row["factor"] = factor
        if frequency:
            row["frequency"] = frequency
    status["final"] = {
        "factor": factor or next((row.get("factor") for row in plan_rows if row.get("factor")), ""),
        "frequency": frequency or next((row.get("frequency") for row in plan_rows if row.get("frequency")), ""),
        "monitor_durations": sorted({str(row.get("monitor_duration") or "") for row in plan_rows if row.get("monitor_duration")}),
    }
    status["frequency_rule"] = build_noise_frequency_status(records, frequency)
    write_json(DEBUG_DIR / "noise_plan_meta_status.json", status)


def infer_noise_factor_from_records(records: List[Dict[str, Any]]) -> str:
    preferred = ["LAeq", "L10", "L50", "L90", "Lmax", "Lmin"]
    found = []
    for record in records:
        for item in record.get("noise_factor_items") or []:
            if item not in found:
                found.append(item)
    ordered = [item for item in preferred if item in found]
    return "、".join(ordered or found)


def noise_factor_items_from_headers(headers: List[str]) -> List[str]:
    items: List[str] = []
    for header in headers:
        text = str(header or "")
        for item in ("LAeq", "L10", "L50", "L90", "Lmax", "Lmin"):
            if re.search(rf"(?<![A-Za-z0-9]){re.escape(item)}(?![A-Za-z0-9])", text, flags=re.IGNORECASE):
                canonical = item
                if canonical not in items:
                    items.append(canonical)
    return items


def extract_noise_monitor_duration(value: Any) -> str:
    text = str(value or "")
    if not text:
        return ""
    if re.search(r"\u4e0d(?:\u5c0f|\u5c11)\u4e8e\s*1\s*h", text, flags=re.IGNORECASE):
        return "\u4e0d\u5c0f\u4e8e1h"
    match = re.search(r"(?:\u6bcf\u6b21(?:\u6d4b\u91cf|\u76d1\u6d4b)?(?:\u65f6\u95f4)?(?:\u4e0d(?:\u5c0f|\u5c11)\u4e8e)?\s*)?(\d+(?:\.\d+)?)\s*(min|\u5206\u949f|h|\u5c0f\u65f6)", text, flags=re.IGNORECASE)
    if not match:
        return ""
    number = match.group(1)
    unit = match.group(2).lower()
    if unit in {"min", "\u5206\u949f"}:
        return f"{number}min"
    if unit in {"h", "\u5c0f\u65f6"}:
        return f"{number}h"
    return ""


def strip_noise_monitor_duration(value: str) -> str:
    text = str(value or "").strip()
    if not text:
        return ""
    text = re.sub(r"(?:\uff1b|;|\uff0c|,)?\s*\u6bcf\u6b21(?:\u6d4b\u91cf|\u76d1\u6d4b)?(?:\u65f6\u95f4)?(?:\u4e0d(?:\u5c0f|\u5c11)\u4e8e)?\s*\d+(?:\.\d+)?\s*(?:min|\u5206\u949f|h|\u5c0f\u65f6)", "", text, flags=re.IGNORECASE)
    text = re.sub(r"(?:\u4e0d(?:\u5c0f|\u5c11)\u4e8e)?\s*\d+(?:\.\d+)?\s*(?:min|\u5206\u949f|h|\u5c0f\u65f6)", "", text, flags=re.IGNORECASE)
    text = re.sub(r"[\uff0c,\uff1b;\uff1a:]?\s*(?:\u8be6\u89c1)?\u8868\d+[^\u3002\uff1b;]*\u76d1\u6d4b\u65f6\u95f4[^\u3002\uff1b;]*", "", text)
    text = re.sub(r"\s+", "", text).strip("\uff0c,\uff1b;\u3002")
    return compact_noise_frequency(text) if text else ""


def compact_noise_frequency(value: str) -> str:
    text = str(value or "")
    if not text:
        return ""
    if "列车" in text or "铁路" in text:
        return "监测2天，昼、夜间各监测1次；铁路边界噪声每次不小于1h，区域环境噪声每次60min。"
    if re.search(r"连续监测\s*2\s*天|监测\s*2\s*天", text) and re.search(r"昼.*夜|夜.*昼", text):
        return "监测2天，昼、夜间各监测1次。"
    if len(text) > 80:
        return text[:80].rstrip("，,；;。") + "。"
    return text


def infer_noise_frequency_from_records(records: List[Dict[str, Any]], plan_frequencies: List[str]) -> str:
    noise_types = {str(record.get("noise_type") or "") for record in records}
    flow_labels = {
        canonical_flow_label(record.get("traffic_flow_unit"))
        for record in records
        if has_record_traffic_flow(record)
    }
    if noise_types or flow_labels:
        return "\u76d1\u6d4b2\u5929\uff0c\u663c\u3001\u591c\u95f4\u5404\u76d1\u6d4b1\u6b21\u3002"
    first_frequency = next((item for item in plan_frequencies if item), "")
    return strip_noise_monitor_duration(first_frequency) or compact_noise_frequency(first_frequency)


def build_noise_frequency_status(records: List[Dict[str, Any]], final_frequency: str) -> Dict[str, Any]:
    return {
        "noise_types": sorted({str(record.get("noise_type") or "") for record in records if record.get("noise_type")}),
        "flow_units": sorted(
            {
                canonical_flow_label(record.get("traffic_flow_unit"))
                for record in records
                if has_record_traffic_flow(record)
            }
        ),
        "final_frequency": final_frequency,
        "source": "result_type_and_plan_rule",
    }


def flow_label_from_header(header: str) -> str:
    text = str(header or "").strip()
    return canonical_flow_label(text) or "车流量（辆/20min）"


def has_noise_point_identity(text: str) -> bool:
    value = str(text or "").strip()
    if len(value) < 4:
        return False
    return any(marker in value for marker in ("点位编号", "敏感点", "背景点", "衰减断面"))


def build_monitor_points_table(plan_rows: List[Dict[str, Any]], records: Optional[List[Dict[str, Any]]] = None) -> Dict[str, Any]:
    headers = ["\u76d1\u6d4b\u70b9\u7f16\u53f7", "\u6869\u53f7", "\u76d1\u6d4b\u70b9\u540d\u79f0", "\u884c\u653f\u533a\u5212", "\u76d1\u6d4b\u70b9\u4f4d\u7f6e", "\u73b0\u72b6\u6807\u51c6", "\u76d1\u6d4b\u56e0\u5b50", "\u76d1\u6d4b\u9891\u6b21"]
    duration_values = sorted({re.sub(r"\s+", "", str(row.get("monitor_duration") or "")) for row in plan_rows if row.get("monitor_duration")})
    include_monitor_time = len(duration_values) >= 2
    if include_monitor_time:
        headers.append("\u76d1\u6d4b\u65f6\u95f4")
    plan_records = first_record_by_plan_row(records or [], plan_rows)
    rows = []
    debug_rows = []
    for row in plan_rows:
        record = plan_records.get(id(row))
        final_name = monitor_point_display_name(row, record)
        final_position = monitor_point_display_position(row, record)
        row_payload = {
            "\u76d1\u6d4b\u70b9\u7f16\u53f7": row["point_code"],
            "\u6869\u53f7": row["station"],
            "\u76d1\u6d4b\u70b9\u540d\u79f0": final_name,
            "\u884c\u653f\u533a\u5212": row["district"],
            "\u76d1\u6d4b\u70b9\u4f4d\u7f6e": final_position,
            "\u73b0\u72b6\u6807\u51c6": noise_standard_display(row),
            "\u76d1\u6d4b\u56e0\u5b50": row["factor"],
            "\u76d1\u6d4b\u9891\u6b21": row["frequency"],
        }
        if include_monitor_time:
            row_payload["\u76d1\u6d4b\u65f6\u95f4"] = row.get("monitor_duration") or "-"
        rows.append(row_payload)
        debug_rows.append(
            {
                "point_code": row.get("point_code"),
                "raw_plan_name": row.get("point_name"),
                "raw_plan_position": row.get("position"),
                "raw_report_point": (record or {}).get("point_text"),
                "position_source": "plan_complete" if complete_floor_position_from_plan(str(row.get("position") or "")) else "report_fallback",
                "monitor_duration": row.get("monitor_duration"),
                "final_name": final_name,
                "final_position": final_position,
            }
        )
    if not any(str(row.get("\u884c\u653f\u533a\u5212") or "").strip() for row in rows):
        headers = [header for header in headers if header != "\u884c\u653f\u533a\u5212"]
        for row in rows:
            row.pop("\u884c\u653f\u533a\u5212", None)
    if all(is_blank_table_value(row.get("\u6869\u53f7")) for row in rows):
        headers = [header for header in headers if header != "\u6869\u53f7"]
        for row in rows:
            row.pop("\u6869\u53f7", None)
    write_json(DEBUG_DIR / "noise_monitor_points_name_position_debug.json", debug_rows)
    return {
        "table_key": NOISE_TABLE_MONITOR_POINTS,
        "caption_suffix": "\u58f0\u73af\u5883\u8d28\u91cf\u73b0\u72b6\u76d1\u6d4b\u70b9",
        "title": "\u58f0\u73af\u5883\u8d28\u91cf\u73b0\u72b6\u76d1\u6d4b\u70b9",
        "headers": headers,
        "rows": rows,
    }


def first_record_by_plan_row(records: List[Dict[str, Any]], plan_rows: List[Dict[str, Any]]) -> Dict[int, Dict[str, Any]]:
    if not records:
        return {}
    plan_row_ids = {id(row) for row in plan_rows}
    result: Dict[int, Dict[str, Any]] = {}
    for record in match_records(records, plan_rows):
        plan = record.get("plan") or {}
        plan_id = id(plan)
        if plan_id not in plan_row_ids or plan_id in result:
            continue
        result[plan_id] = record
    return result


def monitor_point_display_name(plan: Dict[str, Any], record: Optional[Dict[str, Any]]) -> str:
    if is_railway_boundary_point_name(plan.get("point_name")):
        return str(plan.get("point_name") or "").strip()
    if plan.get("is_attenuation"):
        return clean_attenuation_point_name(plan, record)
    candidates = [
        str(plan.get("point_name") or ""),
        str((record or {}).get("point_text") or ""),
    ]
    detail = str(plan.get("code_detail") or "").strip()
    if detail:
        candidates.insert(0, f"{plan.get('point_name') or ''}{detail}")
    for candidate in candidates:
        name = clean_monitor_point_object_name(candidate, plan, record)
        if name:
            return name
    return "-"


def monitor_point_display_position(plan: Dict[str, Any], record: Optional[Dict[str, Any]]) -> str:
    if plan.get("is_attenuation"):
        return clean_attenuation_point_position(plan, record)
    return compose_plan_monitor_position(
        plan.get("position"),
        (record or {}).get("point_text"),
    )


def compose_plan_monitor_position(plan_position: Any, report_text: Any) -> str:
    plan_text = str(plan_position or "").strip()
    report_value = str(report_text or "").strip()
    report_floor = extract_floor_position_fragment(report_value)
    if report_floor:
        context = strip_plan_floor_terms(plan_text)
        marker = (
            "室外" if "室外" in report_value or "室外" in plan_text
            else "室内" if "室内" in report_value or "室内" in plan_text
            else ""
        )
        floor = apply_indoor_outdoor_marker(report_floor, marker)
        context = context.replace("室外", "").replace("室内", "").strip()
        return normalize_floor_position_display(f"{context}{floor}" if context else floor)
    if plan_text:
        return normalize_display_position(plan_text)
    return normalize_display_position(report_value) if report_value else "-"


def clean_monitor_point_object_name(
    value: Any,
    plan: Dict[str, Any],
    record: Optional[Dict[str, Any]] = None,
) -> str:
    text = str(value or "").strip()
    if not text:
        return ""
    raw_code = str((record or {}).get("raw_point_code") or plan.get("point_code") or "").strip()
    station = str(plan.get("station") or "").strip()
    for token in (raw_code, raw_code.replace("NJ", "N", 1) if raw_code.startswith("NJ") else ""):
        if token:
            text = re.sub(rf"(?<![A-Za-z0-9]){re.escape(token)}(?![A-Za-z0-9-])", "", text, flags=re.IGNORECASE)
    if station:
        text = text.replace(station, "")
    text = re.sub(r"K\s*\d+\s*[+-]\s*\d+", "", text, flags=re.IGNORECASE)
    text = re.sub(r"点位编号\s*[A-Za-z]*\d+(?:-\d+)?", "", text, flags=re.IGNORECASE)
    text = re.split(rf"(?:面向|背向|距离|距|{ROW_POSITION_MARKER_PATTERN}|道路红线|道路中心线|公路中心线|铁路|陇海线)", text, maxsplit=1)[0]
    text = re.sub(r"\d+(?:\.\d+)?\s*m", "", text, flags=re.IGNORECASE)
    text = re.sub(r"\d+\s*[层楼](?:室外|室内)?(?:监测)?", "", text)
    text = re.sub(r"(?:室外|室内)\s*(?:监测|噪声监测)?", "", text)
    text = re.sub(r"(?:敏感点|背景点|噪声监测|监测点)", "", text)
    text = re.sub(r"\s+", "", text).strip("、，,;；。()（）")
    text = re.sub(r"(?<=[\u4e00-\u9fff])(\d+)$", r" \1", text)
    return text.strip()


def build_monitor_point_position(plan_position: str, report_position: str) -> str:
    plan_text = str(plan_position or "").strip()
    report_text = str(report_position or "").strip()
    for source in (report_text, plan_text):
        if has_orientation_marker(source) and (re.search(r"\d+\s*[层楼]", source) or is_plain_floor_position(source)):
            return normalize_display_position(source)
    orientation = extract_orientation_context(report_text) or extract_orientation_context(plan_text)
    floor = extract_floor_position_fragment(report_text) or first_floor_position_from_plan(plan_text)
    if orientation and floor:
        return normalize_floor_position_display(f"{orientation} {floor}")
    if orientation:
        return normalize_floor_position_display(orientation)
    if floor:
        return normalize_floor_position_display(floor)
    return normalize_floor_position_display(normalize_display_position(report_text or plan_text))


def extract_orientation_context(text: str) -> str:
    value = str(text or "")
    match = re.search(
        rf"(面向[^，,；;。]*?(?:{ROW_POSITION_MARKER_PATTERN}|本项目|道路|公路|铁路)|背向[^，,；;。]*?(?:{ROW_POSITION_MARKER_PATTERN}|本项目|道路|公路|铁路))",
        value,
    )
    if match:
        result = re.sub(r"\d+\s*[层楼](?:室外|室内)?(?:监测)?", "", match.group(1))
        result = re.sub(r"\s+", "", result).strip("、，,;；。")
        return result
    row_marker = re.search(ROW_POSITION_MARKER_PATTERN, value)
    if row_marker:
        return f"面向本项目{row_marker.group(0)}" if "本项目" in value or "面向" in value else row_marker.group(0)
    return ""


def normalize_display_position(text: str) -> str:
    value = normalize_position_text(str(text or ""))
    value = re.sub(r"\s+", "", value).strip("、，,;；。")
    value = re.sub(r"(面向[^，,；;。]*?)(\d+\s*[层楼])", r"\1 \2", value)
    return normalize_floor_position_display(value)


def normalize_floor_position_display(value: Any) -> str:
    text = re.sub(r"\s+", "", str(value or "")).strip("、，,;；。")
    if not text or text == "-":
        return text
    text = text.replace("楼", "层")
    if re.fullmatch(r"\d+", text):
        return f"{text}层"
    match = re.fullmatch(r"(\d+)(室外|室内)", text)
    if match:
        return f"{match.group(1)}层{match.group(2)}"
    if has_orientation_marker(text) and "m" not in text.lower() and "米" not in text and not re.search(r"\d+\s*层", text):
        text = re.sub(r"(\d+)(室外|室内)?$", lambda m: f"{m.group(1)}层{m.group(2) or ''}", text)
    text = re.sub(r"(\d+)层层", r"\1层", text)
    return text


def is_plain_floor_position(value: Any) -> bool:
    text = re.sub(r"\s+", "", str(value or "")).strip("、，,;；。")
    if re.fullmatch(r"\d+(?:室外|室内)?", text):
        return True
    return bool(has_orientation_marker(text) and "m" not in text.lower() and "米" not in text and re.search(r"\d+(?:室外|室内)?$", text))


def clean_attenuation_point_name(plan: Dict[str, Any], record: Optional[Dict[str, Any]]) -> str:
    for source in (plan.get("point_name"), (record or {}).get("point_text"), plan.get("position")):
        text = str(source or "").strip()
        if not text:
            continue
        text = re.sub(r"K\s*\d+\s*[+-]\s*\d+", "", text, flags=re.IGNORECASE)
        text = re.sub(r"N(?:J)?\d+(?:-\d+)?", "", text, flags=re.IGNORECASE)
        match = re.search(r"([^，,；;。]*断面)", text)
        if match:
            name = clean_monitor_point_object_name(match.group(1), plan, record)
            return name or "衰减断面"
    return "衰减断面"


def clean_attenuation_point_position(plan: Dict[str, Any], record: Optional[Dict[str, Any]]) -> str:
    raw_code = str((record or {}).get("raw_point_code") or plan.get("point_code") or "").strip()
    station = str(plan.get("station") or "").strip()
    point_name = str(plan.get("point_name") or "").strip()
    for source in ((record or {}).get("point_text"), plan.get("position")):
        text = clean_attenuation_position_text(source, raw_code, station, point_name)
        if text:
            return text
    fallback = strip_point_identity(str(plan.get("position") or ""), raw_code, station, point_name)
    fallback = re.sub(r"点位编号\s*[A-Za-z]*\d+(?:-\d+)?", "", fallback, flags=re.IGNORECASE)
    fallback = re.sub(r"N(?:J)?\d+(?:-\d+)?", "", fallback, flags=re.IGNORECASE)
    return re.sub(r"\s+", "", fallback).strip("、，,;；。") or "-"


def clean_attenuation_position_text(value: Any, raw_code: str, station: str, point_name: str) -> str:
    text = strip_point_identity(str(value or ""), raw_code, station, point_name)
    text = re.sub(r"点位编号\s*[A-Za-z]*\d+(?:-\d+)?", "", text, flags=re.IGNORECASE)
    text = re.sub(r"N(?:J)?\d+(?:-\d+)?", "", text, flags=re.IGNORECASE)
    text = re.sub(r"K\s*\d+\s*[+-]\s*\d+", "", text, flags=re.IGNORECASE)
    match = re.search(
        r"((?:距离|距)[^，,；;。]*?\d+(?:\.\d+)?\s*m[^，,；;。]*)",
        text,
        flags=re.IGNORECASE,
    )
    if not match:
        return ""
    return re.sub(r"\s+", "", match.group(1)).strip("、，,;；。")


def clean_result_table_point_name(record: Dict[str, Any], plan: Dict[str, Any]) -> str:
    if is_railway_boundary_point_name(plan.get("point_name")):
        return str(plan.get("point_name") or "").strip()
    display_name = monitor_point_display_name(plan, record)
    if display_name and display_name != "-":
        return display_name
    for source in (plan.get("point_name"), record.get("point_text"), result_name_for_record(record, plan)):
        name = clean_monitor_point_object_name(source, plan, record)
        if name:
            return name
    return "-"


def split_monitor_point_identity(
    point_text: Any,
    plan: Dict[str, Any],
    record: Optional[Dict[str, Any]] = None,
) -> Tuple[str, str]:
    text = str(point_text or "").strip()
    if not text:
        return "", ""
    raw_code = str((record or {}).get("raw_point_code") or "").strip()
    text = strip_noise_point_code_prefix(text, raw_code)
    position = extract_floor_position_fragment(text)
    name = text
    if position:
        name = re.sub(rf"{re.escape(position)}\s*$", "", name)
    name = sanitize_result_point_name(name, {**plan, "position": ""})
    if not name:
        name = monitor_point_display_name({**plan, "code_detail": ""}, None)
    return name, position


def strip_noise_point_code_prefix(text: str, raw_code: str = "") -> str:
    result = str(text or "").strip()
    if raw_code:
        n_code = raw_code.replace("NJ", "N", 1)
        result = re.sub(rf"^\s*(?:{re.escape(raw_code)}|{re.escape(n_code)})\s*", "", result, flags=re.IGNORECASE)
    result = re.sub(r"^\s*N(?:J)?\d+(?:-\d+)?\s*", "", result, flags=re.IGNORECASE)
    return result.strip()


def build_sensitive_result_table(
    records: List[Dict[str, Any]],
    plan_rows: List[Dict[str, Any]],
) -> Dict[str, Any]:
    subtables: List[Dict[str, Any]] = []
    all_rows: List[Dict[str, Any]] = []
    all_warnings: List[str] = []
    type_groups = [
        ("railway_boundary_noise", {"railway_boundary_noise"}),
        ("area_traffic_noise", {"traffic_noise", "area_environment_noise"}),
    ]
    for noise_type, source_types in type_groups:
        type_records = [record for record in records if record.get("noise_type") in source_types]
        if not type_records:
            continue
        for subtable_noise_type, subtable_records in split_sensitive_records_by_flow_unit(noise_type, type_records):
            subtable = build_sensitive_result_subtable(subtable_records, plan_rows, subtable_noise_type)
            if subtable.get("rows"):
                subtables.append(subtable)
                all_rows.extend(subtable["rows"])
                all_warnings.extend(subtable.get("warnings") or [])
    if not subtables:
        subtable = build_sensitive_result_subtable(records, plan_rows, "area_environment_noise")
        subtables.append(subtable)
        all_rows.extend(subtable.get("rows") or [])
        all_warnings.extend(subtable.get("warnings") or [])
    return {
        "table_key": NOISE_TABLE_SENSITIVE,
        "caption_suffix": "项目沿线敏感点声环境现状监测平均值  单位：dB(A)",
        "title": "项目沿线敏感点声环境现状监测平均值  单位：dB(A)",
        "headers": result_headers(include_flow=has_traffic_flow(all_rows), vehicle_breakdown=has_vehicle_flow_breakdown(all_rows)),
        "rows": all_rows,
        "subtables": subtables,
        "flow_label": flow_label_from_rows(all_rows),
        "warnings": all_warnings,
    }


def build_sensitive_result_subtable(
    records: List[Dict[str, Any]],
    plan_rows: List[Dict[str, Any]],
    noise_type: str,
) -> Dict[str, Any]:
    matched = match_records(records, plan_rows)
    grouped = group_sensitive_records_for_result(matched)
    rows: List[Dict[str, Any]] = []
    warnings: List[str] = []
    debug_rows: List[Dict[str, Any]] = []
    for code in sorted(grouped, key=point_sort_key):
        plan = grouped[code]["plan"]
        if plan.get("is_attenuation"):
            continue
        point_records = grouped[code]["records"]
        metrics = compute_point_metrics(point_records, plan)
        row = base_result_row(plan)
        row.update(metrics)
        flows = paired_records(point_records)
        row["traffic_flow_day1_day"] = flow_text(flows, 0, "day")
        row["traffic_flow_day1_night"] = flow_text(flows, 0, "night")
        row["traffic_flow_day2_day"] = flow_text(flows, 1, "day")
        row["traffic_flow_day2_night"] = flow_text(flows, 1, "night")
        row["_flow_label"] = flow_label_from_records(point_records)
        add_vehicle_flow_fields(row, flows)
        row["noise_type"] = noise_type
        row["noise_type_label"] = noise_type_label(noise_type)
        row["noise_types"] = sorted({record["noise_type"] for record in point_records})
        row["needs_review"] = bool(metrics.get("warnings"))
        row["warning"] = "；".join(metrics.get("warnings") or [])
        rows.append(row)
        warnings.extend(metrics.get("warnings") or [])
        debug_rows.append(result_name_position_debug_item(plan, point_records, noise_type_label(noise_type)))
    include_flow = has_traffic_flow(rows)
    if not include_flow:
        strip_traffic_flow_fields(rows)
    headers = result_headers(include_flow=include_flow, vehicle_breakdown=has_vehicle_flow_breakdown(rows))
    return {
        "table_key": f"{NOISE_TABLE_SENSITIVE}_{noise_type}",
        "caption_suffix": f"{noise_type_label(noise_type)}监测结果  单位：dB(A)",
        "title": f"{noise_type_label(noise_type)}监测结果  单位：dB(A)",
        "noise_type": noise_type,
        "noise_type_label": noise_type_label(noise_type),
        "headers": headers,
        "rows": rows,
        "flow_label": flow_label_from_rows(rows),
        "warnings": warnings,
        "name_position_debug": debug_rows,
    }


def split_sensitive_records_by_flow_unit(
    noise_type: str,
    records: List[Dict[str, Any]],
) -> List[Tuple[str, List[Dict[str, Any]]]]:
    if noise_type != "area_traffic_noise":
        return [(noise_type, records)]
    road_or_area: List[Dict[str, Any]] = []
    rail_flow: List[Dict[str, Any]] = []
    for record in records:
        label = canonical_flow_label(record.get("traffic_flow_unit"))
        if has_record_traffic_flow(record) and "列/60min" in label:
            rail_flow.append(record)
        else:
            road_or_area.append(record)
    result: List[Tuple[str, List[Dict[str, Any]]]] = []
    if road_or_area:
        result.append(("area_traffic_noise", road_or_area))
    if rail_flow:
        result.append(("area_traffic_train_flow_noise", rail_flow))
    return result


def noise_type_label(noise_type: str) -> str:
    return {
        "railway_boundary_noise": "铁路边界噪声",
        "traffic_noise": "交通噪声",
        "area_environment_noise": "区域环境噪声",
        "area_traffic_noise": "区域环境及交通噪声",
        "area_traffic_train_flow_noise": "列车经过交通噪声",
    }.get(noise_type, "声环境噪声")


def build_attenuation_result_table(
    records: List[Dict[str, Any]],
    plan_rows: List[Dict[str, Any]],
) -> Dict[str, Any]:
    matched = match_records(records, plan_rows)
    grouped = group_records_by_point(matched)
    rows: List[Dict[str, Any]] = []
    warnings: List[str] = []
    debug_rows: List[Dict[str, Any]] = []
    for code in sorted(grouped, key=point_sort_key):
        plan = grouped[code]["plan"]
        if not plan.get("is_attenuation"):
            continue
        point_records = grouped[code]["records"]
        metrics = compute_point_metrics(point_records, plan)
        row = base_result_row(plan)
        row["监测点位置"] = attenuation_position_for_record(point_records[0], plan)
        row.update(metrics)
        flows = paired_records(point_records)
        row["traffic_flow_day1_day"] = flow_text(flows, 0, "day")
        row["traffic_flow_day1_night"] = flow_text(flows, 0, "night")
        row["traffic_flow_day2_day"] = flow_text(flows, 1, "day")
        row["traffic_flow_day2_night"] = flow_text(flows, 1, "night")
        row["_flow_label"] = flow_label_from_records(point_records)
        add_vehicle_flow_fields(row, flows)
        row["needs_review"] = bool(metrics.get("warnings"))
        row["warning"] = "；".join(metrics.get("warnings") or [])
        rows.append(row)
        warnings.extend(metrics.get("warnings") or [])
        debug_rows.append(result_name_position_debug_item(plan, point_records, "交通噪声衰减断面"))
    include_flow = has_traffic_flow(rows)
    if not include_flow:
        strip_traffic_flow_fields(rows)
    headers = result_headers(include_flow=include_flow, vehicle_breakdown=has_vehicle_flow_breakdown(rows))
    return {
        "table_key": NOISE_TABLE_ATTENUATION,
        "caption_suffix": "现状公路交通噪声衰减断面监测结果  单位：dB(A)",
        "title": "现状公路交通噪声衰减断面监测结果  单位：dB(A)",
        "headers": headers,
        "rows": rows,
        "flow_label": flow_label_from_rows(rows),
        "warnings": warnings,
        "name_position_debug": debug_rows,
    }


def result_name_position_debug_item(plan: Dict[str, Any], records: List[Dict[str, Any]], table_label: str) -> Dict[str, Any]:
    first_record = records[0] if records else {}
    marker, marker_source = record_indoor_outdoor_marker(first_record, plan)
    return {
        "table": table_label,
        "point_code": plan.get("point_code"),
        "raw_plan_name": plan.get("point_name"),
        "raw_plan_position": plan.get("position"),
        "raw_report_point": first_record.get("point_text"),
        "source_headers": first_record.get("source_headers") or [],
        "noise_type": first_record.get("noise_type"),
        "final_name": plan.get("point_name"),
        "final_position": plan.get("position"),
        "indoor_outdoor_marker": marker,
        "indoor_outdoor_source": marker_source,
    }


def collect_result_name_position_debug(table2: Dict[str, Any], table3: Dict[str, Any]) -> List[Dict[str, Any]]:
    items: List[Dict[str, Any]] = []
    for subtable in table2.get("subtables") or []:
        items.extend(subtable.get("name_position_debug") or [])
    items.extend(table3.get("name_position_debug") or [])
    return items


def is_blank_table_value(value: Any) -> bool:
    return str(value or "").strip() in {"", "-", "/", "—", "无", "None", "none", "null"}


def build_noise_plan_meta_warnings(plan_rows: List[Dict[str, Any]]) -> List[str]:
    warnings: List[str] = []
    if any(not is_valid_noise_plan_meta("factor", row.get("factor", "")) for row in plan_rows):
        warnings.append("监测因子未识别")
    if any(not is_valid_noise_plan_meta("frequency", row.get("frequency", "")) for row in plan_rows):
        warnings.append("监测频次未识别")
    status_path = DEBUG_DIR / "noise_plan_meta_status.json"
    if status_path.exists():
        try:
            status = read_json(status_path)
            for warning in status.get("warnings") or []:
                warning_text = str(warning or "").strip()
                if warning_text and warning_text not in warnings:
                    warnings.append(warning_text)
        except Exception:
            pass
    return warnings


def build_indoor_result_table(table2: Dict[str, Any]) -> Dict[str, Any]:
    indoor_rows: List[Dict[str, Any]] = []
    outdoor_subtables: List[Dict[str, Any]] = []
    outdoor_rows_all: List[Dict[str, Any]] = []
    outdoor_warnings: List[str] = []
    source_subtables = table2.get("subtables") or [
        {
            "table_key": table2.get("table_key", NOISE_TABLE_SENSITIVE),
            "caption_suffix": table2.get("caption_suffix", ""),
            "title": table2.get("title", ""),
            "noise_type": table2.get("noise_type"),
            "noise_type_label": table2.get("noise_type_label"),
            "headers": table2.get("headers") or [],
            "rows": table2.get("rows") or [],
            "flow_label": table2.get("flow_label", ""),
            "warnings": table2.get("warnings") or [],
            "name_position_debug": table2.get("name_position_debug") or [],
        }
    ]
    for subtable in source_subtables:
        subtable_outdoor_rows: List[Dict[str, Any]] = []
        for row in subtable.get("rows") or []:
            if is_indoor_result_row(row):
                indoor_rows.append(row)
            else:
                subtable_outdoor_rows.append(row)
        if not subtable_outdoor_rows:
            continue
        include_subtable_flow = has_traffic_flow(subtable_outdoor_rows)
        if not include_subtable_flow:
            strip_traffic_flow_fields(subtable_outdoor_rows)
        updated_subtable = dict(subtable)
        updated_subtable["rows"] = subtable_outdoor_rows
        updated_subtable["headers"] = result_headers(include_flow=include_subtable_flow, vehicle_breakdown=has_vehicle_flow_breakdown(subtable_outdoor_rows))
        updated_subtable["flow_label"] = flow_label_from_rows(subtable_outdoor_rows)
        updated_subtable["warnings"] = [
            warning
            for row in subtable_outdoor_rows
            for warning in (row.get("warnings") or [])
        ]
        outdoor_subtables.append(updated_subtable)
        outdoor_rows_all.extend(subtable_outdoor_rows)
        outdoor_warnings.extend(updated_subtable["warnings"])

    table2["subtables"] = outdoor_subtables
    table2["rows"] = outdoor_rows_all
    include_flow = has_traffic_flow(outdoor_rows_all)
    if not include_flow:
        strip_traffic_flow_fields(outdoor_rows_all)
    table2["headers"] = result_headers(include_flow=include_flow, vehicle_breakdown=has_vehicle_flow_breakdown(outdoor_rows_all))
    table2["flow_label"] = flow_label_from_rows(outdoor_rows_all)
    table2["warnings"] = outdoor_warnings

    include_indoor_flow = has_traffic_flow(indoor_rows)
    if not include_indoor_flow:
        strip_traffic_flow_fields(indoor_rows)
    return {
        "table_key": NOISE_TABLE_INDOOR,
        "caption_suffix": "室内噪声监测结果统计  单位：dB(A)",
        "title": "室内噪声监测结果统计  单位：dB(A)",
        "headers": result_headers(include_flow=include_indoor_flow, vehicle_breakdown=has_vehicle_flow_breakdown(indoor_rows)),
        "rows": indoor_rows,
        "flow_label": flow_label_from_rows(indoor_rows),
        "warnings": [
            warning
            for row in indoor_rows
            for warning in (row.get("warnings") or [])
        ],
    }


def flow_label_from_records(records: List[Dict[str, Any]]) -> str:
    labels = []
    for record in records:
        value = str(record.get("traffic_flow") or "").strip()
        label = canonical_flow_label(record.get("traffic_flow_unit"))
        if value and value not in {"-", "/", "—", "无", "None", "none", "null"} and label and label not in labels:
            labels.append(label)
    if len(labels) == 1:
        return labels[0]
    if len(labels) > 1:
        return "交通量"
    return ""


def flow_label_from_rows(rows: List[Dict[str, Any]]) -> str:
    labels = []
    for row in rows:
        label = canonical_flow_label(row.get("_flow_label"))
        if label and label not in labels:
            labels.append(label)
    if len(labels) == 1:
        return labels[0]
    if len(labels) > 1:
        return "交通量"
    return "车流量（辆/20min）"


def has_record_traffic_flow(record: Dict[str, Any]) -> bool:
    value = str(record.get("traffic_flow") or "").strip()
    return value not in {"", "-", "/", "—", "无", "None", "none", "null"}


def canonical_flow_label(label: Any) -> str:
    text = re.sub(r"\s+", "", str(label or ""))
    if not text:
        return ""
    if "列/60min" in text or "列车流量" in text or "列车" in text:
        return "列车流量（列/60min）"
    if "辆/20min" in text or "车流量" in text:
        return "车流量（辆/20min）"
    return str(label or "").strip()


def build_flow_unit_warnings(records: List[Dict[str, Any]]) -> List[str]:
    labels = sorted(
        {
            canonical_flow_label(record.get("traffic_flow_unit"))
            for record in records
            if has_record_traffic_flow(record)
            and canonical_flow_label(record.get("traffic_flow_unit"))
        }
    )
    groups: Dict[str, int] = {}
    raw_labels: Dict[str, List[str]] = {}
    for record in records:
        if not has_record_traffic_flow(record):
            continue
        label = canonical_flow_label(record.get("traffic_flow_unit"))
        if not label:
            continue
        groups[label] = groups.get(label, 0) + 1
        raw = str(record.get("traffic_flow_unit") or "").strip()
        raw_labels.setdefault(label, [])
        if raw and raw not in raw_labels[label]:
            raw_labels[label].append(raw)
    write_json(DEBUG_DIR / "noise_flow_unit_status.json", {"labels": labels, "groups": groups, "raw_labels": raw_labels})
    if len(labels) > 1:
        return ["噪声监测结果存在多个交通量单位: " + "、".join(labels)]
    return []


def is_indoor_result_row(row: Dict[str, Any]) -> bool:
    position = str(row.get("监测点位置") or "")
    return "室内" in position


def road_vehicle_flow_headers() -> List[str]:
    return [
        f"{base}_{vehicle}"
        for base in LEGACY_TRAFFIC_FLOW_FIELDS
        for vehicle in ROAD_VEHICLE_KEYS
    ]


def has_vehicle_flow_breakdown(rows: List[Dict[str, Any]]) -> bool:
    empty_values = {"", "-", "/", "\u2014", "None", "none", "null"}
    return any(
        str(row.get(header) or "").strip() not in empty_values
        for row in rows for header in road_vehicle_flow_headers()
    )

def result_headers(include_flow: bool = True, vehicle_breakdown: bool = False) -> List[str]:
    headers = [
        "监测点编号",
        "监测点名称",
        "监测点位置",
        "day1_day",
        "day1_night",
        "day2_day",
        "day2_night",
        "avg_day",
        "avg_night",
        "standard_day",
        "standard_night",
        "exceed_day",
        "exceed_night",
    ]
    if include_flow and vehicle_breakdown:
        headers.extend(road_vehicle_flow_headers())
        return headers
    if include_flow:
        headers.extend(
            [
                "traffic_flow_day1_day",
                "traffic_flow_day1_night",
                "traffic_flow_day2_day",
                "traffic_flow_day2_night",
            ]
        )
    return headers


def has_traffic_flow(rows: List[Dict[str, Any]]) -> bool:
    flow_headers = [
        "traffic_flow_day1_day",
        "traffic_flow_day1_night",
        "traffic_flow_day2_day",
        "traffic_flow_day2_night",
    ]
    empty_values = {"", "-", "/", "—", "无", "None", "none", "null"}
    for row in rows:
        for header in flow_headers:
            value = str(row.get(header) or "").strip()
            if value not in empty_values:
                return True
    return False


def strip_traffic_flow_fields(rows: List[Dict[str, Any]]) -> None:
    for row in rows:
        for header in (
            "traffic_flow_day1_day",
            "traffic_flow_day1_night",
            "traffic_flow_day2_day",
            "traffic_flow_day2_night",
        ):
            row.pop(header, None)
        for header in road_vehicle_flow_headers():
            row.pop(header, None)


def match_records(records: List[Dict[str, Any]], plan_rows: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    result = []
    for record in records:
        plan = match_plan(record, plan_rows)
        item = dict(record)
        item["plan"] = plan
        result.append(item)
    return result


def validate_plan_report_match(records: List[Dict[str, Any]], plan_rows: List[Dict[str, Any]]) -> List[str]:
    matched = match_records(records, plan_rows)
    plan_row_ids = {id(row) for row in plan_rows}
    matched_plan_ids = {
        id(record.get("plan"))
        for record in matched
        if id(record.get("plan")) in plan_row_ids
    }
    missing_plan_rows = [
        {
            "point_code": row.get("point_code"),
            "point_name": row.get("point_name"),
            "position": row.get("position"),
            "standard_class": row.get("standard_class"),
        }
        for row in plan_rows
        if id(row) not in matched_plan_ids
    ]
    missing_in_report = sorted(
        {str(row.get("point_code") or "") for row in missing_plan_rows if row.get("point_code")},
        key=point_sort_key,
    )
    extra_in_report = sorted(
        {
            str(record.get("raw_point_code") or record.get("point_text") or "")
            for record in matched
            if id(record.get("plan")) not in plan_row_ids
        },
        key=point_sort_key,
    )
    unmatched_records = [
        {
            "point_text": record.get("point_text"),
            "raw_point_code": record.get("raw_point_code"),
            "noise_type": record.get("noise_type"),
        }
        for record in matched
        if id(record.get("plan")) not in plan_row_ids
    ]
    payload = {
        "missing_in_report": missing_in_report,
        "missing_plan_rows": missing_plan_rows,
        "extra_in_report": extra_in_report,
        "unmatched_records": unmatched_records,
        "plan_count": len(plan_rows),
        "matched_plan_count": len(matched_plan_ids),
        "report_record_count": len(records),
    }
    write_json(DEBUG_DIR / "noise_plan_report_mismatch.json", payload)
    warnings: List[str] = []
    for row in missing_plan_rows:
        warnings.append(
            f"监测方案点位{row.get('point_code')}（{row.get('point_name')}，{row.get('position')}）"
            "未在噪声监测报告中匹配到对应数据"
        )
    for record in unmatched_records:
        warnings.append(
            f"噪声监测报告点位{record.get('raw_point_code') or record.get('point_text')}未匹配到监测方案"
        )
    return warnings


def build_plan_report_position_warnings(
    records: List[Dict[str, Any]],
    plan_rows: List[Dict[str, Any]],
) -> List[str]:
    warnings: List[str] = []
    seen_codes = set()
    for plan in plan_rows:
        code = str(plan.get("point_code") or "").strip()
        if not code or code in seen_codes or plan.get("is_attenuation"):
            continue
        seen_codes.add(code)
        plan_floors = floor_tokens(str(plan.get("position") or ""))
        report_floors = set()
        for record in records:
            if str(record.get("raw_point_code") or "").strip() != code:
                continue
            if "背景点" in str(record.get("point_text") or ""):
                continue
            report_floors.update(floor_tokens(str(record.get("point_text") or "")))
        comparable_plan_floors, comparable_report_floors = normalize_top_floor_equivalence(
            plan_floors,
            report_floors,
        )
        if comparable_plan_floors and comparable_report_floors and comparable_plan_floors != comparable_report_floors:
            warnings.append(
                f"{code} 监测方案楼层({format_floor_tokens(plan_floors)})"
                f"与监测报告楼层({format_floor_tokens(report_floors)})不一致"
            )
    for plan in plan_rows:
        if plan.get("is_attenuation"):
            continue
        code = str(plan.get("point_code") or "").strip()
        matching_record = next(
            (
                record for record in records
                if str(record.get("raw_point_code") or "").strip() == code
                and "\u80cc\u666f\u70b9" not in str(record.get("point_text") or "")
            ),
            None,
        )
        if matching_record is None:
            continue
        plan_context = strip_plan_floor_terms(str(plan.get("position") or ""))
        plan_context = plan_context.replace("\u5ba4\u5916", "").replace("\u5ba4\u5185", "").strip()
        if not plan_context:
            continue
        final_position = monitor_point_display_position(plan, matching_record)
        normalized_context = re.sub(r"\s+", "", normalize_position_text(plan_context))
        normalized_final = re.sub(r"\s+", "", normalize_position_text(final_position))
        if normalized_context not in normalized_final:
            warnings.append(
                f"{code} \u6700\u7ec8\u76d1\u6d4b\u70b9\u4f4d\u7f6e\u672a\u4fdd\u7559\u65b9\u6848\u4f4d\u7f6e\u5173\u7cfb: {plan_context}"
            )
    return warnings


def normalize_top_floor_equivalence(plan_floors: set, report_floors: set) -> Tuple[set, set]:
    normalized_plan = set(plan_floors)
    normalized_report = set(report_floors)
    if "顶层" in normalized_report:
        numeric_plan = [item for item in normalized_plan if str(item).isdigit()]
        if numeric_plan:
            normalized_report.remove("顶层")
            normalized_report.add(max(numeric_plan, key=int))
    return normalized_plan, normalized_report


def format_floor_tokens(tokens: set) -> str:
    numeric = sorted((item for item in tokens if str(item).isdigit()), key=int)
    other = sorted(item for item in tokens if not str(item).isdigit())
    return "、".join([*(f"{item}层" for item in numeric), *other])


def match_plan(record: Dict[str, Any], plan_rows: List[Dict[str, Any]]) -> Dict[str, Any]:
    raw_code = record.get("raw_point_code")
    point_text = record.get("point_text") or ""
    background_plan = match_background_point_plan(point_text, plan_rows)
    if background_plan:
        return background_plan
    if raw_code:
        by_detail_position = match_plan_by_floor_position(
            point_text,
            related_plan_candidates(raw_code, plan_rows),
            require_detail_match=True,
        )
        if by_detail_position:
            return by_detail_position
        candidates = [row for row in plan_rows if row["point_code"] == raw_code]
        if len(candidates) == 1:
            return candidates[0]
        if len(candidates) > 1:
            by_position = match_plan_by_floor_position(point_text, candidates)
            if by_position:
                return by_position
            if plan_rows_equivalent(candidates):
                return candidates[0]
        by_position = match_plan_by_floor_position(point_text, related_plan_candidates(raw_code, plan_rows))
        if by_position:
            return by_position
        if "-" in raw_code:
            parent_code = raw_code.split("-", 1)[0]
            candidates = [row for row in plan_rows if row["point_code"] == parent_code]
            if len(candidates) == 1:
                return candidates[0]
        candidates = [row for row in plan_rows if row["point_code"].startswith(raw_code + "-")]
        if len(candidates) == 1:
            return candidates[0]
        if len(candidates) > 1:
            by_name = [row for row in candidates if row["point_name"] and row["point_name"] in point_text]
            if by_name:
                return by_name[0]
            by_name = [
                row for row in candidates
                if row["point_name"] and len(row["point_name"].rstrip("1234567890")) >= 2 and row["point_name"].rstrip("1234567890") in point_text
            ]
            if by_name:
                return by_name[0]
    for row in plan_rows:
        name = row["point_name"]
        station = row["station"]
        if name and name in point_text:
            return row
        if station and station in point_text and name and name.rstrip("12") in point_text:
            return row
    if is_attenuation_text(point_text):
        distance = extract_distance(point_text)
        for row in plan_rows:
            if row["is_attenuation"] and distance and distance in row["position"]:
                return row
    return {
        "point_code": raw_code or "",
        "station": "",
        "point_name": point_text,
        "district": "",
        "position": point_text,
        "standard_class": "",
        "factor": "",
        "frequency": "",
        "is_attenuation": is_attenuation_text(point_text),
    }


def plan_rows_equivalent(rows: List[Dict[str, Any]]) -> bool:
    if not rows:
        return False
    keys = (
        "point_code",
        "station",
        "point_name",
        "district",
        "position",
        "standard_class",
        "factor",
        "frequency",
        "is_attenuation",
    )
    signatures = {
        tuple(str(row.get(key) or "").strip() for key in keys)
        for row in rows
    }
    return len(signatures) == 1


def related_plan_candidates(raw_code: str, plan_rows: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    if not raw_code:
        return []
    parent_code = raw_code.split("-", 1)[0]
    return [
        row
        for row in plan_rows
        if row.get("point_code") == raw_code
        or str(row.get("point_code") or "").startswith(parent_code + "-")
    ]


def match_plan_by_floor_position(
    point_text: str,
    candidates: List[Dict[str, Any]],
    require_detail_match: bool = False,
) -> Optional[Dict[str, Any]]:
    point_floors = floor_tokens(point_text)
    if not point_floors or not candidates:
        return None
    scored: List[Tuple[int, int, int, Dict[str, Any]]] = []
    any_detail_match = any(plan_detail_matches_point_text(row, point_text) for row in candidates)
    for row in candidates:
        detail_match = plan_detail_matches_point_text(row, point_text)
        if require_detail_match and not detail_match:
            continue
        if any_detail_match and row.get("code_detail") and not detail_match:
            continue
        row_floors = floor_tokens(str(row.get("position") or ""))
        score = len(point_floors & row_floors)
        if score:
            scored.append((score, 1 if detail_match else 0, -len(row_floors), row))
    if not scored:
        return None
    scored.sort(key=lambda item: (item[0], item[1], item[2], point_sort_key(item[3].get("point_code"))), reverse=True)
    return scored[0][3]


def extract_code_detail(text: str) -> str:
    value = str(text or "")
    match = re.search(r"[（(]\s*([^）)]+?)\s*[）)]", value)
    return normalize_plan_detail(match.group(1)) if match else ""


def normalize_plan_detail(text: str) -> str:
    return re.sub(r"\s+", "", str(text or "").strip())


def plan_detail_matches_point_text(row: Dict[str, Any], point_text: str) -> bool:
    detail = normalize_plan_detail(str(row.get("code_detail") or ""))
    if not detail:
        return False
    text = normalize_plan_detail(point_text)
    return detail in text


def floor_tokens(text: str) -> set:
    value = re.sub(r"\s+", "", str(text or ""))
    tokens = set()
    if re.fullmatch(r"\d+", value):
        tokens.add(value)
    tokens.update(re.findall(r"\d+(?=[层楼])", value))
    for match in re.finditer(r"((?:\d+[、,，/和及])+(\d+))[层楼]", value):
        tokens.update(re.findall(r"\d+", match.group(1)))
    if "顶层" in value:
        tokens.add("顶层")
    return tokens


def match_background_point_plan(point_text: str, plan_rows: List[Dict[str, Any]]) -> Optional[Dict[str, Any]]:
    text = str(point_text or "")
    if "背景点" not in text:
        return None
    background_candidates = [
        row
        for row in plan_rows
        if any(
            marker in " ".join(str(row.get(key) or "") for key in ("point_name", "position", "station"))
            for marker in ("背景点", "中部", "空旷", "区域环境")
        )
    ]
    if not background_candidates:
        return None
    text_tokens = set(re.findall(r"[\u4e00-\u9fff]{2,}", text))
    scored: List[Tuple[int, Dict[str, Any]]] = []
    for row in background_candidates:
        row_text = " ".join(str(row.get(key) or "") for key in ("point_name", "position", "station"))
        score = sum(1 for token in text_tokens if token and token in row_text)
        if row.get("point_name") and str(row.get("point_name")).replace("小区中部", "") in text:
            score += 5
        scored.append((score, row))
    scored = sorted(scored, key=lambda item: (item[0], point_sort_key(item[1].get("point_code"))), reverse=True)
    return scored[0][1] if scored and scored[0][0] > 0 else None


def group_records_by_point(records: List[Dict[str, Any]]) -> Dict[str, Dict[str, Any]]:
    grouped: "OrderedDict[str, Dict[str, Any]]" = OrderedDict()
    for record in records:
        plan = record["plan"]
        code = plan.get("point_code") or record.get("raw_point_code") or record.get("point_text")
        if code not in grouped:
            grouped[code] = {"plan": plan, "records": []}
        grouped[code]["records"].append(record)
    for value in grouped.values():
        value["records"].sort(key=noise_record_sort_key)
    return grouped


def group_sensitive_records_for_result(records: List[Dict[str, Any]]) -> Dict[str, Dict[str, Any]]:
    grouped: "OrderedDict[str, Dict[str, Any]]" = OrderedDict()
    for record in records:
        plan = record["plan"]
        raw_code = str(record.get("raw_point_code") or "").strip()
        plan_code = str(plan.get("point_code") or "").strip()
        remapped_background = is_remapped_background_record(record, plan)
        display_code = plan_code if remapped_background else raw_code or plan_code or record.get("point_text")
        group_key = sensitive_result_group_key(record, plan)
        if group_key not in grouped:
            result_plan = dict(plan)
            result_plan["source_plan_code"] = plan_code
            result_plan["point_code"] = display_code
            result_plan["point_name"] = clean_result_point_name(record, plan)
            result_plan["position"] = clean_result_position(
                result_position_for_record(record, plan),
                record,
                plan,
            )
            result_plan["result_sort_key"] = group_key
            grouped[group_key] = {"plan": result_plan, "records": []}
        grouped[group_key]["records"].append(record)
    for value in grouped.values():
        value["records"].sort(key=noise_record_sort_key)
    return grouped


def noise_record_sort_key(record: Dict[str, Any]) -> Tuple[int, int]:
    table_order = record.get("source_table_order")
    source_order = record.get("source_order")
    return (
        int(table_order) if table_order is not None else 9999,
        int(source_order) if source_order is not None else 0,
    )


def is_remapped_background_record(record: Dict[str, Any], plan: Dict[str, Any]) -> bool:
    point_text = str(record.get("point_text") or "")
    raw_code = str(record.get("raw_point_code") or "").strip()
    plan_code = str(plan.get("point_code") or "").strip()
    return "背景点" in point_text and bool(raw_code and plan_code and raw_code != plan_code)


def sensitive_result_group_key(record: Dict[str, Any], plan: Dict[str, Any]) -> str:
    raw_code = str(record.get("raw_point_code") or plan.get("point_code") or record.get("point_text") or "").strip()
    if is_remapped_background_record(record, plan):
        raw_code = str(plan.get("point_code") or raw_code)
    report_sample_code = extract_report_sample_code(str(record.get("point_text") or ""))
    if report_sample_code:
        return f"{raw_code}|{report_sample_code}"
    position = result_position_for_record(record, plan)
    return f"{raw_code}|{position}"


def extract_report_sample_code(text: str) -> str:
    match = re.search(r"点位编号\s*([A-Za-z]\d+(?:-\d+)?)", str(text or ""), flags=re.IGNORECASE)
    return match.group(1).upper() if match else ""


def result_position_for_record(record: Dict[str, Any], plan: Dict[str, Any]) -> str:
    point_text = str(record.get("point_text") or "")
    marker, _marker_source = record_indoor_outdoor_marker(record, plan)
    if is_background_record(record):
        return apply_indoor_outdoor_marker(background_position_for_record(record, plan), marker)
    report_position = extract_position_from_point_text(point_text, record, plan)
    position = compose_plan_monitor_position(
        plan.get("position"),
        report_position or point_text,
    )
    return apply_indoor_outdoor_marker(position, marker)


def clean_result_point_name(record: Dict[str, Any], plan: Dict[str, Any]) -> str:
    return clean_result_table_point_name(record, plan)


def sanitize_result_point_name(value: Any, plan: Dict[str, Any]) -> str:
    if is_railway_boundary_point_name(plan.get("point_name")):
        return str(plan.get("point_name") or "").strip()
    text = str(value or "").strip()
    if not text:
        return ""
    station = str(plan.get("station") or "").strip()
    position = str(plan.get("position") or "").strip()
    if station:
        text = text.replace(station, "")
    if position and position in text:
        text = text.replace(position, "")
    text = re.sub(r"K\s*\d+\s*[+-]\s*\d+", "", text, flags=re.IGNORECASE)
    text = re.sub(r"点位编号\s*[A-Za-z]\d+(?:-\d+)?", "", text, flags=re.IGNORECASE)
    text = re.sub(r"\d+\s*[层楼](?:室外|室内)?(?:监测)?", "", text)
    text = re.sub(rf"(?:面向|背向)[^，。；;\r\n]*?(?:{ROW_POSITION_MARKER_PATTERN}|本项目|道路|铁路|陇海线)[^，。；;\r\n]*", "", text)
    text = re.sub(r"(?:敏感点|背景点|衰减断面|室外监测|室内监测|噪声监测)", "", text)
    text = re.sub(r"\s+", " ", text).strip(" 、，,;；")
    return text.strip()


def is_railway_boundary_point_name(value: Any) -> bool:
    return "铁路边界" in str(value or "")


def clean_result_position(value: Any, record: Dict[str, Any], plan: Dict[str, Any]) -> str:
    marker, _marker_source = record_indoor_outdoor_marker(record, plan)
    text = str(value or "").strip()
    if not text:
        _report_name, report_position = split_monitor_point_identity(record.get("point_text"), plan, record)
        text = report_position
    if not text:
        return "-"
    point_text = str(record.get("point_text") or "")
    raw_code = str(record.get("raw_point_code") or "").strip()
    station = str(plan.get("station") or "").strip()
    point_name = str(plan.get("point_name") or "").strip()
    if text == point_text:
        text = strip_point_identity(text, raw_code, station, point_name)
    else:
        text = strip_point_identity(text, raw_code, station, point_name)
    text = re.sub(r"点位编号\s*[A-Za-z]*\d+(?:-\d+)?", "", text, flags=re.IGNORECASE)
    text = normalize_position_text(text)
    text = re.sub(r"\s+", "", text).strip("、，,;；")
    if is_background_record(record):
        return normalize_floor_position_display(apply_indoor_outdoor_marker(append_background_marker(extract_floor_position_fragment(text) or text), marker))
    if not has_orientation_marker(text):
        text = extract_floor_position_fragment(text) or text
    return normalize_floor_position_display(apply_indoor_outdoor_marker(text or "-", marker))


def extract_position_from_point_text(text: str, record: Dict[str, Any], plan: Dict[str, Any]) -> str:
    cleaned = strip_point_identity(
        str(text or ""),
        str(record.get("raw_point_code") or "").strip(),
        str(plan.get("station") or "").strip(),
        str(plan.get("point_name") or "").strip(),
    )
    cleaned = re.sub(r"点位编号\s*[A-Za-z]*\d+(?:-\d+)?", "", cleaned, flags=re.IGNORECASE)
    cleaned = normalize_position_text(cleaned)
    cleaned = re.sub(r"\s+", "", cleaned).strip("、，,;；")
    if re.search(rf"(面向|背向|{ROW_POSITION_MARKER_PATTERN}|[0-9]+[层楼]|顶层|室外|室内)", cleaned):
        if not has_orientation_marker(cleaned):
            return extract_floor_position_fragment(cleaned) or cleaned
        return cleaned
    return ""


def normalize_position_text(text: str) -> str:
    value = str(text or "")
    value = re.sub(r"(室外|室内)\s*监测", r"\1", value)
    value = re.sub(r"(室外|室内)\s*噪声\s*监测", r"\1", value)
    value = re.sub(r"(?:噪声监测|监测|进行)", "", value)
    value = re.sub(r"敏感点", "", value)
    return value


def is_background_record(record: Dict[str, Any]) -> bool:
    return "背景点" in str(record.get("point_text") or "")


def background_position_for_record(record: Dict[str, Any], plan: Dict[str, Any]) -> str:
    point_text = str(record.get("point_text") or "")
    report_floor = extract_floor_position(point_text)
    if report_floor:
        return append_background_marker(report_floor)
    plan_floor = first_floor_position_from_plan(str(plan.get("position") or ""))
    if plan_floor:
        return append_background_marker(plan_floor)
    return "背景点"


def append_background_marker(position: str) -> str:
    value = str(position or "").strip()
    if not value:
        return "背景点"
    if "背景点" in value:
        return value
    return f"{value}背景点"


def complete_floor_position_from_plan(text: str) -> str:
    value = str(text or "").strip()
    if not value:
        return ""
    compact = re.sub(r"\s+", "", value)
    blocked_markers = ("距离", "距", "断面", "红线", "中心线")
    if any(marker in compact for marker in blocked_markers):
        return ""
    outdoor = "室外"
    indoor = "室内"
    floor = "层"
    suffix = outdoor if outdoor in compact else indoor if indoor in compact else ""
    if compact.startswith(outdoor):
        suffix = outdoor
    elif compact.startswith(indoor):
        suffix = indoor
    numbers: List[str] = []
    multi_match = re.search(r"((?:\d+[、,，和及])+\d+)层", compact)
    if multi_match:
        numbers = re.findall(r"\d+", multi_match.group(1))
    elif re.fullmatch(r"(?:敏感点)?\d+层(?:室外|室内)?(?:监测)?", compact):
        numbers = re.findall(r"\d+", compact)
    elif re.fullmatch(r"(?:室外|室内)?\d+层(?:监测)?", compact):
        numbers = re.findall(r"\d+", compact)
    if numbers:
        floors = [f"{number}{floor}" for number in numbers]
        if suffix:
            floors[-1] = f"{floors[-1]}{suffix}"
        return "、".join(floors)
    if "顶层" in compact:
        return f"顶层{suffix}" if suffix and suffix not in compact else "顶层"
    return ""


def first_floor_position_from_plan(text: str) -> str:
    value = str(text or "")
    multi_floor = re.search(r"((?:\d+\s*[、,，和及]\s*)+\d+)\s*[层楼]", value)
    if multi_floor:
        first = re.search(r"\d+", multi_floor.group(1))
        if first:
            suffix = "室外" if "室外" in value else "室内" if "室内" in value else ""
            return f"{first.group(0)}层{suffix}"
    match = re.search(r"(\d+\s*[层楼](?:室外|室内)?|顶层(?:室外|室内)?)", value)
    if match:
        return re.sub(r"\s+", "", match.group(1))
    if "室外" in value:
        return "室外"
    if "室内" in value:
        return "室内"
    return ""


def extract_floor_position_fragment(text: str) -> str:
    value = str(text or "")
    matches = re.findall(r"\d+\s*[层楼](?:室外|室内)?|顶层(?:室外|室内)?", value)
    if matches:
        return re.sub(r"\s+", "", matches[-1])
    if "室外" in value:
        return "室外"
    if "室内" in value:
        return "室内"
    return ""


def has_indoor_outdoor_marker(text: str) -> bool:
    return "室内" in str(text or "") or "室外" in str(text or "")


def record_indoor_outdoor_marker(record: Dict[str, Any], plan: Optional[Dict[str, Any]] = None) -> Tuple[str, str]:
    header_text = " ".join(
        str(item or "")
        for item in [
            record.get("laeq_key"),
            *(record.get("source_headers") or []),
            *((record.get("raw") or {}).keys()),
        ]
    )
    if "室内噪声级" in header_text or "室内" in header_text:
        return "室内", "result_header"
    if "室外噪声级" in header_text or "室外" in header_text:
        return "室外", "result_header"

    point_text = str(record.get("point_text") or "")
    if "室内" in point_text:
        return "室内", "point_text"
    if "室外" in point_text:
        return "室外", "point_text"

    plan_position = str((plan or {}).get("position") or "")
    if "室内" in plan_position:
        return "室内", "plan_position"
    if "室外" in plan_position:
        return "室外", "plan_position"
    return "", ""


def apply_indoor_outdoor_marker(position: Any, marker: str) -> str:
    text = normalize_floor_position_display(position)
    if not text or text == "-":
        return text or "-"
    marker = str(marker or "").strip()
    if marker not in {"室内", "室外"}:
        return text
    opposite = "室外" if marker == "室内" else "室内"
    if marker in text:
        return text
    if opposite in text:
        return text.replace(opposite, marker)
    if re.fullmatch(r"\d+\s*[层楼]|顶层", text):
        return f"{text}{marker}"
    floor = extract_floor_position_fragment(text)
    if floor and marker not in floor and opposite not in floor:
        return text.replace(floor, f"{floor}{marker}", 1)
    return f"{text}{marker}"


def has_orientation_marker(text: str) -> bool:
    return bool(re.search(rf"(面向|背向|{ROW_POSITION_MARKER_PATTERN}|本项目|道路|铁路|陇海线)", str(text or "")))


def enrich_floor_position_with_plan_context(position: str, plan: Dict[str, Any]) -> str:
    value = str(position or "").strip()
    plan_position = str(plan.get("position") or "").strip()
    if not value or has_indoor_outdoor_marker(value) or has_orientation_marker(value):
        return value
    if re.fullmatch(r"\d+\s*[层楼]|顶层", value) and has_orientation_marker(plan_position):
        context = strip_plan_floor_terms(plan_position)
        if context:
            return f"{context}{value}"
    return value


def strip_plan_floor_terms(text: str) -> str:
    value = str(text or "")
    value = re.sub(r"\d+(?:[、,，和及]\d+)*\s*[层楼](?:室外|室内)?(?:监测)?", "", value)
    value = normalize_position_text(value)
    value = re.sub(r"\s+", "", value).strip("、，,;；")
    return value


def strip_point_identity(text: str, raw_code: str, station: str, point_name: str) -> str:
    result = str(text or "")
    if raw_code:
        result = re.sub(rf"\b{re.escape(raw_code)}\b", "", result, flags=re.IGNORECASE)
    if station:
        result = result.replace(station, "")
    result = re.sub(r"K\s*\d+\s*[+-]\s*\d+", "", result, flags=re.IGNORECASE)
    if point_name:
        result = result.replace(point_name, "")
        base_name = re.sub(r"\d+$", "", point_name).strip()
        if len(base_name) >= 2:
            result = result.replace(base_name, "")
    return result


def attenuation_position_for_record(record: Dict[str, Any], plan: Dict[str, Any]) -> str:
    return clean_attenuation_point_position(plan, record)


def result_name_for_record(record: Dict[str, Any], plan: Dict[str, Any]) -> str:
    point_text = str(record.get("point_text") or "")
    raw_code = str(record.get("raw_point_code") or "")
    text = point_text
    if raw_code and text.startswith(raw_code):
        text = text[len(raw_code):].strip()
    text = re.sub(r"点位编号\s*[A-Za-z]\d+(?:-\d+)?", "", text, flags=re.IGNORECASE).strip()
    text = re.sub(r"(敏感点|背景点).*", "", text).strip()
    text = re.sub(r"\d+\s*层.*", "", text).strip()
    return text or plan.get("point_name") or "-"


def extract_floor_position(text: str) -> str:
    text = str(text or "")
    match = re.search(r"(\d+\s*层(?:进行)?(?:室外|室内)?(?:噪声监测)?)", text)
    if match:
        value = re.sub(r"\s+", "", match.group(1))
        return value.replace("进行", "").replace("噪声监测", "")
    if "顶层" in text:
        return "顶层"
    return ""


def split_floor_positions(text: str) -> List[str]:
    text = re.sub(r"\s+", "", str(text or ""))
    if not text or "层" not in text:
        return []
    suffix = "室外" if "室外" in text else "室内" if "室内" in text else ""
    prefix = text.split("层", 1)[0]
    numbers = re.findall(r"\d+", prefix)
    if len(numbers) > 1:
        return [f"{number}层{suffix}" for number in numbers]
    return [extract_floor_position(text)] if extract_floor_position(text) else []


def base_result_row(plan: Dict[str, Any]) -> Dict[str, Any]:
    return {
        "监测点编号": plan.get("point_code") or "-",
        "监测点名称": plan.get("point_name") or "-",
        "监测点位置": normalize_floor_position_display(plan.get("position") or "-"),
    }


def compute_point_metrics(records: List[Dict[str, Any]], plan: Dict[str, Any]) -> Dict[str, Any]:
    pairs = paired_records(records)
    warnings: List[str] = []
    day_values: List[Decimal] = []
    night_values: List[Decimal] = []
    result: Dict[str, Any] = {}
    for day_index in range(2):
        day_record = pairs[day_index].get("day") if day_index < len(pairs) else None
        night_record = pairs[day_index].get("night") if day_index < len(pairs) else None
        day_value = parse_decimal(day_record.get("laeq")) if day_record else None
        night_value = parse_decimal(night_record.get("laeq")) if night_record else None
        result[f"day{day_index + 1}_day"] = value_or_dash(day_value)
        result[f"day{day_index + 1}_night"] = value_or_dash(night_value)
        if day_value is not None:
            day_values.append(day_value)
        else:
            warnings.append(f"{plan.get('point_code')} 第{day_index + 1}天昼间缺失")
        if night_value is not None:
            night_values.append(night_value)
        else:
            warnings.append(f"{plan.get('point_code')} 第{day_index + 1}天夜间缺失")
    avg_day = quantize_noise_decimal(sum(day_values) / Decimal(len(day_values))) if day_values else None
    avg_night = quantize_noise_decimal(sum(night_values) / Decimal(len(night_values))) if night_values else None
    standard_day, standard_night = noise_standard_limits(plan)
    if standard_day is None or standard_night is None:
        warnings.append(f"{plan.get('point_code')} 未识别现状标准: {plan.get('standard_class_raw') or plan.get('standard_class')}")
    result["avg_day"] = value_or_dash(avg_day)
    result["avg_night"] = value_or_dash(avg_night)
    result["standard_day"] = value_or_dash(standard_day)
    result["standard_night"] = value_or_dash(standard_night)
    result["exceed_day"] = exceed_text(avg_day, standard_day)
    result["exceed_night"] = exceed_text(avg_night, standard_night)
    result["warnings"] = warnings
    return result


def paired_records(records: List[Dict[str, Any]]) -> List[Dict[str, Dict[str, Any]]]:
    day_records = [record for record in records if record.get("period") == "day"]
    night_records = [record for record in records if record.get("period") == "night"]
    unknown_records = [record for record in records if record.get("period") not in {"day", "night"}]
    pairs: List[Dict[str, Dict[str, Any]]] = []
    for index in range(2):
        pair: Dict[str, Dict[str, Any]] = {}
        if index < len(day_records):
            pair["day"] = day_records[index]
        if index < len(night_records):
            pair["night"] = night_records[index]
        if index < len(unknown_records):
            pair["unknown"] = unknown_records[index]
        pairs.append(pair)
    return pairs


def build_compliance_summary(table2: Dict[str, Any], table3: Dict[str, Any]) -> Dict[str, Any]:
    return {
        "sensitive": summarize_rows(table2["rows"]),
        "attenuation": summarize_rows(table3["rows"]),
    }


def summarize_rows(rows: List[Dict[str, Any]]) -> Dict[str, Any]:
    exceed_items = []
    exceed_row_count = 0
    for row in rows:
        row_exceeded = False
        for period, field in [("昼间", "exceed_day"), ("夜间", "exceed_night")]:
            value = row.get(field)
            numeric_value = parse_decimal(value)
            if numeric_value is not None and numeric_value > 0:
                row_exceeded = True
                exceed_items.append(
                    {
                        "point_code": row.get("监测点编号"),
                        "point_name": row.get("监测点名称"),
                        "period": period,
                        "exceed": value,
                    }
                )
        if row_exceeded:
            exceed_row_count += 1
    return {
        "total_count": len(rows),
        "exceed_count": exceed_row_count,
        "exceed_items": exceed_items,
    }


def register_noise_tables(
    numbering: DocxNumbering,
    table1: Dict[str, Any],
    table2: Dict[str, Any],
    table3: Dict[str, Any],
    table_indoor: Optional[Dict[str, Any]] = None,
) -> None:
    tables = [table1]
    sensitive_tables = table2.get("subtables") or [table2]
    tables.extend(sensitive_tables)
    if table_indoor and table_indoor.get("rows"):
        tables.append(table_indoor)
    if table3.get("rows"):
        tables.append(table3)
    for table in tables:
        table["title"] = numbering.register_table(table["table_key"], table["caption_suffix"])


def build_rule_texts(
    table1: Dict[str, Any],
    table2: Dict[str, Any],
    table3: Dict[str, Any],
    summary: Dict[str, Any],
    plan_rows: List[Dict[str, Any]],
    numbering: DocxNumbering,
) -> Dict[str, str]:
    monitoring_meta = summary.get("monitoring_meta") or {}
    return {
        "monitoring_plan_text": build_monitoring_plan_text(table1, plan_rows, numbering),
        "monitoring_result_text": build_monitoring_result_text(monitoring_meta),
        "sensitive_conclusion": sensitive_conclusion(summary["sensitive"]),
        "traffic_intro": build_traffic_intro_text(plan_rows, monitoring_meta, numbering),
        "attenuation_conclusion": attenuation_conclusion(summary["attenuation"], lead_only=False),
    }


def build_monitoring_plan_text(
    table1: Dict[str, Any],
    plan_rows: List[Dict[str, Any]],
    numbering: DocxNumbering,
) -> str:
    total = len(table1["rows"])
    attenuation_count = sum(1 for row in plan_rows if row.get("is_attenuation"))
    text = (
        "根据《环境影响评价技术导则 公路建设项目》（HJ1358-2024）要求，"
        "一级评价应对评价范围内具有代表性的敏感目标的声环境质量进行实测，并对实测结果进行评价。"
        "本项目根据不同路段，贯彻“以点代线、点线结合、以代表性区段为主、反馈全线”的原则，"
        "结合项目沿线敏感点分布、现状噪声源类型及空间特征，"
        f"共设置{total}个声环境监测点"
    )
    if attenuation_count:
        text += f"（含{attenuation_count}个交通噪声衰减断面监测点）"
    text += f"。声环境现状监测方案见{numbering.table_label(NOISE_TABLE_MONITOR_POINTS)}。"
    return text


def build_monitoring_result_text(monitoring_meta: Dict[str, Any]) -> str:
    unit = str(monitoring_meta.get("monitoring_unit") or "").strip()
    date_text = str(monitoring_meta.get("monitor_date_text") or "").strip()
    lead = ""
    if unit and date_text:
        lead = f"{unit}于{date_text}对本项目沿线各监测点位的环境噪声进行了监测。"
    elif unit:
        lead = f"{unit}对本项目沿线各监测点位的环境噪声进行了监测。"
    elif date_text:
        lead = f"本项目于{date_text}对沿线各监测点位开展环境噪声现状监测。"
    tail = (
        "具体测量时间段、测量仪器、测量方法均按规范要求执行。"
        "测量结果以等效连续A声级和统计噪声级给出，并以等效A声级作为最终评价量。"
    )
    return f"{lead}{tail}" if lead else tail


def build_traffic_intro_text(
    plan_rows: List[Dict[str, Any]],
    monitoring_meta: Dict[str, Any],
    numbering: DocxNumbering,
) -> str:
    label = numbering.table_label(NOISE_TABLE_ATTENUATION)
    road_hint = str(monitoring_meta.get("road_hint") or "").strip()
    if road_hint:
        return f"本次评价在{road_hint}开展交通噪声衰减断面监测，监测结果见{label}。"
    return f"本次评价对现状公路交通噪声衰减断面进行监测，监测结果见{label}。"


def build_noise_monitoring_meta(
    records: List[Dict[str, Any]],
    plan_rows: List[Dict[str, Any]],
) -> Dict[str, Any]:
    report_path = find_report_path()
    chunks = load_docx_chunks(report_path) if report_path else []
    monitoring_unit = extract_monitoring_unit(chunks)
    report_ranges = extract_report_noise_date_ranges(chunks)
    record_dates = collect_monitor_dates_from_records(records)
    record_ranges = date_tuples_to_ranges(record_dates)
    merged_ranges = merge_date_ranges([*report_ranges, *record_ranges])
    return {
        "monitoring_unit": monitoring_unit,
        "monitor_date_ranges": [
            {"start": format_date_cn(*start), "end": format_date_cn(*end)}
            for start, end in merged_ranges
        ],
        "monitor_date_text": format_date_ranges_text(merged_ranges),
        "monitor_dates": [format_date_cn(*item) for item in record_dates],
        "road_hint": infer_road_hint(plan_rows),
        "source_file": str(report_path or next(INPUT_DIR.glob("*.xlsx"), "")),
    }


def find_report_path() -> Optional[Path]:
    return next(
        (
            path for path in INPUT_DIR.glob("*.docx")
            if "报告" in path.name and not path.name.startswith("~$")
        ),
        None,
    )


def extract_monitoring_unit(chunks: List[Dict[str, Any]]) -> Optional[str]:
    for chunk in chunks:
        if chunk.get("kind") != "paragraphs":
            continue
        for line in str(chunk.get("text") or "").splitlines():
            candidate = line.strip()
            if LAB_NAME_RE.match(candidate):
                return candidate
    for chunk in chunks:
        if chunk.get("kind") != "table":
            continue
        text = str(chunk.get("text") or "")
        if "样品类型" not in text or "噪声" not in text:
            continue
        match = re.search(r"检测中心[\u4e00-\u9fffA-Za-z0-9（）()·\-]{0,40}", text)
        if match:
            return match.group(0)
    return None


def extract_report_noise_date_ranges(chunks: List[Dict[str, Any]]) -> List[Tuple[DATE_TUPLE, DATE_TUPLE]]:
    ranges: List[Tuple[DATE_TUPLE, DATE_TUPLE]] = []
    for chunk in chunks:
        if chunk.get("kind") != "table":
            continue
        text = str(chunk.get("text") or "")
        if "样品类型" not in text or "噪声" not in text:
            continue
        for match in DATE_RANGE_RE.finditer(text):
            start = (int(match.group(1)), int(match.group(2)), int(match.group(3)))
            end = (int(match.group(4)), int(match.group(5)), int(match.group(6)))
            ranges.append((start, end))
    return ranges


def collect_monitor_dates_from_records(records: List[Dict[str, Any]]) -> List[DATE_TUPLE]:
    dates: List[DATE_TUPLE] = []
    for record in records:
        parsed = parse_monitor_date(str(record.get("monitor_time") or ""))
        if parsed:
            dates.append(parsed)
    return sorted(set(dates))


def parse_monitor_date(text: str) -> Optional[DATE_TUPLE]:
    match = DATE_TOKEN_RE.search(str(text or ""))
    if not match:
        return None
    return int(match.group(1)), int(match.group(2)), int(match.group(3))


def date_tuples_to_ranges(dates: List[DATE_TUPLE]) -> List[Tuple[DATE_TUPLE, DATE_TUPLE]]:
    if not dates:
        return []
    ranges: List[Tuple[DATE_TUPLE, DATE_TUPLE]] = []
    start = end = dates[0]
    for current in dates[1:]:
        if date_ordinal(current) - date_ordinal(end) <= 1:
            end = current
            continue
        ranges.append((start, end))
        start = end = current
    ranges.append((start, end))
    return ranges


def merge_date_ranges(ranges: List[Tuple[DATE_TUPLE, DATE_TUPLE]]) -> List[Tuple[DATE_TUPLE, DATE_TUPLE]]:
    normalized = sorted(
        ((min(start, end), max(start, end)) for start, end in ranges),
        key=lambda item: date_ordinal(item[0]),
    )
    if not normalized:
        return []
    merged: List[Tuple[DATE_TUPLE, DATE_TUPLE]] = [normalized[0]]
    for start, end in normalized[1:]:
        last_start, last_end = merged[-1]
        if date_ordinal(start) <= date_ordinal(last_end) + 1:
            merged[-1] = (last_start, max([last_end, end], key=date_ordinal))
        else:
            merged.append((start, end))
    return merged


def format_date_cn(year: int, month: int, day: int) -> str:
    return f"{year}年{month}月{day}日"


def format_date_ranges_text(ranges: List[Tuple[DATE_TUPLE, DATE_TUPLE]]) -> str:
    parts: List[str] = []
    for start, end in ranges:
        if start == end:
            parts.append(format_date_cn(*start))
        else:
            parts.append(f"{format_date_cn(*start)}~{format_date_cn(*end)}")
    return "、".join(parts)


def date_ordinal(value: DATE_TUPLE) -> int:
    return date(value[0], value[1], value[2]).toordinal()


def infer_road_hint(plan_rows: List[Dict[str, Any]]) -> str:
    for row in plan_rows:
        if not row.get("is_attenuation"):
            continue
        text = " ".join(
            str(row.get(key) or "")
            for key in ("point_name", "position", "station", "point_code")
        )
        match = re.search(r"([\u4e00-\u9fff]{2,8}高速)", text)
        if match:
            return f"现状{match.group(1)}空旷地段"
    return ""


NOISE_POLISH_FIELD_BATCHES = (
    ("monitoring_plan_text", "monitoring_result_text"),
    ("sensitive_conclusion", "traffic_intro", "attenuation_conclusion"),
)


def _noise_polish_extra_rules(table_labels: Dict[str, str]) -> str:
    return (
        f"必须保留表号 {', '.join(table_labels.values())}。\n"
        "必须保留 summary 中所有超标点位编号和超标量。\n"
        "monitoring_result_text 必须保留 monitoring_meta 中的监测单位和监测日期，不得删改或编造。\n"
        "monitoring_plan_text 不得明显短于 style_exemplars 示例长度。\n"
        "若存在超标，不得写“均满足”或“均达标”。"
    )


def _call_noise_polish_batch(
    payload: Dict[str, Any],
    batch_keys: Tuple[str, ...],
    table_labels: Dict[str, str],
) -> Dict[str, str]:
    batch_rule_texts = {
        key: str(payload["rule_texts"].get(key, ""))
        for key in batch_keys
    }
    batch_payload = {**payload, "rule_texts": batch_rule_texts}
    polished = chat_completion_json_object_with_recovery(
        [
            {
                "role": "system",
                "content": "你是严谨的环评报告文本润色助手。必须输出严格 JSON，禁止编造事实。",
            },
            {
                "role": "user",
                "content": build_text_polish_prompt(
                    batch_payload,
                    role="声环境章节",
                    output_keys=", ".join(batch_keys),
                    extra_rules=_noise_polish_extra_rules(table_labels),
                ),
            },
        ],
        profile=LlmProfile.text_polish,
        label=f"noise_text_polish_{batch_keys[0]}",
    )
    return {
        key: str(polished.get(key, "")).strip() or batch_rule_texts[key]
        for key in batch_keys
    }


def polish_noise_text_with_llm(
    rule_texts: Dict[str, str],
    table1: Dict[str, Any],
    table2: Dict[str, Any],
    table3: Dict[str, Any],
    summary: Dict[str, Any],
    numbering: DocxNumbering,
) -> Dict[str, str]:
    table_labels = {
        table["table_key"]: numbering.table_label(table["table_key"])
        for table in (table1, table2, table3)
    }
    payload = {
        "enabled": ENABLE_LLM_TEXT_POLISH,
        "text_guidance": load_text_polish_guidance(TEXT_POLISH_GUIDANCE_PATH),
        "rule_texts": rule_texts,
        "summary": summary,
        "table_labels": table_labels,
        "monitoring_meta": summary.get("monitoring_meta") or {},
        "table_stats": {
            "table4_2_1_rows": len(table1["rows"]),
            "table4_2_2_rows": len(table2["rows"]),
            "table4_2_3_rows": len(table3["rows"]),
        },
    }
    write_json(DEBUG_DIR / "noise_llm_text_input.json", payload)

    if not ENABLE_LLM_TEXT_POLISH:
        validation = {"used_llm": False, "valid": True, "warnings": ["ENABLE_LLM_TEXT_POLISH=false"]}
        write_json(DEBUG_DIR / "noise_llm_text_output.json", {})
        write_json(DEBUG_DIR / "noise_llm_text_validation.json", validation)
        return rule_texts

    if not os.getenv("EIA_LLM_API_KEY"):
        validation = {"used_llm": False, "valid": False, "warnings": ["EIA_LLM_API_KEY is missing"]}
        write_json(DEBUG_DIR / "noise_llm_text_output.json", {})
        write_json(DEBUG_DIR / "noise_llm_text_validation.json", validation)
        return rule_texts

    try:
        polished_fields = dict(rule_texts)
        extra_rules = _noise_polish_extra_rules(table_labels)
        for batch_index, batch_keys in enumerate(NOISE_POLISH_FIELD_BATCHES):
            batch_result = _call_noise_polish_batch(payload, batch_keys, table_labels)
            polished_fields.update(batch_result)
            if batch_index + 1 < len(NOISE_POLISH_FIELD_BATCHES) and LLM_TEXT_BATCH_PAUSE_SECONDS > 0:
                time.sleep(LLM_TEXT_BATCH_PAUSE_SECONDS)
        merged = ensure_noise_table_refs(polished_fields, rule_texts, numbering)
        validation = validate_polished_text(merged, summary, numbering, rule_texts)
        write_json(DEBUG_DIR / "noise_llm_text_output.json", merged)
        write_json(DEBUG_DIR / "noise_llm_text_validation.json", validation)
        if validation["valid"]:
            return clean_polished_texts(merged, numbering)
        return rule_texts
    except Exception as exc:
        if is_network_error(exc):
            try:
                print("noise_text_polish retrying as single request after batch failure", flush=True)
                time.sleep(LLM_TEXT_BATCH_PAUSE_SECONDS)
                polished = chat_completion_json_object_with_recovery(
                    [
                        {
                            "role": "system",
                            "content": "你是严谨的环评报告文本润色助手。必须输出严格 JSON，禁止编造事实。",
                        },
                        {
                            "role": "user",
                            "content": build_text_polish_prompt(
                                payload,
                                role="声环境章节",
                                output_keys="monitoring_plan_text, monitoring_result_text, sensitive_conclusion, traffic_intro, attenuation_conclusion",
                                extra_rules=extra_rules,
                            ),
                        },
                    ],
                    profile=LlmProfile.text_polish,
                    label="noise_text_polish_single",
                    recovery_attempts=2,
                )
                polished_fields = {key: str(polished.get(key, "")).strip() for key in rule_texts}
                merged = ensure_noise_table_refs({**rule_texts, **polished_fields}, rule_texts, numbering)
                validation = validate_polished_text(merged, summary, numbering, rule_texts)
                write_json(DEBUG_DIR / "noise_llm_text_output.json", merged)
                write_json(DEBUG_DIR / "noise_llm_text_validation.json", validation)
                if validation["valid"]:
                    return clean_polished_texts(merged, numbering)
            except Exception:
                pass
        write_json(DEBUG_DIR / "noise_llm_text_output.json", {})
        write_json(
            DEBUG_DIR / "noise_llm_text_validation.json",
            build_rule_text_fallback_validation(exc),
        )
        return rule_texts


def clean_polished_texts(texts: Dict[str, str], numbering: DocxNumbering) -> Dict[str, str]:
    cleaned = dict(texts)
    sensitive_label = numbering.table_label(NOISE_TABLE_SENSITIVE)
    for prefix in (f"监测结果见{sensitive_label}。", f"监测结果见{sensitive_label}，"):
        value = cleaned.get("sensitive_conclusion", "")
        if value.startswith(prefix):
            cleaned["sensitive_conclusion"] = value[len(prefix):].lstrip()
    return cleaned


def ensure_noise_table_refs(
    texts: Dict[str, str],
    rule_texts: Dict[str, str],
    numbering: DocxNumbering,
) -> Dict[str, str]:
    result = dict(texts)
    field_table_keys = {
        "monitoring_plan_text": NOISE_TABLE_MONITOR_POINTS,
        "traffic_intro": NOISE_TABLE_ATTENUATION,
    }
    for field, table_key in field_table_keys.items():
        label = numbering.table_label(table_key)
        current = str(result.get(field, "")).strip()
        if label in current:
            continue
        rule = str(rule_texts.get(field, "")).strip()
        if label not in rule:
            continue
        match = re.search(rf"[^。]*{re.escape(label)}[^。]*。?", rule)
        if match and current:
            clause = match.group(0).strip()
            if not current.endswith("。"):
                current += "。"
            result[field] = current + clause
        else:
            result[field] = rule
    return result


def validate_polished_text(
    polished: Dict[str, Any],
    summary: Dict[str, Any],
    numbering: DocxNumbering,
    rule_texts: Dict[str, str],
) -> Dict[str, Any]:
    required = [
        "monitoring_plan_text",
        "monitoring_result_text",
        "sensitive_conclusion",
        "traffic_intro",
        "attenuation_conclusion",
    ]
    warnings: List[str] = []
    for key in required:
        if not isinstance(polished.get(key), str) or not polished.get(key, "").strip():
            warnings.append(f"missing or empty field: {key}")

    all_text = "\n".join(str(polished.get(key, "")) for key in required)
    table_field_checks = {
        NOISE_TABLE_MONITOR_POINTS: "monitoring_plan_text",
        NOISE_TABLE_ATTENUATION: "traffic_intro",
    }
    for table_key, field in table_field_checks.items():
        label = numbering.table_label(table_key)
        field_text = str(polished.get(field, ""))
        rule_text = str(rule_texts.get(field, ""))
        if label in rule_text and label not in field_text:
            warnings.append(f"missing table number: {label}")

    for item in required_exceed_checks(summary):
        point_code = str(item.get("point_code") or "")
        exceed_value = item.get("exceed")
        exceed_text = f"{exceed_value}dB(A)"
        if point_code and point_code not in all_text:
            warnings.append(f"missing exceed point: {point_code}")
        if not contains_exceed_amount(all_text, exceed_value):
            warnings.append(f"missing exceed amount: {exceed_text}")

    for key in ("sensitive", "attenuation"):
        if summary.get(key, {}).get("exceed_items"):
            target_field = "sensitive_conclusion" if key == "sensitive" else "attenuation_conclusion"
            text = str(polished.get(target_field, ""))
            if "均满足" in text or "均达标" in text:
                warnings.append(f"{target_field} says all compliant despite exceed items")

    monitoring_meta = summary.get("monitoring_meta") or {}
    result_text = str(polished.get("monitoring_result_text", ""))
    unit = str(monitoring_meta.get("monitoring_unit") or "").strip()
    if unit and unit not in result_text:
        warnings.append("monitoring_result_text missing monitoring unit from monitoring_meta")
    date_text = str(monitoring_meta.get("monitor_date_text") or "").strip()
    if date_text:
        years = sorted({str(item[0]) for item in re.findall(r"(\d{4})年", date_text)})
        if years and not any(year in result_text for year in years):
            warnings.append("monitoring_result_text missing monitoring dates from monitoring_meta")
    plan_text = str(polished.get("monitoring_plan_text", ""))
    if len(plan_text) < 150:
        warnings.append("monitoring_plan_text shorter than expected reference length")

    return {"used_llm": True, "valid": not warnings, "llm_applied": True, "warnings": warnings}


def required_exceed_checks(summary: Dict[str, Any]) -> List[Dict[str, Any]]:
    grouped: Dict[Tuple[str, str, str], Dict[str, Any]] = {}
    for section in ("sensitive", "attenuation"):
        for item in summary.get(section, {}).get("exceed_items", []):
            point_code = str(item.get("point_code") or "")
            point_name = str(item.get("point_name") or "")
            period = str(item.get("period") or "")
            key = (point_code, point_name, period)
            try:
                exceed = int(item.get("exceed"))
            except (TypeError, ValueError):
                exceed = 0
            current = grouped.get(key)
            if current is None or exceed > int(current.get("exceed") or 0):
                grouped[key] = {**item, "exceed": exceed}
    return list(grouped.values())


def contains_exceed_amount(text: str, exceed: Any) -> bool:
    try:
        value = int(exceed)
    except (TypeError, ValueError):
        return False
    normalized = (
        str(text or "")
        .replace(" ", "")
        .replace("（", "(")
        .replace("）", ")")
        .replace("－", "-")
        .replace("—", "-")
        .replace("~", "～")
    )
    exact = rf"{value}(?:\.0)?dB\(A\)"
    if re.search(exact, normalized, flags=re.IGNORECASE):
        return True
    range_pattern = r"(\d+(?:\.0)?)[～\-至到](\d+(?:\.0)?)dB\(A\)"
    for match in re.finditer(range_pattern, normalized, flags=re.IGNORECASE):
        start = int(float(match.group(1)))
        end = int(float(match.group(2)))
        if min(start, end) <= value <= max(start, end):
            return True
    return False


def build_docx(
    table1: Dict[str, Any],
    table2: Dict[str, Any],
    table3: Dict[str, Any],
    table_indoor: Dict[str, Any],
    summary: Dict[str, Any],
    texts: Dict[str, str],
    numbering: DocxNumbering,
) -> Document:
    doc = create_section_document()
    setup_document(doc)
    add_chapter_title(doc, numbering.section_title("noise", "声环境现状调查与评价"))
    add_section_heading(doc, numbering.next_level2_heading("noise", "监测方案"))
    add_body_paragraph(doc, texts["monitoring_plan_text"])
    add_caption(doc, numbering.table_caption(NOISE_TABLE_MONITOR_POINTS))
    add_monitor_points_table(doc, table1["headers"], table1["rows"])

    add_section_heading(doc, numbering.next_level2_heading("noise", "监测结果"))
    add_body_paragraph(doc, texts["monitoring_result_text"])
    add_level3_heading(doc, numbering.next_level3_heading("noise", "敏感点声环境质量现状"))
    add_landscape_section(doc)
    sensitive_tables = table2.get("subtables") or [table2]
    labels = [numbering.table_label(item["table_key"]) for item in sensitive_tables if item.get("rows")]
    if labels:
        add_body_paragraph(doc, "监测结果见" + "、".join(labels) + "。")
    for subtable in sensitive_tables:
        if not subtable.get("rows"):
            continue
        add_caption(doc, numbering.table_caption(subtable["table_key"]))
        add_result_table(doc, subtable["rows"], include_flow=True, flow_label=subtable.get("flow_label"))
    add_body_paragraph(doc, texts["sensitive_conclusion"])

    if table_indoor.get("rows"):
        indoor_label = numbering.table_label(NOISE_TABLE_INDOOR)
        add_body_paragraph(doc, f"室内噪声监测结果统计见{indoor_label}。")
        add_caption(doc, numbering.table_caption(NOISE_TABLE_INDOOR))
        add_result_table(doc, table_indoor["rows"], include_flow=True, flow_label=table_indoor.get("flow_label"))

    if table3.get("rows"):
        add_level3_heading(doc, numbering.next_level3_heading("noise", "交通噪声监测结果"))
        add_body_paragraph(doc, texts["traffic_intro"])
        add_caption(doc, numbering.table_caption(NOISE_TABLE_ATTENUATION))
        add_result_table(
            doc,
            table3["rows"],
            include_flow=True,
            flow_label=table3.get("flow_label"),
            merge_name_only=True,
        )
        add_body_paragraph(doc, texts["attenuation_conclusion"])
    return doc


def sensitive_conclusion(summary: Dict[str, Any]) -> str:
    if not summary["exceed_items"]:
        return "根据监测结果，项目沿线敏感点的昼间、夜间监测值均满足《声环境质量标准》相应标准限值。"
    return "根据监测结果，" + format_exceed_items(summary["exceed_items"]) + "；其他敏感点的昼间、夜间监测值满足《声环境质量标准》相应标准限值。"


def attenuation_conclusion(summary: Dict[str, Any], lead_only: bool = False) -> str:
    if lead_only:
        return ""
    if not summary["exceed_items"]:
        return "根据断面监测结果，现状公路交通噪声衰减断面昼间、夜间监测值均满足相应标准限值。"
    return "根据衰减断面监测结果，" + format_exceed_items(summary["exceed_items"]) + "。"


def format_exceed_items(items: List[Dict[str, Any]]) -> str:
    return "；".join(
        f"{item['point_name']}（{item['point_code']}）{item['period']}超标{item['exceed']}dB(A)"
        for item in items
    )


def configure_result_table_geometry(table: Any, vehicle_breakdown: bool) -> None:
    if not vehicle_breakdown:
        return
    table.autofit = False
    widths_cm = [1.2, 2.5, 2.8] + [0.78] * 10 + [0.86] * 12
    for column, width_cm in zip(table.columns, widths_cm):
        column.width = Cm(width_cm)


def compact_vehicle_table_text(table: Any, vehicle_breakdown: bool) -> None:
    if not vehicle_breakdown:
        return
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)


def add_result_table(
    doc: Document,
    rows: List[Dict[str, Any]],
    include_flow: bool,
    flow_label: Optional[str] = None,
    merge_name_only: bool = False,
) -> None:
    include_flow = include_flow and has_traffic_flow(rows)
    vehicle_breakdown = include_flow and has_vehicle_flow_breakdown(rows)
    header_row_count = 4 if vehicle_breakdown else 3
    headers = result_headers(include_flow=include_flow, vehicle_breakdown=vehicle_breakdown)
    table = doc.add_table(rows=header_row_count, cols=len(headers))
    table.style = TABLE_STYLE
    configure_result_table_geometry(table, vehicle_breakdown)
    if vehicle_breakdown:
        build_road_vehicle_result_table_header(table, flow_label or flow_label_from_rows(rows))
    else:
        build_result_table_header(table, include_flow=include_flow, flow_label=flow_label or flow_label_from_rows(rows))
    for row in rows:
        cells = table.add_row().cells
        for i, header in enumerate(headers):
            value = row.get(header, "-")
            cells[i].text = "-" if value == 0 and header.startswith("exceed_") else str(value)
    merge_blocks = merge_result_identity_cells(table, headers, rows, merge_name_only=merge_name_only, header_row_count=header_row_count)
    finalize_table(table, header_row_count=header_row_count)
    compact_vehicle_table_text(table, vehicle_breakdown)
    for column_index, start_row, end_row in merge_blocks:
        suppress_vertical_merge_inner_borders(table, column_index, start_row, end_row)
    doc.add_paragraph()


def merge_result_identity_cells(
    table: Any,
    headers: List[str],
    rows: List[Dict[str, Any]],
    merge_name_only: bool = False,
    header_row_count: int = 3,
) -> List[Tuple[int, int, int]]:
    if len(rows) <= 1:
        return []
    code_header = "监测点编号"
    name_header = "监测点名称"
    if code_header not in headers or name_header not in headers:
        return []
    code_col = headers.index(code_header)
    name_col = headers.index(name_header)
    data_start = header_row_count
    merge_blocks: List[Tuple[int, int, int]] = []
    start = 0
    while start < len(rows):
        code = str(rows[start].get(code_header) or "").strip()
        name = str(rows[start].get(name_header) or "").strip()
        end = start
        if merge_name_only:
            while (
                end + 1 < len(rows)
                and str(rows[end + 1].get(name_header) or "").strip() == name
            ):
                end += 1
        else:
            while (
                end + 1 < len(rows)
                and str(rows[end + 1].get(code_header) or "").strip() == code
                and str(rows[end + 1].get(name_header) or "").strip() == name
            ):
                end += 1
        if end > start and code and name:
            if not merge_name_only:
                code_cell = table.cell(data_start + start, code_col).merge(table.cell(data_start + end, code_col))
                code_cell.text = code
                merge_blocks.append((code_col, data_start + start, data_start + end))
            name_cell = table.cell(data_start + start, name_col).merge(table.cell(data_start + end, name_col))
            name_cell.text = name
            merge_blocks.append((name_col, data_start + start, data_start + end))
        start = end + 1
    return merge_blocks


def add_monitor_points_table(doc: Document, headers: List[str], rows: List[Dict[str, Any]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = TABLE_STYLE
    table.autofit = True
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        header_cells[index].text = str(header)
        if HEADER_FILL:
            shade_cell(header_cells[index], HEADER_FILL)
        set_cell_text_style(header_cells[index], bold=True)

    for row in rows:
        cells = table.add_row().cells
        for index, header in enumerate(headers):
            cells[index].text = str(row.get(header, "-"))
            set_cell_text_style(cells[index], bold=False)

    merge_blocks = []
    merge_blocks.extend(merge_consecutive_data_columns(table, headers, rows, ["监测点编号", "监测点名称"]))
    merge_blocks.extend(merge_repeated_data_columns(table, headers, rows, ["监测因子", "监测频次"]))
    finalize_table(table, header_row_count=1)
    for column_index, start_row, end_row in merge_blocks:
        suppress_vertical_merge_inner_borders(table, column_index, start_row, end_row)
    doc.add_paragraph()


def merge_consecutive_data_columns(
    table: Any,
    headers: List[str],
    rows: List[Dict[str, Any]],
    merge_headers: List[str],
) -> List[Tuple[int, int, int]]:
    if len(rows) <= 1:
        return []
    merge_blocks: List[Tuple[int, int, int]] = []
    for header in merge_headers:
        if header not in headers:
            continue
        column_index = headers.index(header)
        start = 0
        while start < len(rows):
            value = str(rows[start].get(header, "")).strip()
            end = start
            while end + 1 < len(rows) and str(rows[end + 1].get(header, "")).strip() == value:
                end += 1
            if end > start and value and value not in {"-", "/", "—"}:
                merged_cell = table.cell(1 + start, column_index).merge(table.cell(1 + end, column_index))
                merged_cell.text = value
                set_cell_text_style(merged_cell, bold=False)
                merge_blocks.append((column_index, 1 + start, 1 + end))
            start = end + 1
    return merge_blocks


def suppress_vertical_merge_inner_borders(table: Any, column_index: int, start_row: int, end_row: int) -> None:
    """Hide inner horizontal borders for a vertical merge block.

    Some Word/WPS renderers still draw table-level insideH borders through
    vertically merged cells. Setting cell-level top/bottom borders to nil makes
    the merge visually explicit.
    """
    if end_row <= start_row:
        return
    for row_index in range(start_row, end_row + 1):
        tc = table._tbl.tr_lst[row_index].tc_lst[column_index]
        tc_pr = tc.get_or_add_tcPr()
        borders = tc_pr.find(qn("w:tcBorders"))
        if borders is None:
            borders = OxmlElement("w:tcBorders")
            tc_pr.append(borders)
        if row_index > start_row:
            set_cell_border_nil(borders, "top")
        if row_index < end_row:
            set_cell_border_nil(borders, "bottom")


def set_cell_border_nil(borders: Any, edge: str) -> None:
    tag = qn(f"w:{edge}")
    element = borders.find(tag)
    if element is None:
        element = OxmlElement(f"w:{edge}")
        borders.append(element)
    element.set(qn("w:val"), "nil")


def merge_repeated_data_columns(
    table: Any,
    headers: List[str],
    rows: List[Dict[str, Any]],
    merge_headers: List[str],
) -> List[Tuple[int, int, int]]:
    if len(rows) <= 1:
        return []
    merge_blocks: List[Tuple[int, int, int]] = []
    for header in merge_headers:
        if header not in headers:
            continue
        values = [str(row.get(header, "")).strip() for row in rows]
        if not values or not values[0] or any(value != values[0] for value in values):
            continue
        column_index = headers.index(header)
        merged_cell = table.cell(1, column_index).merge(table.cell(len(rows), column_index))
        merged_cell.text = values[0]
        merge_blocks.append((column_index, 1, len(rows)))
    return merge_blocks


def build_road_vehicle_result_table_header(table: Any, flow_label: str) -> None:
    final_header_row = 3
    fixed_labels = (
        (0, "\u76d1\u6d4b\u70b9\u7f16\u53f7"),
        (1, "\u76d1\u6d4b\u70b9\u540d\u79f0"),
        (2, "\u76d1\u6d4b\u70b9\u4f4d\u7f6e"),
    )
    for column, label in fixed_labels:
        table.cell(0, column).merge(table.cell(final_header_row, column)).text = label

    table.cell(0, 3).merge(table.cell(0, 6)).text = "LAeq"
    table.cell(0, 7).merge(table.cell(1, 8)).text = "\u5e73\u5747\u503c"
    table.cell(0, 9).merge(table.cell(1, 10)).text = "\u6807\u51c6\u503c"
    table.cell(0, 11).merge(table.cell(1, 12)).text = "\u8d85\u6807\u91cf"

    table.cell(1, 3).merge(table.cell(1, 4)).text = "\u7b2c\u4e00\u5929"
    table.cell(1, 5).merge(table.cell(1, 6)).text = "\u7b2c\u4e8c\u5929"

    day_night_labels = {
        3: "\u663c",
        4: "\u591c",
        5: "\u663c",
        6: "\u591c",
        7: "\u663c",
        8: "\u591c",
        9: "\u663c",
        10: "\u591c",
        11: "\u663c",
        12: "\u591c",
    }
    for column, label in day_night_labels.items():
        table.cell(2, column).merge(table.cell(3, column)).text = label

    table.cell(0, 13).merge(table.cell(0, 24)).text = flow_label or "\u8f66\u6d41\u91cf\uff08\u8f86/20min\uff09"
    table.cell(1, 13).merge(table.cell(1, 18)).text = "\u7b2c\u4e00\u5929"
    table.cell(1, 19).merge(table.cell(1, 24)).text = "\u7b2c\u4e8c\u5929"
    for start_column, label in ((13, "\u663c"), (16, "\u591c"), (19, "\u663c"), (22, "\u591c")):
        table.cell(2, start_column).merge(table.cell(2, start_column + 2)).text = label
        for offset, vehicle in enumerate(ROAD_VEHICLE_KEYS):
            table.cell(3, start_column + offset).text = ROAD_VEHICLE_LABELS[vehicle]

    for row in table.rows[:4]:
        for cell in row.cells:
            if HEADER_FILL:
                shade_cell(cell, HEADER_FILL)


def build_result_table_header(table: Any, include_flow: bool = True, flow_label: str = "\u8f66\u6d41\u91cf\uff08\u8f86/20min\uff09") -> None:

    for col, text in [(0, "监测点编号"), (1, "监测点名称"), (2, "监测点位置")]:
        cell = table.cell(0, col).merge(table.cell(2, col))
        cell.text = text

    # Top-level groups.
    table.cell(0, 3).merge(table.cell(0, 6)).text = "LAeq"
    table.cell(0, 7).merge(table.cell(1, 8)).text = "平均值"
    table.cell(0, 9).merge(table.cell(1, 10)).text = "标准值"
    table.cell(0, 11).merge(table.cell(1, 12)).text = "超标量"

    # Second-level groups.
    table.cell(1, 3).merge(table.cell(1, 4)).text = "第一天"
    table.cell(1, 5).merge(table.cell(1, 6)).text = "第二天"

    day_night_labels = {
        3: "昼",
        4: "夜",
        5: "昼",
        6: "夜",
        7: "昼",
        8: "夜",
        9: "昼",
        10: "夜",
        11: "昼",
        12: "夜",
    }
    if include_flow:
        table.cell(0, 13).merge(table.cell(0, 16)).text = flow_label or "车流量（辆/20min）"
        table.cell(1, 13).merge(table.cell(1, 14)).text = "第一天"
        table.cell(1, 15).merge(table.cell(1, 16)).text = "第二天"
        day_night_labels.update(
            {
                13: "昼",
                14: "夜",
                15: "昼",
                16: "夜",
            }
        )

    for col, text in day_night_labels.items():
        table.cell(2, col).text = text

    for row in table.rows[:3]:
        for cell in row.cells:
            if HEADER_FILL:
                shade_cell(cell, HEADER_FILL)


def parse_table_rows(table_text: str) -> List[List[str]]:
    rows: List[List[str]] = []
    for line in table_text.splitlines():
        body = re.sub(r"^row\s+\d+:\s*", "", line.strip(), flags=re.IGNORECASE)
        cells = [
            match.group(2).strip()
            for match in re.finditer(r"col\s+(\d+):\s*(.*?)(?=\s*\|\s*col\s+\d+:|$)", body)
        ]
        if any(cells):
            rows.append(cells)
    return rows


def row_to_dict(headers: List[str], row: List[str]) -> Dict[str, Any]:
    return {headers[i]: row[i] if i < len(row) else "" for i in range(len(headers))}


def extract_point_code(text: str) -> Optional[str]:
    match = re.search(r"(?<![A-Za-z0-9])N(?:J)?\d+(?:-\d+){0,2}(?![A-Za-z0-9-])", text or "", flags=re.IGNORECASE)
    if not match:
        return None
    code = match.group(0).upper()
    if code.startswith("N") and not code.startswith("NJ"):
        return "NJ" + code[1:]
    return code


def is_valid_noise_point_code(code: Any) -> bool:
    return bool(re.fullmatch(r"NJ\d+(?:-\d+){0,2}", str(code or "").strip(), flags=re.IGNORECASE))


def extract_distance(text: str) -> Optional[str]:
    match = re.search(r"(\d+)\s*m", text or "", flags=re.IGNORECASE)
    return match.group(1) + "m" if match else None


def is_attenuation_text(text: str) -> bool:
    return any(item in str(text or "") for item in ["衰减断面", "距离中心线", "距离公路中心线", "距离宁徐高速公路中心线", "距离淮徐高速公路中心线"])


def infer_period(time_text: str) -> str:
    match = re.search(r"(\d{1,2}):(\d{2})", str(time_text or ""))
    if not match:
        return "unknown"
    hour = int(match.group(1))
    return "day" if 6 <= hour < 22 else "night"


def parse_decimal(value: Any) -> Optional[Decimal]:
    match = re.search(r"\d+(?:\.\d+)?", str(value or ""))
    if not match:
        return None
    return Decimal(match.group(0))


def parse_int(value: Any) -> Optional[int]:
    parsed = parse_decimal(value)
    return int(parsed.quantize(Decimal("1"), rounding=ROUND_HALF_UP)) if parsed is not None else None


def round_half_up(value: float) -> int:
    return int(Decimal(str(value)).quantize(Decimal("1"), rounding=ROUND_HALF_UP))


def value_or_dash(value: Any) -> Any:
    return "-" if value is None else format_decimal(value)


def quantize_noise_decimal(value: Decimal) -> Decimal:
    return value.quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)


def format_decimal(value: Any) -> Any:
    if value is None:
        return None
    decimal_value = value if isinstance(value, Decimal) else Decimal(str(value))
    text = format(decimal_value, "f")
    return text.rstrip("0").rstrip(".") if "." in text else text


def exceed_text(avg: Optional[Decimal], limit: Optional[Decimal]) -> Any:
    if avg is None or limit is None:
        return "-"
    rounded_avg = quantize_noise_decimal(avg)
    rounded_limit = quantize_noise_decimal(limit)
    return format_decimal(max(Decimal("0"), rounded_avg - rounded_limit))


def vehicle_flow_text(
    pairs: List[Dict[str, Dict[str, Any]]],
    day_index: int,
    period: str,
    vehicle: str,
) -> str:
    if day_index >= len(pairs):
        return "/"
    record = pairs[day_index].get(period)
    if not record:
        return "/"
    value = str(record.get(f"traffic_flow_{vehicle}") or "").strip()
    if value not in {"", "-", "/", "\u2014"}:
        return value
    breakdown = parse_road_flow_breakdown(record.get("traffic_flow"))
    return str(breakdown.get(vehicle) or "/")


def add_vehicle_flow_fields(row: Dict[str, Any], pairs: List[Dict[str, Dict[str, Any]]]) -> None:
    for day_number, day_index in ((1, 0), (2, 1)):
        for period in ("day", "night"):
            base = f"traffic_flow_day{day_number}_{period}"
            for vehicle in ROAD_VEHICLE_KEYS:
                row[f"{base}_{vehicle}"] = vehicle_flow_text(
                    pairs, day_index, period, vehicle
                )

def flow_text(pairs: List[Dict[str, Dict[str, Any]]], day_index: int, period: str) -> str:
    if day_index >= len(pairs):
        return "/"
    record = pairs[day_index].get(period)
    if not record:
        return "/"
    return str(record.get("traffic_flow") or "/").replace(" ", "/")


def point_sort_key(code: Any) -> Tuple[int, int, int, str]:
    text = str(code or "")
    match = re.search(r"N(?:J)?(\d+)(?:-(\d+))?(?:-(\d+))?", text, flags=re.IGNORECASE)
    if not match:
        return (9999, 9999, 9999, text)
    return (int(match.group(1)), int(match.group(2) or 0), int(match.group(3) or 0), text)


def write_json(path: Path, payload: Any) -> None:
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def read_json(path: Path) -> Any:
    return json.loads(path.read_text(encoding="utf-8"))


def load_noise_section_texts(output_dir: Path) -> Dict[str, str]:
    debug_dir = output_dir / "debug_tables"
    texts_path = debug_dir / "noise_section_texts.json"
    if texts_path.exists():
        payload = read_json(texts_path)
        if isinstance(payload, dict) and payload:
            return {key: str(value) for key, value in payload.items()}
    output_path = debug_dir / "noise_llm_text_output.json"
    if output_path.exists():
        payload = read_json(output_path)
        if isinstance(payload, dict) and payload:
            return {key: str(value) for key, value in payload.items()}
    input_path = debug_dir / "noise_llm_text_input.json"
    if input_path.exists():
        payload = read_json(input_path)
        rule_texts = payload.get("rule_texts") if isinstance(payload, dict) else None
        if isinstance(rule_texts, dict) and rule_texts:
            return {key: str(value) for key, value in rule_texts.items()}
    raise FileNotFoundError("未找到可重建声环境章节的文本 JSON")


def rebuild_docx_from_output(
    output_dir: Path,
    numbering: DocxNumbering,
    label_mapping: Optional[Dict[str, str]] = None,
) -> Path:
    output_dir = Path(output_dir)
    debug_dir = output_dir / "debug_tables"
    table1 = read_json(debug_dir / "noise_monitor_points_table.json")
    table2 = read_json(debug_dir / "noise_sensitive_points_result_table.json")
    indoor_path = debug_dir / "noise_indoor_result_table.json"
    table_indoor = read_json(indoor_path) if indoor_path.exists() else empty_indoor_result_table()
    table3 = read_json(debug_dir / "traffic_noise_attenuation_table.json")
    summary = read_json(debug_dir / "noise_compliance_summary.json")
    ensure_noise_table_metadata(table1, table2, table3, table_indoor)

    numbering.begin_section("noise", "声环境现状调查与评价")
    numbering.reset_section_heading_counters("noise")
    register_noise_tables(numbering, table1, table2, table3, table_indoor)

    texts = load_noise_section_texts(output_dir)
    if label_mapping:
        texts = numbering.remap_text_fields(texts, label_mapping)

    doc = build_docx(table1, table2, table3, table_indoor, summary, texts, numbering)
    finalize_section_document(doc)
    doc_path = output_dir / "noise_section.docx"
    doc.save(doc_path)
    return doc_path


def ensure_noise_table_metadata(
    table1: Dict[str, Any],
    table2: Dict[str, Any],
    table3: Dict[str, Any],
    table_indoor: Optional[Dict[str, Any]] = None,
) -> None:
    defaults = [
        (table1, NOISE_TABLE_MONITOR_POINTS, "声环境质量现状监测点"),
        (table2, NOISE_TABLE_SENSITIVE, "项目沿线敏感点声环境现状监测平均值  单位：dB(A)"),
        (table_indoor, NOISE_TABLE_INDOOR, "室内噪声监测结果统计  单位：dB(A)"),
        (table3, NOISE_TABLE_ATTENUATION, "现状公路交通噪声衰减断面监测结果  单位：dB(A)"),
    ]
    for table, table_key, caption_suffix in defaults:
        if table is None:
            continue
        table.setdefault("table_key", table_key)
        table.setdefault("caption_suffix", caption_suffix)
        table.setdefault("title", caption_suffix)


def empty_indoor_result_table() -> Dict[str, Any]:
    return {
        "table_key": NOISE_TABLE_INDOOR,
        "caption_suffix": "室内噪声监测结果统计  单位：dB(A)",
        "title": "室内噪声监测结果统计  单位：dB(A)",
        "headers": result_headers(include_flow=False),
        "rows": [],
        "warnings": [],
    }


if __name__ == "__main__":
    main()
