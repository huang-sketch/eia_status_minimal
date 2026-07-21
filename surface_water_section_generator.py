import json
import os
import re
import subprocess
import sys
import time
import urllib.error
import urllib.parse
import urllib.request
from collections import OrderedDict, defaultdict
from datetime import datetime
from html import unescape
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Tuple

from dotenv import load_dotenv
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from docx_layout import (
    HEADER_FILL,
    TABLE_STYLE,
    add_body_paragraph,
    add_caption,
    add_chapter_title,
    add_formula,
    add_landscape_section,
    add_level3_heading,
    add_portrait_section,
    add_section_heading,
    add_table,
    create_section_document,
    finalize_section_document,
    finalize_table,
    setup_document,
    set_cell_text_style,
    shade_cell,
    table_needs_landscape,
)
from docx_numbering import DocxNumbering, load_numbering
from formal_text_skill import build_surface_water_formal_text_validation, write_formal_text_validation
from llm_client import LlmProfile, SSL_CONTEXT, build_rule_text_fallback_validation, chat_completion_json_object_with_recovery
from table_schema_mapper import record_data_validation
from text_polish_utils import build_text_polish_prompt, load_text_polish_guidance

load_dotenv()

BASE_DIR = Path(__file__).resolve().parent
INPUT_DIR = Path(os.getenv("EIA_INPUT_DIR", "input"))
OUTPUT_DIR = Path(os.getenv("EIA_OUTPUT_DIR", "output"))
DEBUG_DIR = OUTPUT_DIR / "debug_tables"
SECTION_PATH = OUTPUT_DIR / "surface_water_section.docx"
ENABLE_LLM_TEXT_POLISH = os.getenv("ENABLE_LLM_TEXT_POLISH", "false").lower() == "true"
TEXT_POLISH_GUIDANCE_PATH = Path(
    os.getenv("EIA_SURFACE_WATER_TEXT_POLISH_GUIDANCE", BASE_DIR / "config" / "surface_water_text_polish_guidance.json")
)
LOCAL_STATUS_CACHE_PATH = Path(
    os.getenv("EIA_SURFACE_WATER_LOCAL_STATUS_CACHE", BASE_DIR / "config" / "surface_water_local_status_cache.json")
)
LOCAL_STATUS_WEB_SEARCH = os.getenv("EIA_SURFACE_WATER_WEB_SEARCH", "true").lower() != "false"
LOCAL_STATUS_WEB_TIMEOUT_SECONDS = int(os.getenv("EIA_SURFACE_WATER_WEB_TIMEOUT_SECONDS", "12"))

FACTOR_ALIASES = {
    "pH": "pH值",
    "PH": "pH值",
    "pH值": "pH值",
    "DO": "溶解氧",
    "溶解氧": "溶解氧",
    "TP": "总磷",
    "总磷": "总磷",
    "SS": "悬浮物",
    "悬浮物": "悬浮物",
    "CODMn": "高锰酸盐指数",
    "高锰酸盐指数": "高锰酸盐指数",
    "NH3-N": "氨氮",
    "氨氮": "氨氮",
}

NOT_APPLICABLE_FACTORS = {"水温", "悬浮物"}

ADMIN_LEVEL_WEIGHTS = {
    "county": 300,
    "county_city": 300,
    "district": 300,
    "city": 200,
    "province": 100,
    "unknown": 0,
}

SURFACE_TABLE_MONITOR_POINTS = "surface_monitor_points"
SURFACE_TABLE_MONITOR_RESULTS = "surface_monitor_results"
SURFACE_TABLE_COMPLIANCE = "surface_compliance"


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    DEBUG_DIR.mkdir(parents=True, exist_ok=True)
    ensure_inputs()

    monitoring_records = read_json(OUTPUT_DIR / "monitoring_records.json")
    compliance_results = read_json(OUTPUT_DIR / "compliance_results.json")
    standard_config = read_json(OUTPUT_DIR / "standard_config.json")

    surface_records = [
        record for record in monitoring_records
        if record.get("monitor_type") == "surface_water"
    ]
    surface_results = [
        result for result in compliance_results
        if result.get("monitor_type") == "surface_water"
    ]

    factor_list = detect_factors(surface_records)
    detected_payload = {
        "factor_list": factor_list,
        "source": "output/monitoring_records.json",
        "normalization": FACTOR_ALIASES,
    }
    write_json(DEBUG_DIR / "detected_factors.json", detected_payload)

    table1 = build_monitor_points_table(surface_records, standard_config, factor_list)
    table2 = build_monitor_results_table(surface_records, standard_config, factor_list)
    evaluated_factors = detect_evaluated_factors(surface_results)
    table3 = build_compliance_table(surface_results, standard_config, evaluated_factors)

    numbering = load_numbering(OUTPUT_DIR)
    numbering.begin_section("surface_water", "地表水环境现状调查与评价")
    register_surface_tables(numbering, table1, table2, table3)

    write_json(DEBUG_DIR / "surface_water_monitor_points_table.json", table1)
    write_json(DEBUG_DIR / "surface_water_monitor_results_table.json", table2)
    write_json(DEBUG_DIR / "surface_water_compliance_table.json", table3)

    project_meta = read_project_meta()
    local_status = build_local_water_status(project_meta.get("admin_division", ""))
    conclusion = build_conclusion(surface_results)
    rule_texts = build_rule_texts(table1, table2, factor_list, surface_results, conclusion, numbering)
    rule_texts["local_status_text"] = local_status.get("text", "")
    write_json(DEBUG_DIR / "surface_water_local_status.json", local_status)
    texts = polish_surface_water_text_with_llm(
        rule_texts,
        table1,
        table2,
        table3,
        factor_list,
        evaluated_factors,
        surface_results,
        numbering,
    )
    formal_validation = build_surface_water_formal_text_validation(
        table1,
        table2,
        table3,
        standard_config,
        surface_results,
        texts,
    )
    write_formal_text_validation(DEBUG_DIR, formal_validation)
    doc = build_docx(table1, table2, table3, factor_list, evaluated_factors, texts, numbering)
    finalize_section_document(doc)
    doc.save(SECTION_PATH)
    write_json(DEBUG_DIR / "surface_water_section_texts.json", texts)
    print(f"generated: {SECTION_PATH}")
    print(f"detected_factors: {len(factor_list)}")
    print(f"table1 rows: {len(table1['rows'])}")
    print(f"table2 rows: {len(table2['rows'])}")
    print(f"table3 rows: {len(table3['rows'])}")


def ensure_inputs() -> None:
    required = [
        OUTPUT_DIR / "monitoring_records.json",
        OUTPUT_DIR / "standard_config.json",
    ]
    compliance_path = OUTPUT_DIR / "compliance_results.json"
    if all(path.exists() for path in required) and compliance_path.exists():
        return

    engine = Path("compliance_engine.py")
    if engine.exists():
        subprocess.run([sys.executable, str(engine)], check=True)
    else:
        subprocess.run([sys.executable, "surface_water_pipeline.py"], check=True)


def read_json(path: Path) -> Any:
    return json.loads(path.read_text(encoding="utf-8"))


def read_project_meta() -> Dict[str, Any]:
    path = INPUT_DIR / "project_meta.json"
    if not path.exists():
        return {}
    try:
        payload = json.loads(path.read_text(encoding="utf-8"))
    except json.JSONDecodeError:
        return {}
    return payload if isinstance(payload, dict) else {}


def write_json(path: Path, payload: Any) -> None:
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def detect_factors(records: List[Dict[str, Any]]) -> List[str]:
    factors: List[str] = []
    seen: set[str] = set()
    for record in records:
        factor = normalize_factor(record.get("factor"))
        if not factor or factor in seen:
            continue
        seen.add(factor)
        factors.append(factor)
    return factors


def detect_evaluated_factors(results: List[Dict[str, Any]]) -> List[str]:
    factors: List[str] = []
    seen: set[str] = set()
    for result in results:
        factor = normalize_factor(result.get("factor"))
        method = result.get("method")
        not_applicable = bool(result.get("not_applicable")) or method == "not_applicable"
        if not factor or factor in seen or not_applicable or factor in NOT_APPLICABLE_FACTORS:
            continue
        seen.add(factor)
        factors.append(factor)
    return factors


def normalize_factor(value: Any) -> str:
    text = str(value or "").strip()
    return FACTOR_ALIASES.get(text, text)


def points_config(standard_config: Dict[str, Any]) -> Dict[str, Dict[str, Any]]:
    return standard_config.get("points") or {}


def build_monitor_points_table(
    records: List[Dict[str, Any]],
    standard_config: Dict[str, Any],
    factor_list: List[str],
) -> Dict[str, Any]:
    configs = points_config(standard_config)
    rows: List[Dict[str, Any]] = []
    for index, point_code in enumerate(sorted_point_codes(records, configs), start=1):
        config = configs.get(point_code, {})
        first_record = first_record_for_point(records, point_code)
        point_text = (first_record or {}).get("point") or ""
        parsed = parse_point_text(point_text)
        evidence = config.get("evidence") or {}
        dates = unique_ordered(
            record.get("sample_date")
            for record in records
            if record.get("point_code") == point_code
        )
        rows.append(
            {
                "序号": point_code,
                "河流名称": config.get("river_name") or parsed.get("river_name") or "-",
                "中心桩号": monitor_center_station(evidence, parsed),
                "取样断面": monitor_sampling_section(config, evidence, parsed),
                "取样频次": evidence.get("取样频次") or infer_frequency(dates),
                "监测因子": "、".join(factor_list),
                "point_code": point_code,
                "_point_order": index,
            }
        )
    headers = ["序号", "河流名称", "中心桩号", "取样断面", "取样频次", "监测因子"]
    warnings = hide_empty_monitor_point_columns(headers, rows)
    if warnings:
        write_json(DEBUG_DIR / "surface_water_monitor_points_warnings.json", warnings)
        record_data_validation(
            OUTPUT_DIR,
            "surface_water_monitor_points",
            valid=False,
            warnings=warnings,
        )
    return {
        "table_key": SURFACE_TABLE_MONITOR_POINTS,
        "caption_suffix": "水质监测断面布置",
        "title": "水质监测断面布置",
        "headers": headers,
        "rows": rows,
        "warnings": warnings,
    }


def monitor_sampling_section(config: Dict[str, Any], evidence: Dict[str, Any], parsed: Dict[str, str]) -> str:
    value = (
        config.get("section_name")
        or evidence.get("取样断面")
        or evidence.get("监测断面")
        or evidence.get("断面")
        or evidence.get("断面位置")
        or evidence.get("取样位置")
        or evidence.get("采样位置")
        or parsed.get("sampling_section")
        or ""
    )
    text = str(value or "").strip()
    return text or "-"


def hide_empty_monitor_point_columns(headers: List[str], rows: List[Dict[str, Any]]) -> List[str]:
    warnings: List[str] = []
    optional_columns = {
        "河流名称": "表3.1未显示河流名称列：监测方案和监测报告未提供河流名称。",
        "中心桩号": "表3.1未显示中心桩号列：监测方案和监测报告未提供中心桩号。",
    }
    for header, warning in optional_columns.items():
        if header not in headers:
            continue
        if all(is_blank_table_value(row.get(header)) for row in rows):
            headers.remove(header)
            for row in rows:
                row.pop(header, None)
            warnings.append(warning)
    return warnings


def is_blank_table_value(value: Any) -> bool:
    text = str(value or "").strip()
    return text in {"", "-", "/", "—", "无", "未提供", "None", "none", "null"}


def monitor_center_station(evidence: Dict[str, Any], parsed: Dict[str, str]) -> str:
    value = (
        evidence.get("中心桩号")
        or evidence.get("桩号")
        or evidence.get("监测点位置")
        or evidence.get("采样位置")
        or evidence.get("取样位置")
        or parsed.get("center_station")
        or ""
    )
    text = str(value or "").strip()
    if not text or "点位编号" in text:
        return "-"
    return text


def build_monitor_results_table(
    records: List[Dict[str, Any]],
    standard_config: Dict[str, Any],
    factor_list: List[str],
) -> Dict[str, Any]:
    configs = points_config(standard_config)
    point_orders = {code: index for index, code in enumerate(sorted_point_codes(records, configs), start=1)}
    grouped: Dict[Tuple[str, str], Dict[str, Any]] = {}
    for record in records:
        point_code = record.get("point_code")
        sample_date = record.get("sample_date")
        if not point_code or not sample_date:
            continue
        key = (point_code, sample_date)
        if key not in grouped:
            config = configs.get(point_code, {})
            river_name = config.get("river_name") or parse_point_text(record.get("point") or "").get("river_name") or "-"
            grouped[key] = {
                "序号": point_orders.get(point_code, len(point_orders) + 1),
                "河流": f"{river_name}（{point_code}）",
                "监测时间": sample_date,
                "point_code": point_code,
            }
            for factor in factor_list:
                grouped[key][factor] = "-"
        factor = normalize_factor(record.get("factor"))
        if factor in factor_list:
            grouped[key][factor] = record.get("value") or "-"
    headers = ["序号", "河流", "监测时间", *factor_list]
    return {
        "table_key": SURFACE_TABLE_MONITOR_RESULTS,
        "caption_suffix": "现状监测结果表",
        "title": "现状监测结果表",
        "headers": headers,
        "rows": [grouped[key] for key in sorted(grouped, key=point_date_sort_key)],
    }


def build_compliance_table(
    results: List[Dict[str, Any]],
    standard_config: Dict[str, Any],
    evaluated_factors: List[str],
) -> Dict[str, Any]:
    configs = points_config(standard_config)
    grouped: Dict[Tuple[str, str], Dict[str, Any]] = {}
    review_cells = defaultdict(set)
    for result in results:
        point_code = result.get("point_code")
        sample_date = result.get("sample_date")
        if not point_code or not sample_date:
            continue
        key = (point_code, sample_date)
        if key not in grouped:
            config = configs.get(point_code, {})
            river_name = result.get("river_name") or config.get("river_name") or "-"
            grouped[key] = {
                "编号": f"{river_name}（{point_code}）",
                "监测时间": sample_date,
                "point_code": point_code,
            }
            for factor in evaluated_factors:
                grouped[key][factor] = "-"
        factor = normalize_factor(result.get("factor"))
        if factor not in evaluated_factors:
            continue
        method = result.get("method")
        not_applicable = bool(result.get("not_applicable")) or method == "not_applicable"
        if not_applicable:
            grouped[key][factor] = "-"
        elif result.get("method") == "missing_standard":
            grouped[key][factor] = "缺少水质目标"
            review_cells[key].add(factor)
        elif result.get("needs_review"):
            grouped[key][factor] = "需复核"
            review_cells[key].add(factor)
        else:
            index = result.get("standard_index")
            grouped[key][factor] = format_index(index)
    headers = ["编号", "监测时间", *evaluated_factors]
    return {
        "table_key": SURFACE_TABLE_COMPLIANCE,
        "caption_suffix": "地表水环境现状评价结果",
        "title": "地表水环境现状评价结果",
        "headers": headers,
        "rows": [grouped[key] for key in sorted(grouped, key=point_date_sort_key)],
        "evaluated_factors": evaluated_factors,
        "review_cells": [
            {"point_code": key[0], "sample_date": key[1], "factor": factor}
            for key, factors in review_cells.items()
            for factor in sorted(factors)
        ],
    }


def build_conclusion(results: List[Dict[str, Any]]) -> str:
    missing_standard_items = [
        result for result in results
        if result.get("method") == "missing_standard"
    ]
    if missing_standard_items:
        points = sorted({str(result.get("point_code") or "") for result in missing_standard_items if result.get("point_code")})
        point_text = "、".join(points) if points else "相关监测断面"
        return (
            f"根据监测结果，{point_text}暂未识别到对应的水质目标或执行标准类别，"
            "评价结果表中相应因子列标注为“缺少水质目标”。"
            "需补充明确各监测断面的水质目标后，再按《地表水环境质量标准》（GB3838-2002）相应类别进行达标评价。"
        )
    exceed_items = [
        result for result in results
        if result.get("is_compliant") is False
    ]
    compliant_text = build_compliant_factor_summary(results)
    if not exceed_items:
        return (
            f"根据监测及评价结果，项目地表水各监测断面水质总体满足相应水质类别标准要求。"
            f"各监测点位的{compliant_text}均满足《地表水环境质量标准》（GB3838-2002）"
            "相应水质类别标准要求。"
        )

    exceed_text = build_exceed_summary(exceed_items)
    reason_text = build_surface_water_exceed_reason(exceed_items)
    return (
        f"根据现状监测及评价结果，项目地表水监测断面水质总体以达标为主，局部断面存在个别因子超标。"
        f"除超标因子外，各监测点位的{compliant_text}"
        "均满足《地表水环境质量标准》（GB3838-2002）相应水质类别标准要求。"
        f"超标情况为：{exceed_text}。{reason_text}"
    )


def build_compliant_factor_summary(results: List[Dict[str, Any]]) -> str:
    compliant_factors = unique_ordered(
        normalize_factor(result.get("factor"))
        for result in results
        if result.get("is_compliant") is True
    )
    if not compliant_factors:
        return "相关评价因子"
    return "、".join(compliant_factors) + "等评价因子"


def build_exceed_summary(exceed_items: List[Dict[str, Any]]) -> str:
    details = []
    for item in exceed_items:
        details.append(
            f"{item.get('river_name') or '-'}（{item.get('point_code') or '-'}）"
            f"{item.get('sample_date') or '-'} {normalize_factor(item.get('factor')) or '-'}"
        )
    return "；".join(details)


def build_surface_water_exceed_reason(exceed_items: List[Dict[str, Any]]) -> str:
    factors = unique_ordered(normalize_factor(item.get("factor")) for item in exceed_items)
    if not factors:
        return ""
    factor_text = "、".join(factors)
    return (
        f"超标原因可能与项目所在区域现状水体受上游来水、沿线生活及农业面源等综合影响有关，"
        f"其中{factor_text}指标在局部断面出现波动。"
    )


def build_compliant_summary(results: List[Dict[str, Any]]) -> str:
    grouped: "OrderedDict[str, List[str]]" = OrderedDict()
    for result in results:
        if result.get("is_compliant") is not True:
            continue
        key = f"{result.get('river_name') or '-'}（{result.get('point_code') or '-'}）{result.get('sample_date') or '-'}"
        factor = normalize_factor(result.get("factor"))
        if not factor:
            continue
        grouped.setdefault(key, [])
        if factor not in grouped[key]:
            grouped[key].append(factor)
    if not grouped:
        return "各监测断面相关评价因子"
    return "；".join(f"{key}的{'、'.join(factors)}" for key, factors in grouped.items()) + "等评价因子"


def register_surface_tables(
    numbering: DocxNumbering,
    table1: Dict[str, Any],
    table2: Dict[str, Any],
    table3: Dict[str, Any],
) -> None:
    for table in (table1, table2, table3):
        table["title"] = numbering.register_table(table["table_key"], table["caption_suffix"])


def build_rule_texts(
    table1: Dict[str, Any],
    table2: Dict[str, Any],
    factor_list: List[str],
    results: List[Dict[str, Any]],
    conclusion: str,
    numbering: DocxNumbering,
) -> Dict[str, str]:
    monitor_points_label = numbering.table_label(SURFACE_TABLE_MONITOR_POINTS)
    monitor_results_label = numbering.table_label(SURFACE_TABLE_MONITOR_RESULTS)
    compliance_label = numbering.table_label(SURFACE_TABLE_COMPLIANCE)
    return {
        "monitor_points_text": (
            f"根据项目所在区域的水文特征、河流水体规模，共计在评价范围设置了"
            f"{len(table1['rows'])}个监测断面进行水质监测。监测断面概况详见{monitor_points_label}。"
        ),
        "monitoring_time_method_text": build_monitoring_description(table1, factor_list),
        "monitoring_result_text": build_surface_water_result_intro(table2, numbering),
        "evaluation_result_intro": f"地表水监测点位环境现状评价结果见{compliance_label}。",
        "conclusion": conclusion,
    }


def build_surface_water_result_intro(table2: Dict[str, Any], numbering: DocxNumbering) -> str:
    dates = unique_ordered(row.get("监测时间") for row in table2.get("rows", []) if row.get("监测时间"))
    date_text = "、".join(dates)
    label = numbering.table_label(SURFACE_TABLE_MONITOR_RESULTS)
    if date_text:
        return f"本项目地表水于{date_text}开展现状监测，监测结果详见{label}。"
    return f"本项目地表水监测结果详见{label}。"


def build_local_water_status(admin_division: str) -> Dict[str, Any]:
    admin_text = str(admin_division or "").strip()
    started = time.time()
    cache = read_local_status_cache()
    matches: List[Dict[str, Any]] = []
    for item in cache.get("items", []):
        if not isinstance(item, dict):
            continue
        match = build_cache_match(admin_text, item)
        if match.get("matched"):
            matches.append(match)

    if matches:
        selected = sorted(
            matches,
            key=lambda match: cache_match_sort_key(match),
            reverse=True,
        )[0]
        selected_item = selected["item"]
        text = str(selected_item.get("text") or "").strip()
        match_candidates = [
            {
                "source_name": match["item"].get("source_name"),
                "year": match["item"].get("year"),
                "admin_level": match["item"].get("admin_level"),
                "matched_alias": match.get("matched_alias"),
                "matched_candidate": match.get("matched_candidate"),
                "score": match.get("score"),
                "score_detail": match.get("score_detail"),
            }
            for match in sorted(matches, key=lambda match: cache_match_sort_key(match), reverse=True)[:8]
        ]
        return {
            "admin_division": admin_text,
            "admin_candidates": extract_admin_search_candidates(admin_text),
            "status": "cache_hit",
            "cache_hit_used_directly": True,
            "year": selected_item.get("year"),
            "match_level": selected_item.get("admin_level"),
            "matched_alias": selected.get("matched_alias"),
            "matched_candidate": selected.get("matched_candidate"),
            "match_score": selected.get("score"),
            "score_detail": selected.get("score_detail"),
            "source_name": selected_item.get("source_name"),
            "source_url": selected_item.get("source_url"),
            "search_query": build_local_status_search_query(admin_text),
            "web_search": {
                "status": "skipped",
                "web_search_skipped_reason": "cache_hit_used_directly",
                "text": "",
            },
            "cache_candidates": match_candidates,
            "elapsed_ms": int((time.time() - started) * 1000),
            "text": text,
        }

    web_started = time.time()
    web_result = search_local_water_status_web(admin_text)
    web_result["elapsed_ms"] = int((time.time() - web_started) * 1000)
    if web_result.get("text"):
        web_result["cache_hit_used_directly"] = False
        web_result["elapsed_ms"] = int((time.time() - started) * 1000)
        return web_result

    payload = {
        "admin_division": admin_text,
        "admin_candidates": extract_admin_search_candidates(admin_text),
        "status": "missing_cache",
        "cache_hit_used_directly": False,
        "search_query": build_local_status_search_query(admin_text),
        "web_search": web_result,
        "elapsed_ms": int((time.time() - started) * 1000),
        "text": "",
        "warning": "联网检索和本地缓存均未得到可用公报文本，未生成区域地表水环境质量现状段落。",
    }
    return payload


def search_local_water_status_web(admin_division: str) -> Dict[str, Any]:
    admin_text = str(admin_division or "").strip()
    if not LOCAL_STATUS_WEB_SEARCH:
        return {
            "admin_division": admin_text,
            "status": "web_disabled",
            "search_query": build_local_status_search_query(admin_text),
            "text": "",
        }

    candidates = extract_admin_search_candidates(admin_text)
    current_year = datetime.now().year
    years = [current_year - 1, current_year - 2, current_year]
    searched: List[Dict[str, Any]] = []
    usable: List[Dict[str, Any]] = []

    for candidate in candidates:
        for year in years:
            query = f"{candidate} {year}年 生态环境状况公报 地表水 水质"
            try:
                results = search_bing(query)
            except Exception as exc:
                searched.append({"query": query, "error": str(exc)})
                continue
            searched.append({"query": query, "result_count": len(results)})
            for result in results[:5]:
                page_payload = build_local_status_from_search_result(
                    admin_text,
                    candidate,
                    year,
                    result,
                )
                if page_payload.get("text"):
                    usable.append(page_payload)

    if not usable:
        return {
            "admin_division": admin_text,
            "status": "web_no_usable_result",
            "search_query": build_local_status_search_query(admin_text),
            "searched": searched,
            "text": "",
        }

    selected = sorted(
        usable,
        key=lambda item: (
            int(item.get("match_score") or 0),
            int(item.get("year") or 0),
            int(item.get("source_score") or 0),
        ),
        reverse=True,
    )[0]
    selected["status"] = "web_hit"
    selected["searched"] = searched
    return selected


def search_bing(query: str) -> List[Dict[str, str]]:
    url = "https://www.bing.com/search?q=" + urllib.parse.quote(query)
    html_text = fetch_url_text(url)
    results: List[Dict[str, str]] = []
    for match in re.finditer(
        r'<li class="b_algo".*?<h2.*?<a\b[^>]*href="([^"]+)"[^>]*>(.*?)</a>.*?(?:<p>(.*?)</p>)?',
        html_text,
        flags=re.IGNORECASE | re.DOTALL,
    ):
        result_url = normalize_search_result_url(unescape(match.group(1)))
        title = clean_html_text(match.group(2))
        snippet = clean_html_text(match.group(3) or "")
        if result_url and title:
            results.append({"title": title, "url": result_url, "snippet": snippet})
    return results


def normalize_search_result_url(url: str) -> str:
    text = unescape(url or "")
    if "bing.com/ck/" not in text:
        return text
    parsed = urllib.parse.urlparse(text)
    params = urllib.parse.parse_qs(parsed.query)
    encoded = (params.get("u") or [""])[0]
    if encoded.startswith("a1"):
        encoded = encoded[2:]
    try:
        import base64

        padding = "=" * (-len(encoded) % 4)
        decoded = base64.urlsafe_b64decode((encoded + padding).encode("ascii")).decode("utf-8", errors="ignore")
        if decoded.startswith("http"):
            return decoded
    except Exception:
        pass
    return text


def build_local_status_from_search_result(
    admin_division: str,
    candidate: str,
    query_year: int,
    result: Dict[str, str],
) -> Dict[str, Any]:
    title = result.get("title", "")
    url = result.get("url", "")
    snippet = result.get("snippet", "")
    page_text = ""
    fetch_error = ""
    try:
        page_text = fetch_url_text(url)
    except Exception as exc:
        fetch_error = str(exc)

    clean_text = html_to_text(page_text) if page_text else snippet
    year = extract_year(title, snippet, clean_text) or query_year
    extracted = extract_water_status_sentences(clean_text or snippet)
    if not extracted:
        return {
            "admin_division": admin_division,
            "candidate": candidate,
            "year": year,
            "source_name": clean_source_name(title, candidate, year),
            "source_url": url,
            "fetch_error": fetch_error,
            "text": "",
        }

    source_name = clean_source_name(title, candidate, year)
    text = format_local_status_text(source_name, year, extracted)
    return {
        "admin_division": admin_division,
        "candidate": candidate,
        "year": year,
        "match_score": admin_candidate_score(admin_division, candidate),
        "source_score": source_url_score(url),
        "source_name": source_name,
        "source_url": url,
        "source_title": title,
        "source_snippet": snippet,
        "fetch_error": fetch_error,
        "text": text,
    }


def fetch_url_text(url: str) -> str:
    req = urllib.request.Request(
        url,
        headers={
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36",
            "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
        },
    )
    with urllib.request.urlopen(req, timeout=LOCAL_STATUS_WEB_TIMEOUT_SECONDS, context=SSL_CONTEXT) as response:
        raw = response.read(2_000_000)
    return raw.decode("utf-8", errors="ignore")


def clean_html_text(value: str) -> str:
    text = re.sub(r"<.*?>", "", value or "", flags=re.DOTALL)
    return unescape(re.sub(r"\s+", " ", text)).strip()


def html_to_text(html_text: str) -> str:
    text = re.sub(r"(?is)<script.*?</script>|<style.*?</style>", " ", html_text or "")
    text = re.sub(r"(?is)<br\s*/?>|</p>|</div>|</li>|</tr>|</h\d>", "\n", text)
    text = re.sub(r"(?is)<.*?>", " ", text)
    text = unescape(text)
    return re.sub(r"[ \t\r\f\v]+", " ", text)


def extract_water_status_sentences(text: str) -> List[str]:
    normalized = re.sub(r"\s+", "", text or "")
    parts = re.split(r"(?<=[。；;])", normalized)
    keywords = ("地表水", "水环境", "断面", "优Ⅲ", "优III", "饮用水", "水源", "水质", "国考", "省考")
    selected: List[str] = []
    for part in parts:
        if len(part) < 12 or len(part) > 180:
            continue
        if not any(keyword in part for keyword in keywords):
            continue
        if not re.search(r"\d|Ⅲ|III|Ⅱ|II|Ⅴ|V|优|良好|达标", part):
            continue
        if part not in selected:
            selected.append(part)
        if len(selected) >= 5:
            break
    return selected


def extract_year(*texts: str) -> Optional[int]:
    for text in texts:
        for match in re.finditer(r"(20\d{2})年?", text or ""):
            year = int(match.group(1))
            if 2000 <= year <= datetime.now().year:
                return year
    return None


def clean_source_name(title: str, candidate: str, year: int) -> str:
    title_text = re.split(r"[-_|—]", title or "")[0].strip()
    if "公报" in title_text:
        return title_text
    return f"{year}年{candidate}生态环境状况公报"


def format_local_status_text(source_name: str, year: int, sentences: List[str]) -> str:
    body = "".join(sentences).strip()
    body = re.sub(rf"^根据《?{re.escape(source_name)}》?[，,]?", "", body)
    body = re.sub(rf"^{year}年[，,]?", "", body)
    return f"根据《{source_name}》，{year}年，{body}"


def admin_candidate_score(admin_division: str, candidate: str) -> int:
    candidates = extract_admin_search_candidates(admin_division)
    try:
        parsed = parse_admin_division(admin_division)
        level_weight = ADMIN_LEVEL_WEIGHTS.get(parsed.get("candidate_levels", {}).get(candidate, ""), 0)
        return 100 - candidates.index(candidate) * 10 + level_weight + len(candidate)
    except ValueError:
        return len(candidate)


def source_url_score(url: str) -> int:
    host = urllib.parse.urlparse(url or "").netloc.lower()
    if host.endswith(".gov.cn") or ".gov.cn" in host:
        return 30
    if host.endswith(".cn"):
        return 10
    return 0


def build_cache_match(admin_division: str, item: Dict[str, Any]) -> Dict[str, Any]:
    admin_text = str(admin_division or "").strip()
    aliases = local_status_item_aliases(item)
    candidates = extract_admin_search_candidates(admin_text)
    candidate_levels = parse_admin_division(admin_text).get("candidate_levels", {})
    best: Optional[Dict[str, Any]] = None

    for alias in aliases:
        for index, candidate in enumerate(candidates):
            relation = admin_alias_relation(admin_text, candidate, alias)
            if not relation:
                continue
            item_level = normalize_admin_level(item.get("admin_level"))
            candidate_level = candidate_levels.get(candidate, infer_admin_level(candidate))
            score_detail = {
                "relation": relation,
                "candidate_order": index,
                "item_level": item_level,
                "candidate_level": candidate_level,
                "item_level_weight": ADMIN_LEVEL_WEIGHTS.get(item_level, 0),
                "candidate_level_weight": ADMIN_LEVEL_WEIGHTS.get(candidate_level, 0),
                "year": int(item.get("year") or 0),
                "alias_length": len(alias),
            }
            score = (
                relation_score(relation)
                + ADMIN_LEVEL_WEIGHTS.get(item_level, 0)
                + ADMIN_LEVEL_WEIGHTS.get(candidate_level, 0) // 4
                + max(0, 80 - index * 10)
                + len(alias)
            )
            current = {
                "matched": True,
                "item": item,
                "matched_alias": alias,
                "matched_candidate": candidate,
                "score": score,
                "score_detail": score_detail,
            }
            if best is None or cache_match_sort_key(current) > cache_match_sort_key(best):
                best = current

    if best:
        return best
    return {"matched": False, "item": item, "score": 0}


def cache_match_sort_key(match: Dict[str, Any]) -> Tuple[int, int, int, int]:
    item = match.get("item") if isinstance(match.get("item"), dict) else {}
    return (
        int(match.get("score") or 0),
        int(item.get("year") or 0),
        source_url_score(str(item.get("source_url") or "")),
        len(str(match.get("matched_alias") or "")),
    )


def local_status_item_aliases(item: Dict[str, Any]) -> List[str]:
    values: List[Any] = []
    level = normalize_admin_level(item.get("admin_level"))
    values.append(item.get("admin_division"))
    if level in {"county", "county_city", "district"}:
        values.append(item.get("county") or item.get("district") or item.get("county_city"))
    elif level == "city":
        values.append(item.get("city"))
    elif level == "province":
        values.append(item.get("province"))
    values.extend(item.get("admin_division_aliases") or [])
    values.extend(item.get("aliases") or [])
    values.extend(item.get("admin_codes") or [])
    return unique_ordered(values)


def admin_alias_relation(admin_division: str, candidate: str, alias: str) -> str:
    admin_text = normalize_admin_text(admin_division)
    candidate_text = normalize_admin_text(candidate)
    alias_text = normalize_admin_text(alias)
    if not alias_text:
        return ""
    if alias_text == admin_text:
        return "exact_admin"
    if alias_text == candidate_text:
        return "exact_candidate"
    if alias_text in admin_text:
        return "alias_in_admin"
    if admin_text and admin_text in alias_text:
        return "admin_in_alias"
    return ""


def relation_score(relation: str) -> int:
    return {
        "exact_admin": 1000,
        "exact_candidate": 900,
        "alias_in_admin": 700,
        "admin_in_alias": 650,
    }.get(relation, 0)


def admin_match_score(admin_division: str, item: Dict[str, Any]) -> int:
    return int(build_cache_match(admin_division, item).get("score") or 0)


def read_local_status_cache() -> Dict[str, Any]:
    if not LOCAL_STATUS_CACHE_PATH.exists():
        return {"items": []}
    try:
        payload = json.loads(LOCAL_STATUS_CACHE_PATH.read_text(encoding="utf-8"))
    except json.JSONDecodeError:
        return {"items": []}
    return payload if isinstance(payload, dict) else {"items": []}


def build_local_status_search_query(admin_division: str) -> str:
    admin_text = str(admin_division or "").strip()
    candidates = extract_admin_search_candidates(admin_text)
    return "；".join(f"{candidate} 最新 生态环境状况公报 地表水 水质" for candidate in candidates)


def extract_admin_search_candidates(admin_division: str) -> List[str]:
    admin_text = normalize_admin_text(admin_division)
    if not admin_text:
        return ["当地"]
    parsed = parse_admin_division(admin_text)
    candidates = [admin_text]
    for key in ("county", "city", "province"):
        value = parsed.get(key)
        if value and value not in candidates:
            candidates.append(value)
    return candidates


def parse_admin_division(admin_division: str) -> Dict[str, Any]:
    admin_text = normalize_admin_text(admin_division)
    province = regex_first(r"^(.+?(?:省|自治区))", admin_text)
    rest = admin_text[len(province):] if province else admin_text
    county = regex_first(r"([^省市县区]+?(?:县|区|旗))$", rest)
    city_part = rest[: -len(county)] if county else rest
    cities = re.findall(r"[^市]+?市|[^州]+?自治州|[^区]+?地区|[^盟]+?盟", city_part)
    city = ""
    if len(cities) >= 2:
        county = cities[-1]
        city = cities[-2]
    elif len(cities) == 1:
        city = cities[0]

    raw_level = infer_admin_level(county) if county else ("city" if city else ("province" if province else infer_admin_level(admin_text)))
    candidate_levels = {
        admin_text: raw_level,
    }
    if county:
        candidate_levels[county] = infer_admin_level(county)
    if city:
        candidate_levels[city] = "city"
    if province:
        candidate_levels[province] = "province"
    return {
        "raw": admin_text,
        "province": province,
        "city": city,
        "county": county,
        "candidate_levels": candidate_levels,
    }


def regex_first(pattern: str, text: str) -> str:
    match = re.search(pattern, text or "")
    return match.group(1) if match else ""


def normalize_admin_text(value: Any) -> str:
    text = str(value or "").strip()
    text = re.sub(r"\s+", "", text)
    return text


def infer_admin_level(value: Any) -> str:
    text = normalize_admin_text(value)
    if not text:
        return "unknown"
    if text.endswith(("县", "旗")):
        return "county"
    if text.endswith("区"):
        return "district"
    if text.endswith("市"):
        return "county_city" if len(text) <= 4 else "city"
    if text.endswith(("省", "自治区")):
        return "province"
    return "unknown"


def normalize_admin_level(value: Any) -> str:
    text = str(value or "").strip()
    aliases = {
        "county-level city": "county_city",
        "county_city": "county_city",
        "county": "county",
        "district": "district",
        "city": "city",
        "province": "province",
        "县级市": "county_city",
        "县": "county",
        "区": "district",
        "市": "city",
        "省": "province",
    }
    return aliases.get(text, infer_admin_level(text))


def build_docx(
    table1: Dict[str, Any],
    table2: Dict[str, Any],
    table3: Dict[str, Any],
    factor_list: List[str],
    evaluated_factors: List[str],
    texts: Dict[str, str],
    numbering: DocxNumbering,
) -> Document:
    doc = create_section_document()
    setup_document(doc)

    add_chapter_title(doc, numbering.section_title("surface_water", "地表水环境现状调查与评价"))
    local_status_text = texts.get("local_status_text", "").strip()
    if local_status_text:
        add_section_heading(doc, numbering.next_level2_heading("surface_water", "区域地表水环境质量现状"))
        add_body_paragraph(doc, local_status_text)

    add_section_heading(doc, numbering.next_level2_heading("surface_water", "水环境质量现状调查"))

    landscape_active = False

    add_level3_heading(doc, numbering.next_level3_heading("surface_water", "现状监测点布置"))
    add_body_paragraph(doc, texts.get("monitor_points_text") or "-")
    add_caption(doc, numbering.table_caption(SURFACE_TABLE_MONITOR_POINTS))
    add_monitor_points_table(doc, table1["headers"], table1["rows"])

    add_level3_heading(doc, numbering.next_level3_heading("surface_water", "监测时间、频率和方法"))
    add_body_paragraph(doc, texts.get("monitoring_time_method_text") or "-")

    add_level3_heading(doc, numbering.next_level3_heading("surface_water", "现状监测结果"))
    add_body_paragraph(doc, texts.get("monitoring_result_text") or "-")
    if table_needs_landscape(table2["headers"]):
        add_landscape_section(doc)
        landscape_active = True
    add_caption(doc, numbering.table_caption(SURFACE_TABLE_MONITOR_RESULTS))
    add_surface_water_grouped_table(doc, table2["headers"], table2["rows"], ["序号", "河流"])
    if landscape_active and not table_needs_landscape(table3["headers"]):
        add_portrait_section(doc)
        landscape_active = False

    add_level3_heading(doc, numbering.next_level3_heading("surface_water", "现状评价结果"))
    add_body_paragraph(doc, "①评价方法")
    add_evaluation_method_text(doc)
    add_body_paragraph(doc, "②评价结果")
    add_body_paragraph(doc, texts.get("evaluation_result_intro") or "-")
    if table_needs_landscape(table3["headers"]) and not landscape_active:
        add_landscape_section(doc)
        landscape_active = True
    add_caption(doc, numbering.table_caption(SURFACE_TABLE_COMPLIANCE))
    add_surface_water_grouped_table(doc, table3["headers"], table3["rows"], ["编号"])
    if landscape_active:
        add_portrait_section(doc)

    add_body_paragraph(doc, texts.get("conclusion") or "-")
    return doc


def add_monitor_points_table(doc: Document, headers: List[str], rows: List[Dict[str, Any]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = TABLE_STYLE
    table.autofit = True
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        header_cells[index].text = str(header)
        if HEADER_FILL:
            shade_cell(header_cells[index], HEADER_FILL)
        set_cell_text_style(header_cells[index], bold=True)

    for row in rows:
        cells = table.add_row().cells
        for index, header in enumerate(headers):
            cells[index].text = str(row.get(header, "-"))
            set_cell_text_style(cells[index], bold=False)

    merge_monitor_points_repeated_columns(table, headers, rows)
    finalize_table(table, header_row_count=1)
    doc.add_paragraph()


def add_surface_water_grouped_table(
    doc: Document,
    headers: List[str],
    rows: List[Dict[str, Any]],
    merge_headers: List[str],
) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = TABLE_STYLE
    table.autofit = True
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        header_cells[index].text = str(header)
        if HEADER_FILL:
            shade_cell(header_cells[index], HEADER_FILL)
        set_cell_text_style(header_cells[index], bold=True)

    for row in rows:
        cells = table.add_row().cells
        for index, header in enumerate(headers):
            cells[index].text = str(row.get(header, "-"))
            set_cell_text_style(cells[index], bold=False)

    merge_surface_water_consecutive_columns(table, headers, rows, merge_headers)
    finalize_table(table, header_row_count=1)
    doc.add_paragraph()


def merge_surface_water_consecutive_columns(
    table: Any,
    headers: List[str],
    rows: List[Dict[str, Any]],
    merge_headers: List[str],
) -> None:
    if len(rows) <= 1:
        return
    merge_blocks: List[Tuple[int, int, int]] = []
    start = 0
    while start < len(rows):
        point_code = normalize_merge_value(rows[start].get("point_code"))
        end = start
        while end + 1 < len(rows) and normalize_merge_value(rows[end + 1].get("point_code")) == point_code:
            end += 1
        if point_code and end > start:
            for header in merge_headers:
                if header not in headers:
                    continue
                column_index = headers.index(header)
                values = [normalize_merge_value(rows[index].get(header)) for index in range(start, end + 1)]
                if not values or not values[0] or any(value != values[0] for value in values):
                    continue
                merged_cell = table.cell(1 + start, column_index).merge(table.cell(1 + end, column_index))
                merged_cell.text = str(rows[start].get(header) or "-")
                set_cell_text_style(merged_cell, bold=False)
                merge_blocks.append((column_index, 1 + start, 1 + end))
        start = end + 1
    for column_index, start_row, end_row in merge_blocks:
        suppress_vertical_merge_inner_borders(table, column_index, start_row, end_row)


def suppress_vertical_merge_inner_borders(table: Any, column_index: int, start_row: int, end_row: int) -> None:
    if end_row <= start_row:
        return
    for row_index in range(start_row, end_row + 1):
        cell = table.cell(row_index, column_index)
        tc_pr = cell._tc.get_or_add_tcPr()
        borders = tc_pr.find(qn("w:tcBorders"))
        if borders is None:
            borders = OxmlElement("w:tcBorders")
            tc_pr.append(borders)
        for edge_name in ("top", "bottom"):
            if (row_index == start_row and edge_name == "top") or (row_index == end_row and edge_name == "bottom"):
                continue
            tag = qn(f"w:{edge_name}")
            border = borders.find(tag)
            if border is None:
                border = OxmlElement(f"w:{edge_name}")
                borders.append(border)
            border.set(qn("w:val"), "nil")


def merge_monitor_points_repeated_columns(table: Any, headers: List[str], rows: List[Dict[str, Any]]) -> None:
    if len(rows) <= 1:
        return
    for header in ("取样断面", "取样频次", "监测因子"):
        if header not in headers:
            continue
        values = [normalize_merge_value(row.get(header)) for row in rows]
        if not values or not values[0] or any(value != values[0] for value in values):
            continue
        column_index = headers.index(header)
        merged_cell = table.cell(1, column_index).merge(table.cell(len(rows), column_index))
        merged_cell.text = str(rows[0].get(header) or "-")
        set_cell_text_style(merged_cell, bold=False)


def normalize_merge_value(value: Any) -> str:
    return re.sub(r"\s+", "", str(value or "").strip())


def polish_surface_water_text_with_llm(
    rule_texts: Dict[str, str],
    table1: Dict[str, Any],
    table2: Dict[str, Any],
    table3: Dict[str, Any],
    factor_list: List[str],
    evaluated_factors: List[str],
    results: List[Dict[str, Any]],
    numbering: DocxNumbering,
) -> Dict[str, str]:
    table_labels = {
        table["table_key"]: numbering.table_label(table["table_key"])
        for table in (table1, table2, table3)
    }
    payload = {
        "enabled": ENABLE_LLM_TEXT_POLISH,
        "text_guidance": load_text_polish_guidance(TEXT_POLISH_GUIDANCE_PATH),
        "rule_texts": rule_texts,
        "table_labels": table_labels,
        "summary": build_surface_water_summary(table1, table2, factor_list, evaluated_factors, results),
        "table_stats": {
            "monitor_points": len(table1.get("rows", [])),
            "monitor_result_rows": len(table2.get("rows", [])),
            "compliance_rows": len(table3.get("rows", [])),
        },
    }
    write_json(DEBUG_DIR / "surface_water_llm_text_input.json", payload)

    if not ENABLE_LLM_TEXT_POLISH:
        validation = {"used_llm": False, "valid": True, "warnings": ["ENABLE_LLM_TEXT_POLISH=false"]}
        write_json(DEBUG_DIR / "surface_water_llm_text_output.json", {})
        write_json(DEBUG_DIR / "surface_water_llm_text_validation.json", validation)
        return rule_texts

    if not os.getenv("EIA_LLM_API_KEY"):
        validation = {"used_llm": False, "valid": False, "warnings": ["EIA_LLM_API_KEY is missing"]}
        write_json(DEBUG_DIR / "surface_water_llm_text_output.json", {})
        write_json(DEBUG_DIR / "surface_water_llm_text_validation.json", validation)
        return rule_texts

    try:
        polished = chat_completion_json_object_with_recovery(
            [
                {
                    "role": "system",
                    "content": "你只输出合法 JSON 对象，不输出解释。",
                },
                {
                    "role": "user",
                    "content": build_text_polish_prompt(
                        payload,
                        role="地表水环境章节",
                        output_keys="local_status_text, monitor_points_text, monitoring_time_method_text, monitoring_result_text, evaluation_result_intro, conclusion",
                    ),
                },
            ],
            profile=LlmProfile.text_polish,
            label="surface_water_text_polish",
        )
        polished_fields = {key: str(polished.get(key, "")).strip() for key in rule_texts}
        merged = ensure_surface_table_refs({**rule_texts, **polished_fields}, rule_texts, numbering)
        validation = validate_polished_text(merged, rule_texts, results, numbering)
        write_json(DEBUG_DIR / "surface_water_llm_text_output.json", merged)
        write_json(DEBUG_DIR / "surface_water_llm_text_validation.json", validation)
        if validation["valid"]:
            return merged
        return rule_texts
    except Exception as exc:
        write_json(DEBUG_DIR / "surface_water_llm_text_output.json", {})
        write_json(
            DEBUG_DIR / "surface_water_llm_text_validation.json",
            build_rule_text_fallback_validation(exc),
        )
        return rule_texts


def build_surface_water_summary(
    table1: Dict[str, Any],
    table2: Dict[str, Any],
    factor_list: List[str],
    evaluated_factors: List[str],
    results: List[Dict[str, Any]],
) -> Dict[str, Any]:
    exceed_items = [
        {
            "river_name": item.get("river_name"),
            "point_code": item.get("point_code"),
            "sample_date": item.get("sample_date"),
            "factor": normalize_factor(item.get("factor")),
            "standard_index": item.get("standard_index"),
        }
        for item in results
        if item.get("is_compliant") is False
    ]
    return {
        "monitor_point_count": len(table1.get("rows", [])),
        "monitor_dates": unique_ordered(row.get("监测时间") for row in table2.get("rows", [])),
        "factor_list": factor_list,
        "evaluated_factors": evaluated_factors,
        "exceed_items": exceed_items,
    }


def ensure_surface_table_refs(
    texts: Dict[str, str],
    rule_texts: Dict[str, str],
    numbering: DocxNumbering,
) -> Dict[str, str]:
    result = dict(texts)
    field_table_keys = {
        "monitor_points_text": SURFACE_TABLE_MONITOR_POINTS,
        "monitoring_result_text": SURFACE_TABLE_MONITOR_RESULTS,
        "evaluation_result_intro": SURFACE_TABLE_COMPLIANCE,
    }
    for field, table_key in field_table_keys.items():
        label = numbering.table_label(table_key)
        current = str(result.get(field, "")).strip()
        if label in current:
            continue
        rule = str(rule_texts.get(field, "")).strip()
        if label not in rule:
            continue
        match = re.search(rf"[^。]*{re.escape(label)}[^。]*。?", rule)
        if match and current:
            clause = match.group(0).strip()
            if not current.endswith("。"):
                current += "。"
            result[field] = current + clause
        else:
            result[field] = rule
    return result


def validate_polished_text(
    polished: Dict[str, Any],
    rule_texts: Dict[str, str],
    results: List[Dict[str, Any]],
    numbering: DocxNumbering,
) -> Dict[str, Any]:
    warnings: List[str] = []
    for key in rule_texts:
        if not isinstance(polished.get(key), str) or not polished.get(key, "").strip():
            warnings.append(f"missing or empty field: {key}")

    all_text = "\n".join(str(polished.get(key, "")) for key in rule_texts)
    table_field_checks = {
        SURFACE_TABLE_MONITOR_POINTS: "monitor_points_text",
        SURFACE_TABLE_MONITOR_RESULTS: "monitoring_result_text",
        SURFACE_TABLE_COMPLIANCE: "evaluation_result_intro",
    }
    for table_key, field in table_field_checks.items():
        label = numbering.table_label(table_key)
        field_text = str(polished.get(field, ""))
        rule_text = str(rule_texts.get(field, ""))
        if label in rule_text and label not in field_text:
            warnings.append(f"missing table reference: {label}")

    exceed_items = [item for item in results if item.get("is_compliant") is False]
    conclusion = str(polished.get("conclusion", ""))
    if exceed_items and has_unqualified_all_compliant_claim(conclusion):
        warnings.append("conclusion says all compliant despite exceed items")
    for item in exceed_items:
        point = str(item.get("point_code") or "")
        factor = normalize_factor(item.get("factor"))
        if point and point not in all_text:
            warnings.append(f"missing exceed point: {point}")
        if factor and factor not in all_text:
            warnings.append(f"missing exceed factor: {factor}")

    return {"used_llm": True, "valid": not warnings, "llm_applied": True, "warnings": warnings}


def has_unqualified_all_compliant_claim(conclusion: str) -> bool:
    sentences = [
        sentence.strip()
        for sentence in re.split(r"[。；;]", str(conclusion or ""))
        if sentence.strip()
    ]
    for sentence in sentences:
        if is_qualified_compliant_sentence(sentence):
            continue
        if re.search(r"(全部|全[部市县区]?[^\s，,。；;]{0,12}|所有|各监测点位|各断面|各监测断面).*?(均满足|均达标|均符合|均达到|全部达标)", sentence):
            return True
        if "均满足" in sentence or "均达标" in sentence or "全部达标" in sentence:
            return True
    return False


def is_qualified_compliant_sentence(sentence: str) -> bool:
    text = re.sub(r"\s+", "", sentence or "")
    if re.search(r"除[^。；;]*(超标|不达标|超过)[^。；;]*外", text):
        return True
    if re.search(r"(其余|其他|其余各项|其他评价因子|其余评价因子|非超标因子)", text):
        return True
    return False


def add_evaluation_method_text(doc: Document) -> None:
    add_body_paragraph(doc, "现状监测结果按水质指数法进行单项水质参数评价，计算公式如下：")
    add_formula(doc, "Sᵢ,ⱼ = Cᵢ,ⱼ / Cₛᵢ")
    add_body_paragraph(doc, "式中：Sᵢ,ⱼ——评价因子i的水质指数，大于1表明该水质因子超标；")
    add_body_paragraph(doc, "Cᵢ,ⱼ——评价因子i在j点的实测统计代表值，mg/L；")
    add_body_paragraph(doc, "Cₛᵢ——评价因子i的水质评价标准限值，mg/L。")
    add_body_paragraph(doc, "其中，pH的标准指数为：")
    add_formula(doc, "SₚH,ⱼ = (7.0 - pHⱼ) / (7.0 - pHsd)        （pHⱼ≤7.0）")
    add_formula(doc, "SₚH,ⱼ = (pHⱼ - 7.0) / (pHsu - 7.0)        （pHⱼ＞7.0）")
    add_body_paragraph(doc, "DO的标准指数为：")
    add_formula(doc, "DOsat = 468 / (31.6 + T)")
    add_formula(doc, "SᴅO,ⱼ = |DOf - DOⱼ| / (DOf - DOs)        （DOⱼ≥DOs）")
    add_formula(doc, "SᴅO,ⱼ = 10 - 9DOⱼ / DOs        （DOⱼ＜DOs）")
    add_body_paragraph(doc, "式中：SₚH,ⱼ——pH值的指数，大于1表明该水质因子超标；")
    add_body_paragraph(doc, "pHⱼ——pH值实测统计代表值；")
    add_body_paragraph(doc, "pHsd——评价标准中pH值的下限值；")
    add_body_paragraph(doc, "pHsu——评价标准中pH值的上限值；")
    add_body_paragraph(doc, "SᴅO,ⱼ——溶解氧的标准指数，大于1表明该水质因子超标；")
    add_body_paragraph(doc, "DOⱼ——溶解氧在j点的实测统计代表值，mg/L；")
    add_body_paragraph(doc, "DOs——溶解氧的水质评价标准限值，mg/L；")
    add_body_paragraph(doc, "DOf——饱和溶解氧浓度，mg/L；")
    add_body_paragraph(doc, "T——水温，℃。")
    add_body_paragraph(doc, "水温、悬浮物等无相应限值的指标仅列出现状监测值，不参与标准指数评价。")


def build_monitoring_description(table1: Dict[str, Any], factor_list: List[str]) -> str:
    frequencies = unique_ordered(row.get("取样频次") for row in table1["rows"] if row.get("取样频次"))
    frequency_text = "；".join(frequencies) if frequencies else "-"
    factors = "、".join(factor_list)
    return (
        f"本次地表水现状监测断面按监测方案布设，取样频次为{frequency_text}。"
        f"监测因子包括{factors}。各监测因子检测方法按监测报告及相关水质检测标准执行。"
    )


def sorted_point_codes(
    records: List[Dict[str, Any]],
    configs: Dict[str, Dict[str, Any]],
) -> List[str]:
    codes = unique_ordered([*configs.keys(), *(record.get("point_code") for record in records)])
    return sorted(codes, key=point_sort_key)


def point_sort_key(code: str) -> Tuple[int, str]:
    digits = "".join(ch for ch in str(code) if ch.isdigit())
    return (int(digits) if digits else 9999, str(code))


def point_date_sort_key(key: Tuple[str, str]) -> Tuple[int, str, Tuple[int, ...], str]:
    point_code, sample_date = key
    return (*point_sort_key(point_code), date_sort_key(sample_date), str(sample_date))


def date_sort_key(value: Any) -> Tuple[int, ...]:
    numbers = [int(item) for item in re_find_numbers(str(value or ""))]
    return tuple(numbers) if numbers else (9999,)


def re_find_numbers(text: str) -> List[str]:
    import re

    return re.findall(r"\d+", text)


def first_record_for_point(records: List[Dict[str, Any]], point_code: str) -> Optional[Dict[str, Any]]:
    return next((record for record in records if record.get("point_code") == point_code), None)


def parse_point_text(text: str) -> Dict[str, Optional[str]]:
    parts = str(text or "").split()
    result: Dict[str, Optional[str]] = {
        "center_station": None,
        "river_name": None,
        "sampling_section": None,
    }
    if len(parts) >= 2:
        result["center_station"] = parts[1]
    if len(parts) >= 3:
        result["river_name"] = parts[2]
    if len(parts) >= 4:
        section_parts = []
        for part in parts[3:]:
            if part.startswith("点位编号"):
                break
            section_parts.append(part)
        result["sampling_section"] = "".join(section_parts) if section_parts else None
    return result


def infer_frequency(dates: List[str]) -> str:
    if len(dates) == 3:
        return "连续取样三天，每天一次"
    if len(dates) == 1:
        return "取样一天，一天一次"
    if dates:
        return f"共取样{len(dates)}天"
    return "-"


def unique_ordered(items: Iterable[Any]) -> List[str]:
    result: List[str] = []
    seen: set[str] = set()
    for item in items:
        text = str(item or "").strip()
        if not text or text in seen:
            continue
        seen.add(text)
        result.append(text)
    return result


def format_index(value: Any) -> str:
    if value is None:
        return "-"
    try:
        return f"{float(value):.2f}"
    except (TypeError, ValueError):
        return "需复核"


def load_surface_section_texts(output_dir: Path) -> Dict[str, str]:
    debug_dir = output_dir / "debug_tables"
    texts_path = debug_dir / "surface_water_section_texts.json"
    if texts_path.exists():
        payload = read_json(texts_path)
        if isinstance(payload, dict) and payload:
            return {key: str(value) for key, value in payload.items()}
    output_path = debug_dir / "surface_water_llm_text_output.json"
    if output_path.exists():
        payload = read_json(output_path)
        if isinstance(payload, dict) and payload:
            texts = {key: str(value) for key, value in payload.items()}
            local_status = read_json(debug_dir / "surface_water_local_status.json")
            if isinstance(local_status, dict):
                texts.setdefault("local_status_text", str(local_status.get("text") or ""))
            return texts
    input_path = debug_dir / "surface_water_llm_text_input.json"
    if input_path.exists():
        payload = read_json(input_path)
        rule_texts = payload.get("rule_texts") if isinstance(payload, dict) else None
        if isinstance(rule_texts, dict) and rule_texts:
            texts = {key: str(value) for key, value in rule_texts.items()}
            local_status = read_json(debug_dir / "surface_water_local_status.json")
            if isinstance(local_status, dict):
                texts["local_status_text"] = str(local_status.get("text") or "")
            return texts
    raise FileNotFoundError("未找到可重建地表水章节的文本 JSON")


def rebuild_docx_from_output(
    output_dir: Path,
    numbering: DocxNumbering,
    label_mapping: Optional[Dict[str, str]] = None,
) -> Path:
    output_dir = Path(output_dir)
    debug_dir = output_dir / "debug_tables"
    table1 = read_json(debug_dir / "surface_water_monitor_points_table.json")
    table2 = read_json(debug_dir / "surface_water_monitor_results_table.json")
    table3 = read_json(debug_dir / "surface_water_compliance_table.json")
    ensure_surface_table_metadata(table1, table2, table3)
    detected = read_json(debug_dir / "detected_factors.json")
    factor_list = detected.get("factor_list") if isinstance(detected, dict) else []
    if not isinstance(factor_list, list):
        factor_list = []
    evaluated_factors = [
        header
        for header in table3.get("headers", [])[2:]
        if str(header).strip()
    ]

    numbering.begin_section("surface_water", "地表水环境现状调查与评价")
    numbering.reset_section_heading_counters("surface_water")
    register_surface_tables(numbering, table1, table2, table3)

    texts = load_surface_section_texts(output_dir)
    if label_mapping:
        texts = numbering.remap_text_fields(texts, label_mapping)

    doc = build_docx(table1, table2, table3, factor_list, evaluated_factors, texts, numbering)
    finalize_section_document(doc)
    doc_path = output_dir / "surface_water_section.docx"
    doc.save(doc_path)
    return doc_path


def ensure_surface_table_metadata(
    table1: Dict[str, Any],
    table2: Dict[str, Any],
    table3: Dict[str, Any],
) -> None:
    defaults = [
        (table1, SURFACE_TABLE_MONITOR_POINTS, "水质监测断面布置"),
        (table2, SURFACE_TABLE_MONITOR_RESULTS, "现状监测结果表"),
        (table3, SURFACE_TABLE_COMPLIANCE, "地表水环境现状评价结果"),
    ]
    for table, table_key, caption_suffix in defaults:
        table.setdefault("table_key", table_key)
        table.setdefault("caption_suffix", caption_suffix)
        table.setdefault("title", caption_suffix)


if __name__ == "__main__":
    main()
